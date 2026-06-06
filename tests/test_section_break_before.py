"""
Unit tests for the section_break_before operation.

section_break_before moves a section break that currently *follows* a heading
to immediately *before* it, so the heading starts its own page/section instead
of being stranded at the tail of the previous (e.g. landscape) section.
"""
import copy
import json
import tempfile

from docx import Document
from docx.oxml.ns import qn

from src.document_processor import DocxBulkUpdater
from src.config import load_operations_from_json


def _make_doc_with_stranded_heading():
    """Heading 'O2 RAW DATA' stranded inside a landscape section whose break
    sits *after* it (mirrors the FTIR template bug)."""
    doc = Document()
    doc.add_paragraph("FTIR RAW DATA")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Spectrum"
    doc.add_paragraph("O2 RAW DATA")

    # Empty paragraph carrying a landscape section break (the stranding break)
    breaker = doc.add_paragraph("")
    land = copy.deepcopy(doc.sections[-1]._sectPr)
    pg = land.find(qn('w:pgSz'))
    w, h = pg.get(qn('w:w')), pg.get(qn('w:h'))
    pg.set(qn('w:w'), h)
    pg.set(qn('w:h'), w)
    pg.set(qn('w:orient'), 'landscape')
    breaker._p.get_or_add_pPr().append(land)

    doc.add_paragraph("after O2 tables")
    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)
    return path


def _para_by_text(doc, text):
    for el in doc.element.body:
        if el.tag == qn('w:p'):
            t = "".join(x.text or "" for x in el.iter(qn('w:t'))).strip()
            if t == text:
                return el
    return None


class TestSectionBreakBefore:

    def test_moves_break_before_heading(self):
        path = _make_doc_with_stranded_heading()
        processor = DocxBulkUpdater([{"op": "section_break_before", "match": "O2 RAW DATA"}])
        assert processor.modify_docx(path) is True

        doc = Document(path)
        heading = _para_by_text(doc, "O2 RAW DATA")
        assert heading is not None

        # The paragraph immediately before the heading now carries the landscape break
        prev = heading.getprevious()
        assert prev is not None and prev.tag == qn('w:p')
        sectPr = prev.find(qn('w:pPr')).find(qn('w:sectPr'))
        assert sectPr is not None
        assert sectPr.find(qn('w:pgSz')).get(qn('w:orient')) == 'landscape'

        # No paragraph-level section break remains after the heading
        el = heading.getnext()
        while el is not None:
            if el.tag == qn('w:p'):
                pPr = el.find(qn('w:pPr'))
                assert pPr is None or pPr.find(qn('w:sectPr')) is None
            el = el.getnext()

    def test_section_count_unchanged(self):
        path = _make_doc_with_stranded_heading()
        before = len(Document(path).sections)
        DocxBulkUpdater([{"op": "section_break_before", "match": "O2 RAW DATA"}]).modify_docx(path)
        assert len(Document(path).sections) == before  # break moved, not added

    def test_idempotent(self):
        path = _make_doc_with_stranded_heading()
        processor = DocxBulkUpdater([{"op": "section_break_before", "match": "O2 RAW DATA"}])
        assert processor.modify_docx(path) is True
        assert processor.modify_docx(path) is False

    def test_no_match_returns_false(self):
        path = _make_doc_with_stranded_heading()
        processor = DocxBulkUpdater([{"op": "section_break_before", "match": "NOPE"}])
        assert processor.modify_docx(path) is False

    def test_exact_match_preferred_over_substring(self):
        # An appendix-list entry contains the heading text as a substring;
        # the exact heading must be the one that gets the break.
        doc = Document()
        doc.add_paragraph("APPENDIX F O2 RAW DATA")  # substring, earlier in doc
        doc.add_table(rows=1, cols=2).rows[0].cells[0].text = "Spectrum"
        doc.add_paragraph("O2 RAW DATA")             # exact heading
        breaker = doc.add_paragraph("")
        land = copy.deepcopy(doc.sections[-1]._sectPr)
        pg = land.find(qn('w:pgSz'))
        pg.set(qn('w:orient'), 'landscape')
        breaker._p.get_or_add_pPr().append(land)
        doc.add_paragraph("tail")
        path = tempfile.mktemp(suffix=".docx")
        doc.save(path)

        DocxBulkUpdater([{"op": "section_break_before", "match": "O2 RAW DATA"}]).modify_docx(path)
        doc = Document(path)
        heading = _para_by_text(doc, "O2 RAW DATA")
        prev = heading.getprevious()
        assert prev.find(qn('w:pPr')).find(qn('w:sectPr')) is not None
        # The appendix-list paragraph was not touched
        appendix = _para_by_text(doc, "APPENDIX F O2 RAW DATA")
        ap_prev = appendix.getprevious()
        assert ap_prev is None or ap_prev.find(qn('w:pPr')) is None or \
            ap_prev.find(qn('w:pPr')).find(qn('w:sectPr')) is None

    def test_config_loading(self, tmp_path):
        config = tmp_path / "cfg.json"
        config.write_text(json.dumps({"section_break_before": {"match": "O2 RAW DATA"}}))
        operations, _ = load_operations_from_json(config)
        assert operations[0]["op"] == "section_break_before"
        assert operations[0]["match"] == "O2 RAW DATA"
