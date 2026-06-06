"""
Unit tests for the divider operation.

divider isolates a matched paragraph on its own vertically-centered page
*without touching the paragraph itself*: it inserts a section break after the
paragraph and sets <w:vAlign w:val="center"/> on the divider's own section.
The paragraph keeps whatever horizontal alignment its style already gives it
(forcing jc/indent on the numbered heading was what broke its centering).
"""
import copy
import json
import tempfile

from docx import Document
from docx.oxml.ns import qn

from src.document_processor import DocxBulkUpdater
from src.config import load_operations_from_json


def _make_sect_break(doc, paragraph):
    """Turn an existing empty paragraph into a section-break paragraph by
    appending a copy of the body-final sectPr to its pPr."""
    final = doc.element.body.find(qn('w:sectPr'))
    paragraph._p.get_or_add_pPr().append(copy.deepcopy(final))


def _make_doc():
    """A divider heading 'O2 RAW DATA' followed by the rawdata table."""
    doc = Document()
    doc.add_paragraph("FTIR RAW DATA")
    doc.add_paragraph("O2 RAW DATA")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Spec"
    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)
    return path


def _para_by_text(doc, text):
    for el in doc.element.body:
        if el.tag == qn('w:p'):
            t = "".join(x.text or "" for x in el.iter(qn('w:t'))).strip()
            if t == text:
                return el
    return None


class TestDivider:

    def test_inserts_vertically_centered_section_after(self):
        path = _make_doc()
        assert DocxBulkUpdater([{"op": "divider", "match": "O2 RAW DATA"}]).modify_docx(path) is True
        doc = Document(path)
        heading = _para_by_text(doc, "O2 RAW DATA")

        # The paragraph immediately after the divider now carries a sectPr
        # (the section break) and that section is vertically centered.
        nxt = heading.getnext()
        assert nxt is not None and nxt.tag == qn('w:p')
        sectPr = nxt.find(qn('w:pPr')).find(qn('w:sectPr'))
        assert sectPr is not None
        assert sectPr.find(qn('w:vAlign')).get(qn('w:val')) == 'center'

    def test_paragraph_is_left_untouched(self):
        # The divider must NOT add jc/ind to the heading paragraph; it relies
        # on the paragraph's existing style for horizontal alignment.
        path = _make_doc()
        DocxBulkUpdater([{"op": "divider", "match": "O2 RAW DATA"}]).modify_docx(path)
        doc = Document(path)
        heading = _para_by_text(doc, "O2 RAW DATA")
        pPr = heading.find(qn('w:pPr'))
        if pPr is not None:
            assert pPr.find(qn('w:jc')) is None
            assert pPr.find(qn('w:ind')) is None

    def test_section_count_increases_by_one(self):
        path = _make_doc()
        before = len(Document(path).sections)
        DocxBulkUpdater([{"op": "divider", "match": "O2 RAW DATA"}]).modify_docx(path)
        assert len(Document(path).sections) == before + 1  # divider gets its own section

    def test_no_page_break_run_added(self):
        # The divider isolates via a section break, not a w:br page break.
        path = _make_doc()
        DocxBulkUpdater([{"op": "divider", "match": "O2 RAW DATA"}]).modify_docx(path)
        doc = Document(path)
        heading = _para_by_text(doc, "O2 RAW DATA")
        assert not any(b.get(qn('w:type')) == 'page' for b in heading.iter(qn('w:br')))

    def test_idempotent(self):
        path = _make_doc()
        processor = DocxBulkUpdater([{"op": "divider", "match": "O2 RAW DATA"}])
        assert processor.modify_docx(path) is True
        assert processor.modify_docx(path) is False

        # Still exactly one inserted section break after the divider.
        doc = Document(path)
        heading = _para_by_text(doc, "O2 RAW DATA")
        nxt = heading.getnext()
        nxt2 = nxt.getnext()
        assert nxt.find(qn('w:pPr')).find(qn('w:sectPr')) is not None
        if nxt2 is not None and nxt2.tag == qn('w:p'):
            p2 = nxt2.find(qn('w:pPr'))
            assert p2 is None or p2.find(qn('w:sectPr')) is None

    def test_collapses_breaks_around_vanishing_block(self):
        # Mirrors the real templates: divider, a section break, a block of
        # vanishing docxtpl directives, another section break, then the table.
        # The two breaks bracket the vanishing block -> blank page on render.
        # The divider must collapse them to a single (vAlign-centered) break.
        doc = Document()
        doc.add_paragraph("O2 RAW DATA")
        break_a = doc.add_paragraph("")
        doc.add_paragraph("{%p set run1 = o2_rawdata %}")
        doc.add_paragraph("{%p if run1 %}")
        doc.add_paragraph("")
        break_b = doc.add_paragraph("")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "O2 %"
        _make_sect_break(doc, break_a)
        _make_sect_break(doc, break_b)
        path = tempfile.mktemp(suffix=".docx")
        doc.save(path)

        assert DocxBulkUpdater([{"op": "divider", "match": "O2 RAW DATA"}]).modify_docx(path) is True

        doc = Document(path)
        els = list(doc.element.body)
        o2 = _para_by_text(doc, "O2 RAW DATA")
        o2_idx = els.index(o2)
        tbl_idx = next(i for i, e in enumerate(els) if e.tag == qn('w:tbl'))
        # Exactly one section break between the divider and the table.
        sect_breaks = [e for e in els[o2_idx + 1:tbl_idx]
                       if e.tag == qn('w:p') and e.find(qn('w:pPr')) is not None
                       and e.find(qn('w:pPr')).find(qn('w:sectPr')) is not None]
        assert len(sect_breaks) == 1
        assert sect_breaks[0].find(qn('w:pPr')).find(qn('w:sectPr')).find(qn('w:vAlign')).get(qn('w:val')) == 'center'

    def test_no_match_returns_false(self):
        path = _make_doc()
        assert DocxBulkUpdater([{"op": "divider", "match": "NOPE"}]).modify_docx(path) is False

    def test_exact_match_preferred_over_substring(self):
        # A TOC/appendix-list entry contains the heading text as a substring;
        # only the exact divider should get the section break after it.
        doc = Document()
        doc.add_paragraph("APPENDIX F O2 RAW DATA")  # substring entry
        doc.add_paragraph("O2 RAW DATA")             # exact divider
        doc.add_table(rows=1, cols=2).rows[0].cells[0].text = "Spec"
        path = tempfile.mktemp(suffix=".docx")
        doc.save(path)

        DocxBulkUpdater([{"op": "divider", "match": "O2 RAW DATA"}]).modify_docx(path)
        doc = Document(path)
        divider = _para_by_text(doc, "O2 RAW DATA")
        nxt = divider.getnext()
        assert nxt.find(qn('w:pPr')).find(qn('w:sectPr')) is not None
        # The appendix-list paragraph was not given a break after it.
        appendix = _para_by_text(doc, "APPENDIX F O2 RAW DATA")
        ap_next = appendix.getnext()
        if ap_next is not None and ap_next.tag == qn('w:p'):
            ap_pPr = ap_next.find(qn('w:pPr'))
            assert ap_pPr is None or ap_pPr.find(qn('w:sectPr')) is None

    def test_config_loading(self, tmp_path):
        config = tmp_path / "cfg.json"
        config.write_text(json.dumps({"divider": {"match": "O2 RAW DATA"}}))
        operations, _ = load_operations_from_json(config)
        assert operations[0]["op"] == "divider"
        assert operations[0]["match"] == "O2 RAW DATA"
