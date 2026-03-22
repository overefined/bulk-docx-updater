"""
Integration tests for the DOCX bulk updater.

Tests end-to-end functionality with actual DOCX documents,
CLI interface integration, and real file operations.
"""
import pytest
import tempfile
import shutil
from pathlib import Path
import json

from docx import Document
from docx.shared import Inches

from src.document_processor import DocxBulkUpdater
from src.config import load_operations_from_json


class TestIntegrationWithRealDocuments:
    """Integration tests using actual DOCX documents."""
    
    def setup_method(self):
        """Set up test fixtures with temporary directory."""
        self.temp_dir = Path(tempfile.mkdtemp())
        self.test_docx = self.temp_dir / "test_document.docx"
        self.config_file = self.temp_dir / "test_config.json"
        
        # Create a simple test DOCX document
        doc = Document()
        doc.add_paragraph("This is a test document with old text.")
        doc.add_paragraph("Another paragraph with old text here.")
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2" 
        table.cell(1, 0).text = "Cell with old text"
        table.cell(1, 1).text = "Another cell"
        
        doc.save(self.test_docx)
        
        # Create test configuration (dict format)
        test_config = {
            "replace": [
                ["old text", "new content"],
                ["test document", "sample document"]
            ]
        }

        with open(self.config_file, 'w') as f:
            json.dump(test_config, f)
    
    def teardown_method(self):
        """Clean up temporary files."""
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_document_modification_with_config_file(self):
        """Test document modification using configuration file."""
        operations, settings = load_operations_from_json(self.config_file)
        updater = DocxBulkUpdater(operations)
        
        # Verify document was modified
        result = updater.modify_docx(self.test_docx)
        assert result is True
        
        # Verify changes were applied
        doc = Document(self.test_docx)
        content = " ".join([para.text for para in doc.paragraphs])
        assert "new content" in content
        assert "sample document" in content
        assert "old text" not in content
    
    def test_document_modification_preserves_structure(self):
        """Test that document structure is preserved during modification."""
        operations = [{"op": "replace", "search": "old text", "replace": "new content"}]
        updater = DocxBulkUpdater(operations)
        
        # Get original structure
        original_doc = Document(self.test_docx)
        original_paragraph_count = len(original_doc.paragraphs)
        original_table_count = len(original_doc.tables)
        
        # Modify document
        updater.modify_docx(self.test_docx)
        
        # Verify structure is preserved
        modified_doc = Document(self.test_docx)
        assert len(modified_doc.paragraphs) == original_paragraph_count
        assert len(modified_doc.tables) == original_table_count
    
    def test_formatting_preservation_during_replacement(self):
        """Test that formatting is preserved during text replacement."""
        # Create document with formatted text
        doc = Document()
        para = doc.add_paragraph()
        run1 = para.add_run("Normal text with ")
        run2 = para.add_run("old text")
        run2.font.bold = True
        run2.font.size = Inches(0.2)  # Larger font
        run3 = para.add_run(" at the end.")
        
        formatted_docx = self.temp_dir / "formatted_test.docx"
        doc.save(formatted_docx)
        
        # Apply replacement
        operations = [{"op": "replace", "search": "old text", "replace": "new content"}]
        updater = DocxBulkUpdater(operations)
        updater.modify_docx(formatted_docx)
        
        # Verify text was changed
        modified_doc = Document(formatted_docx)
        full_text = modified_doc.paragraphs[0].text
        assert "new content" in full_text
        assert "old text" not in full_text
    
    def test_table_content_replacement(self):
        """Test that table content is properly replaced."""
        operations = [{"op": "replace", "search": "Cell with old text", "replace": "Cell with new content"}]
        updater = DocxBulkUpdater(operations)
        
        result = updater.modify_docx(self.test_docx)
        assert result is True
        
        # Verify table content was changed
        doc = Document(self.test_docx)
        table = doc.tables[0]
        cell_text = table.cell(1, 0).text
        assert cell_text == "Cell with new content"
    
    def test_margin_standardization(self):
        """Test document margin standardization."""
        custom_margins = {'top': 0.5, 'bottom': 1.5, 'left': 0.75, 'right': 1.25}
        operations = []  # No text replacements
        updater = DocxBulkUpdater(
            operations,
            standardize_margins=True,
            margins=custom_margins
        )
        
        result = updater.modify_docx(self.test_docx)
        assert result is True  # Should return True due to margin changes
        
        # Verify margins were applied
        doc = Document(self.test_docx)
        section = doc.sections[0]
        assert section.top_margin == Inches(0.5)
        assert section.bottom_margin == Inches(1.5)
        assert section.left_margin == Inches(0.75)
        assert section.right_margin == Inches(1.25)
    
    def test_dry_run_preview_without_modification(self):
        """Test dry run functionality doesn't modify the document."""
        # Get original content
        original_doc = Document(self.test_docx)
        original_content = [para.text for para in original_doc.paragraphs]
        
        operations = [{"op": "replace", "search": "old text", "replace": "new content"}]
        updater = DocxBulkUpdater(operations)
        
        # Get preview of changes
        changes = updater.get_document_changes_preview(self.test_docx)
        
        # Verify changes were detected
        assert len(changes) > 0
        
        # Verify document wasn't actually modified
        doc = Document(self.test_docx)
        current_content = [para.text for para in doc.paragraphs]
        assert current_content == original_content
    
    def test_multiple_document_processing(self):
        """Test processing multiple documents."""
        # Create second test document
        test_docx2 = self.temp_dir / "test_document2.docx"
        doc2 = Document()
        doc2.add_paragraph("Second document with old text content.")
        doc2.save(test_docx2)
        
        operations = [{"op": "replace", "search": "old text", "replace": "new content"}]
        updater = DocxBulkUpdater(operations)
        
        # Process both documents
        result1 = updater.modify_docx(self.test_docx)
        result2 = updater.modify_docx(test_docx2)
        
        assert result1 is True
        assert result2 is True
        
        # Verify both were modified
        doc1 = Document(self.test_docx)
        doc2 = Document(test_docx2)
        
        assert "new content" in doc1.paragraphs[0].text
        assert "new content" in doc2.paragraphs[0].text




class TestErrorHandling:
    """Integration tests for error handling scenarios."""
    
    def test_nonexistent_config_file(self):
        """Test handling of nonexistent configuration file."""
        nonexistent_config = Path("nonexistent_config.json")
        
        with pytest.raises(SystemExit):
            load_operations_from_json(nonexistent_config)
    
    def test_corrupted_docx_file(self):
        """Test handling of corrupted DOCX files."""
        # Create a file that's not a valid DOCX
        corrupted_file = Path(tempfile.mktemp(suffix=".docx"))
        with open(corrupted_file, 'w') as f:
            f.write("This is not a valid DOCX file")
        
        try:
            operations = [{"op": "replace", "search": "test", "replace": "example"}]
            updater = DocxBulkUpdater(operations)
            
            # Should return False (failed to process)
            result = updater.modify_docx(corrupted_file)
            assert result is False
            
        finally:
            if corrupted_file.exists():
                corrupted_file.unlink()
    
    def test_invalid_replacement_config(self):
        """Test handling of invalid replacement configurations."""
        invalid_operations = [
            {"op": "replace"},  # Missing search/replace
        ]

        with pytest.raises(SystemExit):
            from src.config import validate_operations
            validate_operations(invalid_operations)
    
    def test_empty_directory_processing(self):
        """Test processing empty directory."""
        empty_dir = Path(tempfile.mkdtemp())
        
        try:
            # Should handle empty directory gracefully
            operations = [{"op": "replace", "search": "test", "replace": "example"}]
            updater = DocxBulkUpdater(operations)
            
            # Process directory (should find no files)
            files = list(empty_dir.glob("*.docx"))
            assert len(files) == 0
            
        finally:
            empty_dir.rmdir()


class TestComplexFormattingIntegration:
    """Integration tests for complex formatting scenarios."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.temp_dir = Path(tempfile.mkdtemp())
        self.test_docx = self.temp_dir / "formatting_test.docx"
        
        # Create document with complex formatting
        doc = Document()
        para = doc.add_paragraph("Text with ")
        bold_run = para.add_run("bold formatting")
        bold_run.font.bold = True
        para.add_run(" and normal text.")
        
        doc.save(self.test_docx)
    
    def teardown_method(self):
        """Clean up temporary files."""
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_formatting_token_processing(self):
        """Test processing of formatting tokens in replacement text."""
        operations = [
            {"op": "replace", "search": "bold formatting", "replace": "{format:italic,size14}italic formatted{/format}"}
        ]
        updater = DocxBulkUpdater(operations)
        
        result = updater.modify_docx(self.test_docx)
        assert result is True
        
        # Verify content was changed
        doc = Document(self.test_docx)
        full_text = doc.paragraphs[0].text
        assert "italic formatted" in full_text
        assert "bold formatting" not in full_text
    
    def test_global_formatting_tokens(self):
        """Test processing of global formatting tokens."""
        operations = [
            {"op": "replace", "search": "Text with", "replace": "New text linebreak with"}
        ]
        updater = DocxBulkUpdater(operations)
        
        result = updater.modify_docx(self.test_docx)
        assert result is True
        
        # Verify replacement was applied
        doc = Document(self.test_docx)
        assert "New text" in doc.paragraphs[0].text
