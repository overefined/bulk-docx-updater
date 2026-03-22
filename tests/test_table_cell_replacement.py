"""
Unit tests for table cell replacement functionality.

Tests the replace_table_cell operation that allows targeting specific
table cells by coordinates for precise content replacement.
"""
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Inches

from src.document_processor import DocxBulkUpdater
from src.config import load_operations_from_json


class TestTableCellReplacement:
    """Test cases for table cell replacement functionality."""

    def create_test_document(self):
        """Create a test document with tables for testing."""
        doc = Document()

        # Add a simple table
        table = doc.add_table(rows=3, cols=3)

        # Set header row
        table.rows[0].cells[0].text = "Phase"
        table.rows[0].cells[1].text = "Time"
        table.rows[0].cells[2].text = "Value"

        # Set data rows
        table.rows[1].cells[0].text = "{{ phase1 }}"
        table.rows[1].cells[1].text = "{{ time1 }}"
        table.rows[1].cells[2].text = "{{ value1 }}"

        table.rows[2].cells[0].text = "{{ phase2 }}"
        table.rows[2].cells[1].text = "{{ time2 }}"
        table.rows[2].cells[2].text = "{{ value2 }}"

        # Add a second table
        table2 = doc.add_table(rows=2, cols=2)
        table2.rows[0].cells[0].text = "Name"
        table2.rows[0].cells[1].text = "Age"
        table2.rows[1].cells[0].text = "John"
        table2.rows[1].cells[1].text = "25"

        return doc

    def test_basic_cell_replacement(self, tmp_path):
        """Test basic table cell content replacement."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement
        operations = [{
            'op': 'replace_table_cell',

                'table_index': 0,
                'row': 0,
                'column': 0,
                'replace': 'Stage'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)
        table = modified_doc.tables[0]
        assert table.rows[0].cells[0].text == "Stage"
        assert table.rows[0].cells[1].text == "Time"  # Unchanged

    def test_cell_replacement_with_validation(self, tmp_path):
        """Test table cell replacement with content validation."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement with validation
        operations = [{
            'op': 'replace_table_cell',

                'table_index': 0,
                'row': 0,
                'column': 1,
                'search': 'Time',
                'replace': 'Duration'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)
        table = modified_doc.tables[0]
        assert table.rows[0].cells[1].text == "Duration"

    def test_complete_header_matching(self, tmp_path):
        """Test matching complete header row to distinguish between similar tables."""
        # Create document with multiple tables having similar headers
        doc = Document()

        # First table: Phase, Time, Value
        table1 = doc.add_table(rows=2, cols=3)
        table1.rows[0].cells[0].text = "Phase"
        table1.rows[0].cells[1].text = "Time"
        table1.rows[0].cells[2].text = "Value"
        table1.rows[1].cells[0].text = "Data1"

        # Second table: Phase, Time, O2 %
        table2 = doc.add_table(rows=2, cols=3)
        table2.rows[0].cells[0].text = "Phase"
        table2.rows[0].cells[1].text = "Time"
        table2.rows[0].cells[2].text = "O2 %"
        table2.rows[1].cells[0].text = "Data2"

        test_file = tmp_path / "test_complete_header.docx"
        doc.save(test_file)

        # Configure replacement targeting the second table specifically
        operations = [{
            'op': 'replace_table_cell',

                'table_header': 'Phase, Time, O2 %',  # Complete header match
                'row': 1,
                'column': 0,
                'search': 'Data2',
                'replace': 'Modified Data2'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)

        # First table should be unchanged
        table1 = modified_doc.tables[0]
        assert table1.rows[1].cells[0].text == "Data1"

        # Second table should be changed
        table2 = modified_doc.tables[1]
        assert table2.rows[1].cells[0].text == "Modified Data2"

    def test_cell_replacement_validation_failure(self, tmp_path):
        """Test that replacement fails when validation doesn't match."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement with incorrect validation
        operations = [{
            'op': 'replace_table_cell',

                'table_index': 0,
                'row': 0,
                'column': 0,
                'search': 'WrongContent',
                'replace': 'Stage'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)

        # Test the specific cell replacement method
        doc = Document(test_file)
        cell_config = operations[0]
        result = processor.replace_table_cell(doc, cell_config)

        # Verify replacement was not applied due to validation failure
        assert result is False

        # Content should remain unchanged
        table = doc.tables[0]
        assert table.rows[0].cells[0].text == "Phase"

    def test_formatting_tokens_in_replacement(self, tmp_path):
        """Test table cell replacement with formatting tokens."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement with formatting
        operations = [{
            'op': 'replace_table_cell',

                'table_index': 0,
                'row': 0,
                'column': 0,
                'replace': '{format:left,bold}Phase Header{/format}'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)
        table = modified_doc.tables[0]
        cell_text = table.rows[0].cells[0].text
        assert "Phase Header" in cell_text

        # Check that formatting was applied (bold)
        runs = table.rows[0].cells[0].paragraphs[0].runs
        assert any(run.bold for run in runs if "Phase Header" in run.text)

    def test_multiple_table_targeting(self, tmp_path):
        """Test replacement in specific table when multiple tables exist."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement for second table
        operations = [{
            'op': 'replace_table_cell',

                'table_index': 1,
                'row': 0,
                'column': 0,
                'replace': 'Full Name'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)

        # First table should be unchanged
        table1 = modified_doc.tables[0]
        assert table1.rows[0].cells[0].text == "Phase"

        # Second table should be changed
        table2 = modified_doc.tables[1]
        assert table2.rows[0].cells[0].text == "Full Name"

    def test_invalid_table_index(self, tmp_path):
        """Test handling of invalid table index."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement with invalid table index
        operations = [{
            'op': 'replace_table_cell',

                'table_index': 99,  # Non-existent table
                'row': 0,
                'column': 0,
                'replace': 'Test'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)

        # Test the specific cell replacement method
        doc = Document(test_file)
        cell_config = operations[0]
        result = processor.replace_table_cell(doc, cell_config)

        # Verify replacement failed gracefully
        assert result is False

    def test_invalid_row_index(self, tmp_path):
        """Test handling of invalid row index."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement with invalid row index
        operations = [{
            'op': 'replace_table_cell',

                'table_index': 0,
                'row': 99,  # Non-existent row
                'column': 0,
                'replace': 'Test'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)

        # Test the specific cell replacement method
        doc = Document(test_file)
        cell_config = operations[0]
        result = processor.replace_table_cell(doc, cell_config)

        # Verify replacement failed gracefully
        assert result is False

    def test_invalid_column_index(self, tmp_path):
        """Test handling of invalid column index."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement with invalid column index
        operations = [{
            'op': 'replace_table_cell',

                'table_index': 0,
                'row': 0,
                'column': 99,  # Non-existent column
                'replace': 'Test'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)

        # Test the specific cell replacement method
        doc = Document(test_file)
        cell_config = operations[0]
        result = processor.replace_table_cell(doc, cell_config)

        # Verify replacement failed gracefully
        assert result is False

    def test_empty_content_replacement(self, tmp_path):
        """Test replacement with empty content."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement with empty content
        operations = [{
            'op': 'replace_table_cell',

                'table_index': 0,
                'row': 0,
                'column': 0,
                'replace': ''
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)
        table = modified_doc.tables[0]
        assert table.rows[0].cells[0].text == ""

    def test_config_integration(self, tmp_path):
        """Test table cell replacement through config file integration."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Create config file
        config_data = {
            "replace_table_cell": [
                {"table_index": 0, "row": 0, "column": 0, "search": "Phase", "replace": "{format:left}Time{/format}"},
                {"table_index": 0, "row": 0, "column": 1, "search": "Time", "replace": "{format:left}Phase{/format}"}
            ]
        }

        config_file = tmp_path / "test_config.json"
        import json
        config_file.write_text(json.dumps(config_data))

        # Load config and apply replacements
        operations, _ = load_operations_from_json(config_file)
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacements were applied
        assert result is True

        # Check that headers were swapped
        modified_doc = Document(test_file)
        table = modified_doc.tables[0]
        assert table.rows[0].cells[0].text == "Time"
        assert table.rows[0].cells[1].text == "Phase"

    def test_default_table_index(self, tmp_path):
        """Test that table_index defaults to 0 when not specified."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement without table_index (should default to 0)
        operations = [{
            'op': 'replace_table_cell',

                'row': 0,
                'column': 0,
                'replace': 'Default Table'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied to first table
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)
        table = modified_doc.tables[0]
        assert table.rows[0].cells[0].text == "Default Table"

    def test_table_header_selection(self, tmp_path):
        """Test table selection by header content."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement using table_header to target second table
        operations = [{
            'op': 'replace_table_cell',

                'table_header': 'Name',  # This should match the second table
                'row': 0,
                'column': 0,
                'replace': 'Full Name'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)

        # First table should be unchanged
        table1 = modified_doc.tables[0]
        assert table1.rows[0].cells[0].text == "Phase"

        # Second table should be changed
        table2 = modified_doc.tables[1]
        assert table2.rows[0].cells[0].text == "Full Name"

    def test_table_header_not_found(self, tmp_path):
        """Test handling when table header is not found."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement with non-existent header
        operations = [{
            'op': 'replace_table_cell',

                'table_header': 'NonExistentHeader',
                'row': 0,
                'column': 0,
                'replace': 'Test'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)

        # Test the specific cell replacement method
        doc = Document(test_file)
        cell_config = operations[0]
        result = processor.replace_table_cell(doc, cell_config)

        # Verify replacement failed gracefully
        assert result is False

    def test_table_header_with_search_validation(self, tmp_path):
        """Test table header selection with content validation."""
        # Create test document
        doc = self.create_test_document()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement using table_header with validation
        operations = [{
            'op': 'replace_table_cell',

                'table_header': 'Phase',  # This should match the first table
                'row': 0,
                'column': 1,
                'search': 'Time',  # Validate current content
                'replace': 'Duration'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)
        table = modified_doc.tables[0]
        assert table.rows[0].cells[1].text == "Duration"

    def test_complete_header_matching(self, tmp_path):
        """Test matching complete header row to distinguish between similar tables."""
        # Create document with multiple tables having similar headers
        doc = Document()

        # First table: Phase, Time, Value
        table1 = doc.add_table(rows=2, cols=3)
        table1.rows[0].cells[0].text = "Phase"
        table1.rows[0].cells[1].text = "Time"
        table1.rows[0].cells[2].text = "Value"
        table1.rows[1].cells[0].text = "Data1"

        # Second table: Phase, Time, O2 %
        table2 = doc.add_table(rows=2, cols=3)
        table2.rows[0].cells[0].text = "Phase"
        table2.rows[0].cells[1].text = "Time"
        table2.rows[0].cells[2].text = "O2 %"
        table2.rows[1].cells[0].text = "Data2"

        test_file = tmp_path / "test_complete_header.docx"
        doc.save(test_file)

        # Configure replacement targeting the second table specifically
        operations = [{
            'op': 'replace_table_cell',

                'table_header': 'Phase, Time, O2 %',  # Complete header match
                'row': 1,
                'column': 0,
                'search': 'Data2',
                'replace': 'Modified Data2'
            
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)

        # First table should be unchanged
        table1 = modified_doc.tables[0]
        assert table1.rows[1].cells[0].text == "Data1"

        # Second table should be changed
        table2 = modified_doc.tables[1]
        assert table2.rows[1].cells[0].text == "Modified Data2"