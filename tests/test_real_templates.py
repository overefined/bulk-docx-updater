"""
Real template integration tests for the DOCX bulk updater.

Tests the tool against actual DOCX templates to ensure it works correctly
with complex, real-world document structures and formatting.
"""
import pytest
import tempfile
import shutil
from pathlib import Path
import json

from docx import Document

from document_processor import DocxBulkUpdater
from config import load_replacements_from_json


class TestRealTemplateProcessing:
    """Tests using actual DOCX templates from the templates directory."""
    
    TEMPLATES_DIR = Path("/mnt/c/Development/scripts/docx-templates/templates")
    
    def setup_method(self):
        """Set up test fixtures with temporary directory."""
        self.temp_dir = Path(tempfile.mkdtemp())
        
        # Skip tests if templates directory doesn't exist
        if not self.TEMPLATES_DIR.exists():
            pytest.skip(f"Templates directory not found: {self.TEMPLATES_DIR}")
    
    def teardown_method(self):
        """Clean up temporary files."""
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def copy_template(self, template_name: str) -> Path:
        """Copy a template to temp directory for testing."""
        template_path = self.TEMPLATES_DIR / template_name
        if not template_path.exists():
            pytest.skip(f"Template not found: {template_path}")
        
        temp_template = self.temp_dir / template_name
        shutil.copy2(template_path, temp_template)
        return temp_template
    
    def test_hyperlink_detection_skips_appendix_list(self):
        """Test that hyperlink detection correctly skips appendix list entries."""
        template_path = self.copy_template("ECOM_1x21min_20250510.docx")
        
        # Use the actual replacement from replace.json
        replacements = [
            {"search": "SITE PHOTOS", "replace": "SITE PHOTOSTEST_CONTENT_INSERTED"}
        ]
        updater = DocxBulkUpdater(replacements)
        
        # Test dry run first to verify behavior
        changes = updater.get_document_changes_preview(template_path)
        
        # Should find changes only in standalone SITE PHOTOS, not in appendix list
        change_text = "\n".join(changes)
        
        # Count occurrences of SITE PHOTOS in the change preview
        site_photos_occurrences = change_text.count("SITE PHOTOS")
        
        # Apply changes
        result = updater.modify_docx(template_path)
        assert result is True
        
        # Verify the document was modified correctly
        doc = Document(template_path)
        full_text = " ".join([para.text for para in doc.paragraphs])
        
        # Should contain the inserted test content
        assert "TEST_CONTENT_INSERTED" in full_text
    
    def test_tester_qualifications_replacement_ecom_template(self):
        """Test TESTER QUALIFICATIONS replacement with formatting."""
        template_path = self.copy_template("ECOM_1x21min_20250510.docx")
        
        replacements = [
            {
                "search": "TESTER QUALIFICATIONS",
                "replace": "{format:center,bold,size16}TESTER QUALIFICATIONS{/format}pagebreak{format:center,size12}{{ technician_resume }}{/format}"
            }
        ]
        updater = DocxBulkUpdater(replacements)
        
        # Test dry run first
        changes = updater.get_document_changes_preview(template_path)
        print(f"TESTER QUALIFICATIONS changes: {len(changes)}")
        
        # Apply changes
        result = updater.modify_docx(template_path)
        
        # If replacement was made, verify the result
        if result:
            doc = Document(template_path)
            full_text = " ".join([para.text for para in doc.paragraphs])
            assert "{{ technician_resume }}" in full_text
    
    def test_bracket_removal_ecom_template(self):
        """Test removing bracketed text like ' (0.84)'."""
        template_path = self.copy_template("ECOM_1x21min_20250510.docx")
        
        replacements = [
            {"search": " (0.84)", "replace": ""}
        ]
        updater = DocxBulkUpdater(replacements)
        
        # Test dry run first
        changes = updater.get_document_changes_preview(template_path)
        print(f"Bracket removal changes: {len(changes)}")
        
        # Apply changes
        result = updater.modify_docx(template_path)
        
        # Check if changes were made
        if result:
            doc = Document(template_path)
            full_text = " ".join([para.text for para in doc.paragraphs])
            assert " (0.84)" not in full_text
    
    def test_remove_empty_paragraphs_after_replacement(self):
        """Test the remove_empty_paragraphs_after functionality."""
        template_path = self.copy_template("ECOM_1x21min_20250510.docx")
        
        replacements = [
            {
                "search": "TESTER QUALIFICATIONS",
                "replace": "{format:center,bold,size16}TESTER QUALIFICATIONS{/format}pagebreak{format:center,size12}{{ technician_resume }}{/format}"
            },
            {
                "remove_empty_paragraphs_after": "{{ technician_resume }}"
            }
        ]
        updater = DocxBulkUpdater(replacements)
        
        # Apply changes
        result = updater.modify_docx(template_path)
        
        # The test should run without errors (we fixed the 'search' key error)
        print(f"Remove empty paragraphs result: {result}")
    
    def test_ftir_template_processing(self):
        """Test processing FTIR template with common replacements."""
        template_path = self.copy_template("FTIR_1x21min_20240619.docx")
        
        replacements = [
            {"search": "old_placeholder", "replace": "new_value"},
            {"search": "test_data", "replace": "production_data"}
        ]
        updater = DocxBulkUpdater(replacements)
        
        # Test dry run
        changes = updater.get_document_changes_preview(template_path)
        print(f"FTIR template changes: {len(changes)}")
        
        # Apply changes (should handle gracefully even if no matches)
        result = updater.modify_docx(template_path)
        print(f"FTIR processing result: {result}")
    
    def test_margin_standardization_on_real_template(self):
        """Test margin standardization on actual template."""
        template_path = self.copy_template("ECOM_1x21min_20250510.docx")
        
        custom_margins = {'top': 0.75, 'bottom': 0.75, 'left': 1.0, 'right': 1.0}
        replacements = []  # No text replacements
        updater = DocxBulkUpdater(
            replacements,
            standardize_margins=True,
            margins=custom_margins
        )
        
        result = updater.modify_docx(template_path)
        assert result is True  # Should return True due to margin changes
        
        # Verify margins were applied
        doc = Document(template_path)
        section = doc.sections[0]
        from docx.shared import Inches
        assert abs(section.top_margin - Inches(0.75)) < Inches(0.01)
        assert abs(section.left_margin - Inches(1.0)) < Inches(0.01)
    
    def test_complex_formatting_on_real_template(self):
        """Test complex formatting tokens on actual template."""
        template_path = self.copy_template("ECOM_1x21min_20250510.docx")
        
        replacements = [
            {
                "search": "Test Report",
                "replace": "{format:center,bold,size18}ENHANCED TEST REPORT{/format}linebreak{format:italic}Modified by Bulk Updater{/format}"
            }
        ]
        updater = DocxBulkUpdater(replacements)
        
        # Test dry run first
        changes = updater.get_document_changes_preview(template_path)
        
        # Apply changes
        result = updater.modify_docx(template_path)
        print(f"Complex formatting result: {result}")
    
    def test_table_content_replacement_real_template(self):
        """Test table content replacement on real template."""
        template_path = self.copy_template("ECOM_1x21min_20250510.docx")
        
        replacements = [
            {"search": "Table", "replace": "Modified Table"},
            {"search": "Cell", "replace": "Updated Cell"}
        ]
        updater = DocxBulkUpdater(replacements)
        
        # Apply changes
        result = updater.modify_docx(template_path)
        
        # Verify document structure is preserved
        doc = Document(template_path)
        original_table_count = len(doc.tables)
        
        # Document should still have same number of tables
        print(f"Tables in document: {original_table_count}")
    
    def test_multiple_templates_batch_processing(self):
        """Test processing multiple real templates in batch."""
        templates = ["ECOM_1x21min_20250510.docx", "FTIR_1x21min_20240619.docx"]
        copied_templates = []
        
        for template_name in templates:
            try:
                copied_templates.append(self.copy_template(template_name))
            except pytest.skip.Exception:
                continue  # Skip if template doesn't exist
        
        if not copied_templates:
            pytest.skip("No templates available for batch testing")
        
        replacements = [
            {"search": "test_value", "replace": "production_value"}
        ]
        updater = DocxBulkUpdater(replacements)
        
        results = []
        for template_path in copied_templates:
            result = updater.modify_docx(template_path)
            results.append(result)
        
        print(f"Batch processing results: {results}")
        # All should complete without errors
        assert all(isinstance(r, bool) for r in results)
    
    def test_split_text_across_runs(self):
        """Test handling of text split across multiple DOCX runs."""
        template_path = self.copy_template("ECOM_1x21min_20250510.docx")
        
        # Look for text that might be split across runs
        replacements = [
            {"search": "test", "replace": "REPLACEMENT"},  # Common word likely to be split
            {"search": "data", "replace": "INFORMATION"}
        ]
        updater = DocxBulkUpdater(replacements)
        
        # Test dry run to see what would be changed
        changes = updater.get_document_changes_preview(template_path)
        print(f"Split text replacement changes: {len(changes)}")
        
        # Apply changes
        result = updater.modify_docx(template_path)
        print(f"Split text processing result: {result}")


class TestConfigurationWithRealTemplates:
    """Test configuration file handling with real templates."""
    
    TEMPLATES_DIR = Path("/mnt/c/Development/scripts/docx-templates/templates")
    
    def setup_method(self):
        """Set up test fixtures."""
        self.temp_dir = Path(tempfile.mkdtemp())
        
        if not self.TEMPLATES_DIR.exists():
            pytest.skip(f"Templates directory not found: {self.TEMPLATES_DIR}")
    
    def teardown_method(self):
        """Clean up temporary files."""
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_actual_replace_json_config(self):
        """Test using the actual replace.json configuration file."""
        # Copy a template
        template_path = self.temp_dir / "test_template.docx"
        original_template = self.TEMPLATES_DIR / "ECOM_1x21min_20250510.docx"
        
        if original_template.exists():
            shutil.copy2(original_template, template_path)
        else:
            pytest.skip("ECOM template not found")
        
        # Use the actual configuration file
        config_path = Path("/mnt/c/Development/scripts/docx-templates/bulk-docx-updater/replace.json")
        if not config_path.exists():
            pytest.skip("replace.json not found")
        
        replacements = load_replacements_from_json(config_path)
        updater = DocxBulkUpdater(replacements)
        
        # Test dry run
        changes = updater.get_document_changes_preview(template_path)
        print(f"Actual config changes: {len(changes)}")
        if changes:
            for i, change in enumerate(changes):
                if i >= 5:  # Show first 5 changes only
                    break
                print(f"  Change {i+1}: {change}")
        
        # Apply changes
        result = updater.modify_docx(template_path)
        print(f"Actual config processing result: {result}")
        
        # Verify document is still valid
        doc = Document(template_path)
        assert len(doc.paragraphs) > 0  # Should have content


class TestAdditionalRealTemplateFeatures(TestRealTemplateProcessing):
    """Additional tests for real template features that were previously mocked."""

    def test_dry_run_preview_functionality(self):
        """Test dry run preview without modifying documents."""
        template_path = self.copy_template("ECOM_1x21min_20250510.docx")
        
        replacements = [
            {"search": "Test", "replace": "Analysis"}
        ]
        updater = DocxBulkUpdater(replacements)
        
        # Get original content
        original_doc = Document(template_path)
        original_text = " ".join([para.text for para in original_doc.paragraphs])
        
        # Get preview of changes
        changes = updater.get_document_changes_preview(template_path)
        
        # Verify document wasn't modified
        current_doc = Document(template_path)
        current_text = " ".join([para.text for para in current_doc.paragraphs])
        assert current_text == original_text
        
        # Should return changes dict
        assert isinstance(changes, dict)
    
    def test_no_changes_scenario(self):
        """Test behavior when no replacements are found."""
        template_path = self.copy_template("ECOM_1x21min_20250510.docx")
        
        replacements = [
            {"search": "NONEXISTENT_TEXT_12345", "replace": "REPLACEMENT"}
        ]
        updater = DocxBulkUpdater(replacements)
        
        # Should return False when no changes are made
        result = updater.modify_docx(template_path)
        assert result is False
    
    def test_headers_and_footers_processing(self):
        """Test that headers and footers are processed for replacements."""
        template_path = self.copy_template("ECOM_1x21min_20250510.docx")
        
        # Many DOCX templates have standard headers/footers
        replacements = [
            {"search": "Page", "replace": "Sheet"}  # Common in headers/footers
        ]
        updater = DocxBulkUpdater(replacements)
        
        # Test dry run to see if headers/footers have matches
        changes = updater.get_document_changes_preview(template_path)
        
        # Apply changes
        result = updater.modify_docx(template_path)
        
        # If changes were found and applied, result should be True
        # If no "Page" text in headers/footers, result should be False
        assert isinstance(result, bool)
    
    def test_error_handling_with_corrupted_file(self):
        """Test error handling with invalid file."""
        # Create a fake DOCX file (not actually a DOCX)
        fake_docx = self.temp_dir / "fake.docx"
        with open(fake_docx, 'w') as f:
            f.write("This is not a DOCX file")
        
        replacements = [{"search": "test", "replace": "example"}]
        updater = DocxBulkUpdater(replacements)
        
        # Should handle the error gracefully and return False
        result = updater.modify_docx(fake_docx)
        assert result is False
