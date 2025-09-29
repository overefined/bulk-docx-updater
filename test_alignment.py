#!/usr/bin/env python3
"""
Test script to verify table cell alignment properties in DOCX files.
This script checks if the alignment formatting is actually being applied.
"""

import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def check_table_alignment(docx_path: Path):
    """Check alignment of paragraphs in table cells."""
    print(f"Checking alignment in: {docx_path}")

    try:
        doc = Document(docx_path)

        table_count = 0
        cell_count = 0
        alignment_info = []

        for table_idx, table in enumerate(doc.tables):
            table_count += 1
            print(f"\n--- Table {table_idx + 1} ---")

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_count += 1

                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        # Check if paragraph contains our target text
                        para_text = paragraph.text.strip()
                        if any(target in para_text for target in ["o2.phase_fmtd", "o2.ReadingTimestamp"]):
                            alignment = paragraph.alignment
                            alignment_name = "None"

                            if alignment is not None:
                                alignment_map = {
                                    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
                                    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
                                    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
                                    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY"
                                }
                                alignment_name = alignment_map.get(alignment, f"Unknown({alignment})")

                            info = {
                                'table': table_idx + 1,
                                'row': row_idx + 1,
                                'cell': cell_idx + 1,
                                'paragraph': para_idx + 1,
                                'text': para_text[:50] + ("..." if len(para_text) > 50 else ""),
                                'alignment': alignment_name,
                                'alignment_value': alignment
                            }
                            alignment_info.append(info)

                            print(f"  Row {row_idx+1}, Cell {cell_idx+1}, Para {para_idx+1}: "
                                  f"'{para_text[:30]}...' -> {alignment_name}")

        print(f"\n--- Summary ---")
        print(f"Total tables: {table_count}")
        print(f"Total cells: {cell_count}")
        print(f"Target paragraphs found: {len(alignment_info)}")

        # Count alignment types
        alignment_counts = {}
        for info in alignment_info:
            align_name = info['alignment']
            alignment_counts[align_name] = alignment_counts.get(align_name, 0) + 1

        print(f"\nAlignment distribution:")
        for align_type, count in alignment_counts.items():
            print(f"  {align_type}: {count}")

        # Check if all target paragraphs are LEFT aligned
        left_aligned = [info for info in alignment_info if info['alignment'] == 'LEFT']
        if len(left_aligned) == len(alignment_info):
            print(f"\n✅ SUCCESS: All {len(alignment_info)} target paragraphs are LEFT aligned!")
        else:
            print(f"\n❌ ISSUE: Only {len(left_aligned)}/{len(alignment_info)} target paragraphs are LEFT aligned")
            non_left = [info for info in alignment_info if info['alignment'] != 'LEFT']
            for info in non_left:
                print(f"  - Table {info['table']}, Row {info['row']}, Cell {info['cell']}: "
                      f"'{info['text']}' is {info['alignment']}")

        return alignment_info

    except Exception as e:
        print(f"Error checking alignment: {e}")
        return []

def main():
    if len(sys.argv) != 2:
        print("Usage: python test_alignment.py <docx_file>")
        sys.exit(1)

    docx_path = Path(sys.argv[1])
    if not docx_path.exists():
        print(f"Error: File not found: {docx_path}")
        sys.exit(1)

    alignment_info = check_table_alignment(docx_path)

if __name__ == "__main__":
    main()