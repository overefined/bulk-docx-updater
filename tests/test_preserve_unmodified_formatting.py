"""Test that formatting is preserved for unmodified parts of text."""
import pytest
from docx import Document
from docx.shared import Pt
from src.document_processor import DocxBulkUpdater
import tempfile
from pathlib import Path


def test_preserve_bold_in_unmodified_text():
    """Test that bold formatting is preserved in text that wasn't part of the replacement."""
    # Create a test document
    doc = Document()
    para = doc.add_paragraph()

    # Add "Label: " in bold
    run1 = para.add_run("Label: ")
    run1.font.bold = True

    # Add "old_text" in normal formatting
    run2 = para.add_run("old_text")
    run2.font.bold = False

    # Save to temp file
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        temp_path = Path(tmp.name)

    try:
        # Apply replacement that only changes "old_text"
        config = [{
            "op": "replace",
            "search": "old_text",
            "replace": "new_text"
        }]

        updater = DocxBulkUpdater(config)
        updater.modify_docx(temp_path)

        # Load modified document
        modified_doc = Document(temp_path)
        modified_para = modified_doc.paragraphs[0]

        # Verify the text was changed
        assert modified_para.text == "Label: new_text"

        # Verify formatting: "Label: " should still be bold
        # Find the run(s) containing "Label: "
        full_text = ""
        for run in modified_para.runs:
            full_text += run.text
            if "Label:" in full_text and "Label:" in run.text:
                # This run contains the label
                assert run.font.bold == True, f"Label should be bold but got bold={run.font.bold}"
                break
        else:
            # If we didn't find it in a single run, check if the beginning runs are bold
            beginning_text = ""
            for run in modified_para.runs:
                beginning_text += run.text
                if "Label:" in beginning_text:
                    # We've accumulated enough text to include the label
                    # Check if this run or previous runs contain the label text
                    if "Label" in run.text or ":" in run.text:
                        assert run.font.bold == True, f"Label portion should be bold but run '{run.text}' has bold={run.font.bold}"
                    if "Label:" in beginning_text:
                        break

    finally:
        # Cleanup
        temp_path.unlink()


def test_preserve_bold_label_with_template_variables():
    """Test realistic case: bold label followed by template variables that get replaced."""
    # Create a test document
    doc = Document()
    para = doc.add_paragraph()

    # Add "Combustor:" in bold (split across two runs like in real doc)
    run1 = para.add_run("Combustor")
    run1.font.bold = True
    run2 = para.add_run(":")
    run2.font.bold = True

    # Add template variables in normal formatting
    run3 = para.add_run(" {{ manufacturer }} {{ model }}, {{ serial }}")
    run3.font.bold = False

    # Save to temp file
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        temp_path = Path(tmp.name)

    try:
        # Apply replacement that modifies the serial number part
        config = [{
            "op": "replace",
            "search": ", {{ serial }}",
            "replace": ", Unit {{ unit }}, {{ serial }}"
        }]

        updater = DocxBulkUpdater(config)
        updater.modify_docx(temp_path)

        # Load modified document
        modified_doc = Document(temp_path)
        modified_para = modified_doc.paragraphs[0]

        # Verify the text was changed
        assert "Unit {{ unit }}" in modified_para.text

        # Verify formatting: "Combustor:" should still be bold
        found_bold_combustor = False
        for run in modified_para.runs:
            if "Combustor" in run.text or (":" in run.text and modified_para.runs[0].text == run.text):
                assert run.font.bold == True, f"Combustor/colon should be bold but run '{run.text}' has bold={run.font.bold}"
                found_bold_combustor = True

        assert found_bold_combustor, "Did not find bold 'Combustor' in modified paragraph"

    finally:
        # Cleanup
        temp_path.unlink()


if __name__ == "__main__":
    test_preserve_bold_in_unmodified_text()
    test_preserve_bold_label_with_template_variables()
    print("All tests passed!")
