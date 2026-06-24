"""
Unit tests for the insert_block operation.

insert_block adds brand-new body-level content (paragraphs and/or tables) at an
anchor paragraph located by text — used to add a section that didn't exist
before (e.g. a new raw-data appendix), which replace_table can't do.
"""
import tempfile

from docx import Document
from docx.oxml.ns import qn

from src.document_processor import DocxBulkUpdater
from src.config import load_operations_from_json


BLOCK = (
    "<block>"
    "<w:p><w:pPr></w:pPr><w:r><w:t>NEW HEADING</w:t></w:r></w:p>"
    "<w:tbl><w:tblPr><w:tblW w:w=\"0\" w:type=\"auto\"/></w:tblPr>"
    "<w:tblGrid><w:gridCol w:w=\"100\"/></w:tblGrid>"
    "<w:tr><w:tc><w:tcPr><w:tcW w:w=\"0\" w:type=\"auto\"/></w:tcPr>"
    "<w:p><w:r><w:t>CELL</w:t></w:r></w:p></w:tc></w:tr></w:tbl>"
    "</block>"
)


def _make_doc():
    doc = Document()
    doc.add_paragraph("INTRO")
    doc.add_paragraph("ANCHOR")
    doc.add_paragraph("TAIL")
    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)
    return path


def _body_paragraph_texts(doc):
    out = []
    for el in doc.element.body:
        if el.tag == qn('w:p'):
            out.append("".join(x.text or "" for x in el.iter(qn('w:t'))).strip())
    return out


class TestInsertBlock:

    def test_insert_before_anchor(self):
        path = _make_doc()
        proc = DocxBulkUpdater([{"op": "insert_block", "before": "ANCHOR", "replace": BLOCK}])
        assert proc.modify_docx(path) is True

        doc = Document(path)
        texts = _body_paragraph_texts(doc)
        # Heading lands immediately before the anchor; a new table is created.
        assert "NEW HEADING" in texts
        assert texts.index("NEW HEADING") < texts.index("ANCHOR")
        assert texts.index("NEW HEADING") > texts.index("INTRO")
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "CELL"

    def test_insert_after_anchor_preserves_order(self):
        path = _make_doc()
        block = ("<block>"
                 "<w:p><w:r><w:t>FIRST</w:t></w:r></w:p>"
                 "<w:p><w:r><w:t>SECOND</w:t></w:r></w:p>"
                 "</block>")
        proc = DocxBulkUpdater([{"op": "insert_block", "after": "ANCHOR", "replace": block}])
        assert proc.modify_docx(path) is True

        texts = _body_paragraph_texts(Document(path))
        assert texts.index("ANCHOR") < texts.index("FIRST") < texts.index("SECOND") < texts.index("TAIL")

    def test_skip_if_present_is_idempotent(self):
        path = _make_doc()
        op = {"op": "insert_block", "before": "ANCHOR", "replace": BLOCK,
              "skip_if_present": "NEW HEADING"}
        assert DocxBulkUpdater([op]).modify_docx(path) is True
        # Second run is a no-op because the marker is already present.
        assert DocxBulkUpdater([op]).modify_docx(path) is False
        assert _body_paragraph_texts(Document(path)).count("NEW HEADING") == 1

    def test_missing_anchor_is_noop(self):
        path = _make_doc()
        proc = DocxBulkUpdater([{"op": "insert_block", "before": "NOPE", "replace": BLOCK}])
        assert proc.modify_docx(path) is False

    def test_config_loads_replace_file(self, tmp_path):
        block_file = tmp_path / "block.xml"
        block_file.write_text(BLOCK, encoding="utf-8")
        cfg = tmp_path / "cfg.json"
        cfg.write_text(
            '{"insert_block": {"before": "ANCHOR", "replace_file": "block.xml"}}',
            encoding="utf-8")
        ops, _ = load_operations_from_json(cfg)
        assert len(ops) == 1
        assert ops[0]["op"] == "insert_block"
        assert ops[0]["before"] == "ANCHOR"
        assert "NEW HEADING" in ops[0]["replace"]
