"""
Unit tests for the DocxBulkUpdater class.

Tests document-level operations focusing on initialization and component interaction.
Real document processing is tested in test_real_templates.py and test_integration.py.
"""
import pytest
from pathlib import Path
from docx import Document

from src.document_processor import DocxBulkUpdater


class TestDocxBulkUpdater:
    """Test cases for DocxBulkUpdater class initialization and basic functionality."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.operations = [
            {"op": "replace", "search": "old text", "replace": "new text"},
            {"op": "replace", "search": "test", "replace": "example"}
        ]
    
    def test_init_default_parameters(self):
        """Test initialization with default parameters."""
        updater = DocxBulkUpdater(self.operations)
        
        assert updater.operations == self.operations
        assert updater.preserve_formatting is True
        assert updater.standardize_margins is False
        assert updater.margins == {
            'top': 1.0, 'bottom': 1.0, 'left': 1.0, 'right': 1.0
        }
    
    def test_init_custom_parameters(self):
        """Test initialization with custom parameters."""
        custom_margins = {'top': 0.5, 'bottom': 0.5, 'left': 0.75, 'right': 0.75}
        updater = DocxBulkUpdater(
            self.operations,
            preserve_formatting=False,
            standardize_margins=True,
            margins=custom_margins
        )
        
        assert updater.preserve_formatting is False
        assert updater.standardize_margins is True
        assert updater.margins == custom_margins
    
    def test_components_initialization(self):
        """Test that internal components are properly initialized."""
        updater = DocxBulkUpdater(self.operations)
        
        # Check that formatter and text_replacer are initialized
        assert hasattr(updater, 'formatter')
        assert hasattr(updater, 'text_replacer')
        assert updater.formatter is not None
        assert updater.text_replacer is not None
    
    def test_format_diff(self):
        """Test the format_diff method."""
        updater = DocxBulkUpdater(self.operations)
        original_lines = ["line 1", "line 2", "line 3"]
        modified_lines = ["line 1", "changed line 2", "line 3"]
        
        diff = updater.format_diff(original_lines, modified_lines, "test_section")
        
        assert isinstance(diff, str)
        assert "line 1" in diff
        assert "changed line 2" in diff
        # Should show the change
        assert "-" in diff or "+" in diff
    
    def test_remove_empty_paragraphs_after_pattern(self):
        """Test the remove_empty_paragraphs_after_pattern method."""
        updater = DocxBulkUpdater(self.operations)
        
        # Create a test document
        doc = Document()
        doc.add_paragraph("Some content")
        doc.add_paragraph("SITE PHOTOS")
        
        # Next paragraph with empty run
        next_para = doc.add_paragraph("")
        next_para.add_run("")  # empty run - should be removed
        next_para.add_run("More content")
        
        # Test cleanup
        modified = updater.remove_empty_paragraphs_after_pattern(doc, "SITE PHOTOS")
        
        assert modified is True
        assert next_para.text == "More content"
    
    def test_remove_empty_paragraphs_after_pattern_no_match(self):
        """Test remove_empty_paragraphs_after_pattern when pattern is not found."""
        updater = DocxBulkUpdater(self.operations)
        
        # Create a test document without the pattern
        doc = Document()
        doc.add_paragraph("Some content")
        doc.add_paragraph("Other content")
        doc.add_paragraph("")  # empty
        
        # Test cleanup
        modified = updater.remove_empty_paragraphs_after_pattern(doc, "SITE PHOTOS")
        
        assert modified is False
        assert len(doc.paragraphs) == 3  # No changes made