"""
Tests for cross-paragraph text replacement functionality.

These tests ensure that text patterns spanning multiple paragraphs
are correctly identified and replaced while preserving document structure.
"""
import pytest
from docx import Document
from docx.text.paragraph import Paragraph
from text_replacement import TextReplacer
from formatting import FormattingProcessor


class TestCrossParagraphReplacement:
    """Test cross-paragraph text replacement functionality."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.formatter = FormattingProcessor()
        
    def create_test_document_with_split_text(self, text_parts):
        """Create a test document with text split across multiple paragraphs."""
        doc = Document()
        for text in text_parts:
            para = doc.add_paragraph(text)
        return doc
    
    def test_simple_cross_paragraph_replacement(self):
        """Test basic cross-paragraph text replacement."""
        operations = [
            {
                "op": "replace",
                "search": "Hello World",
                "replace": "Goodbye World"
            }
        ]
        
        replacer = TextReplacer(operations, self.formatter)
        
        # Create document with text split across paragraphs
        doc = self.create_test_document_with_split_text(["Hello ", "World"])
        paragraphs = list(doc.paragraphs)
        
        # Test the replacement
        result = replacer.replace_text_across_paragraphs(paragraphs)
        
        assert result is True
        assert paragraphs[0].text == "Goodbye World"
        assert paragraphs[1].text == ""  # Second paragraph should be cleared
    
    def test_template_pattern_cross_paragraph_replacement(self):
        """Test replacement of template patterns spanning paragraphs."""
        operations = [
            {
                "op": "replace",
                "search": "{% if cylinder_certs != none %}{% for cert in cylinder_certs %}{{ cert }}{% endfor %}{% endif %}",
                "replace": "{% if cylinder_certs != none %}{% for img in cylinder_certs %}{{ img }}{% endfor %}{% endif %}"
            }
        ]
        
        replacer = TextReplacer(operations, self.formatter)
        
        # Create document with template split across paragraphs (like in actual DOCX)
        doc = self.create_test_document_with_split_text([
            "{% if cylinder_certs != none %}{% for cert in cylinder_certs %}{{ cert }}",
            "{% endfor %}{% endif %}"
        ])
        paragraphs = list(doc.paragraphs)
        
        result = replacer.replace_text_across_paragraphs(paragraphs)
        
        assert result is True
        assert "{% for img in cylinder_certs %}{{ img }}" in paragraphs[0].text
        assert paragraphs[1].text == ""  # Second paragraph cleared
    
    def test_no_cross_paragraph_match(self):
        """Test that single-paragraph patterns are not processed."""
        operations = [
            {
                "op": "replace",
                "search": "Complete Pattern",
                "replace": "Replaced Pattern"
            }
        ]
        
        replacer = TextReplacer(operations, self.formatter)
        
        # Create document where pattern is complete in one paragraph
        doc = self.create_test_document_with_split_text([
            "Complete Pattern",
            "Other text"
        ])
        paragraphs = list(doc.paragraphs)
        
        # Cross-paragraph replacement should not process this
        result = replacer.replace_text_across_paragraphs(paragraphs)
        assert result is False  # Should return False, let single-paragraph processing handle it
    
    def test_pattern_not_found(self):
        """Test when search pattern is not found in paragraphs."""
        operations = [
            {
                "op": "replace",
                "search": "NonExistent Pattern",
                "replace": "Replacement"
            }
        ]
        
        replacer = TextReplacer(operations, self.formatter)
        
        doc = self.create_test_document_with_split_text(["Hello", "World"])
        paragraphs = list(doc.paragraphs)
        
        result = replacer.replace_text_across_paragraphs(paragraphs)
        assert result is False
        assert paragraphs[0].text == "Hello"  # Unchanged
        assert paragraphs[1].text == "World"  # Unchanged
    
    def test_three_paragraph_span(self):
        """Test replacement spanning three paragraphs."""
        operations = [
            {
                "op": "replace",
                "search": "First Second Third",
                "replace": "Replaced Text"
            }
        ]
        
        replacer = TextReplacer(operations, self.formatter)
        
        doc = self.create_test_document_with_split_text(["First ", "Second ", "Third"])
        paragraphs = list(doc.paragraphs)
        
        result = replacer.replace_text_across_paragraphs(paragraphs)
        
        assert result is True
        assert paragraphs[0].text == "Replaced Text"
        assert paragraphs[1].text == ""  # Cleared
        assert paragraphs[2].text == ""  # Cleared
    
    def test_multiple_replacements_same_paragraphs(self):
        """Test multiple different cross-paragraph patterns."""
        operations = [
            {
                "op": "replace",
                "search": "Pattern One",
                "replace": "Replacement One"
            },
            {
                "op": "replace",
                "search": "Pattern Two",
                "replace": "Replacement Two"
            }
        ]
        
        replacer = TextReplacer(operations, self.formatter)
        
        # Create document with first pattern split
        doc = self.create_test_document_with_split_text([
            "Pattern ", "One", "Pattern ", "Two"
        ])
        paragraphs = list(doc.paragraphs)
        
        # Should process the first match it finds
        result = replacer.replace_text_across_paragraphs(paragraphs)
        assert result is True
        # First pattern should be replaced
        assert "Replacement One" in paragraphs[0].text
    
    def test_empty_paragraph_list(self):
        """Test handling of empty paragraph list."""
        operations = [{"op": "replace", "search": "test", "replace": "replaced"}]
        replacer = TextReplacer(operations, self.formatter)
        
        result = replacer.replace_text_across_paragraphs([])
        assert result is False
    
    def test_single_paragraph_list(self):
        """Test handling of single paragraph (should not process)."""
        operations = [{"op": "replace", "search": "test", "replace": "replaced"}]
        replacer = TextReplacer(operations, self.formatter)
        
        doc = self.create_test_document_with_split_text(["test"])
        paragraphs = list(doc.paragraphs)
        
        # Single paragraph should not be processed by cross-paragraph method
        result = replacer.replace_text_across_paragraphs(paragraphs)
        assert result is False
    
    def test_formatting_preservation(self):
        """Test that basic formatting is preserved during cross-paragraph replacement."""
        operations = [
            {
                "op": "replace",
                "search": "Bold Text",
                "replace": "{format:bold}Formatted Text{/format}"
            }
        ]
        
        replacer = TextReplacer(operations, self.formatter)
        
        doc = self.create_test_document_with_split_text(["Bold ", "Text"])
        
        # Add some basic formatting to first paragraph
        first_para = doc.paragraphs[0]
        if first_para.runs:
            first_para.runs[0].bold = True
        
        paragraphs = list(doc.paragraphs)
        result = replacer.replace_text_across_paragraphs(paragraphs)
        
        assert result is True
        assert "Formatted Text" in paragraphs[0].text
        # Check that a run exists (formatting was processed)
        assert len(paragraphs[0].runs) > 0
    
    def test_append_after_cross_paragraph(self):
        """Test that appending via replace works with cross-paragraph patterns."""
        operations = [
            {
                "op": "replace",
                "search": "Pattern Text",
                "replace": "Pattern TextInserted Content"
            }
        ]
        
        replacer = TextReplacer(operations, self.formatter)
        
        doc = self.create_test_document_with_split_text(["Pattern ", "Text"])
        paragraphs = list(doc.paragraphs)
        
        # Cross-paragraph replacement should handle this
        result = replacer.replace_text_across_paragraphs(paragraphs)
        assert result is True
        # Should contain both original pattern and inserted content
        combined_text = paragraphs[0].text
        assert "Pattern Text" in combined_text
        assert "Inserted Content" in combined_text
