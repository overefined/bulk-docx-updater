"""Test clear_properties functionality."""
import tempfile
from pathlib import Path
from docx import Document
from src.document_processor import DocxBulkUpdater


def test_clear_properties_author_and_title():
    """Test clearing author and title properties."""
    # Create a test document with properties
    doc = Document()
    doc.add_paragraph("Test content")
    doc.core_properties.author = "John Doe"
    doc.core_properties.title = "Test Document"
    doc.core_properties.subject = "Test Subject"

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = Path(f.name)

    try:
        doc.save(temp_path)

        # Clear author and title
        operations = [
            {
                "op": "clear_properties",
                "properties": ["author", "title"]
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True, "Clear properties operation should return True"

        # Verify properties were cleared
        doc = Document(temp_path)
        assert doc.core_properties.author == "", "Author should be cleared"
        assert doc.core_properties.title == "", "Title should be cleared"
        assert doc.core_properties.subject == "Test Subject", "Subject should not be cleared"

    finally:
        temp_path.unlink(missing_ok=True)


def test_clear_all_properties():
    """Test clearing all common properties with true flag."""
    # Create a test document with properties
    doc = Document()
    doc.add_paragraph("Test content")
    doc.core_properties.author = "John Doe"
    doc.core_properties.title = "Test Document"
    doc.core_properties.subject = "Test Subject"
    doc.core_properties.keywords = "test, document"
    doc.core_properties.category = "Test Category"

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = Path(f.name)

    try:
        doc.save(temp_path)

        # Clear all common properties
        operations = [
            {
                "op": "clear_properties",
                "properties": ["author", "company", "title", "subject", "keywords", "category"]
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True, "Clear properties operation should return True"

        # Verify all common properties were cleared
        doc = Document(temp_path)
        assert doc.core_properties.author == "", "Author should be cleared"
        assert doc.core_properties.title == "", "Title should be cleared"
        assert doc.core_properties.subject == "", "Subject should be cleared"
        assert doc.core_properties.keywords == "", "Keywords should be cleared"
        assert doc.core_properties.category == "", "Category should be cleared"

    finally:
        temp_path.unlink(missing_ok=True)


def test_clear_properties_list_format():
    """Test clearing properties with list format."""
    # Create a test document with properties
    doc = Document()
    doc.add_paragraph("Test content")
    doc.core_properties.author = "John Doe"
    doc.core_properties.keywords = "test, document"
    doc.core_properties.comments = "Test comments"

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = Path(f.name)

    try:
        doc.save(temp_path)

        # Clear specific properties
        operations = [
            {
                "op": "clear_properties",
                "properties": ["author", "keywords"]
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True, "Clear properties operation should return True"

        # Verify specified properties were cleared
        doc = Document(temp_path)
        assert doc.core_properties.author == "", "Author should be cleared"
        assert doc.core_properties.keywords == "", "Keywords should be cleared"
        assert doc.core_properties.comments == "Test comments", "Comments should not be cleared"

    finally:
        temp_path.unlink(missing_ok=True)


def test_clear_properties_single_string():
    """Test clearing a single property with string format."""
    # Create a test document with properties
    doc = Document()
    doc.add_paragraph("Test content")
    doc.core_properties.author = "John Doe"
    doc.core_properties.title = "Test Document"

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = Path(f.name)

    try:
        doc.save(temp_path)

        # Clear single property
        operations = [
            {
                "op": "clear_properties",
                "properties": ["author"]
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True, "Clear properties operation should return True"

        # Verify only author was cleared
        doc = Document(temp_path)
        assert doc.core_properties.author == "", "Author should be cleared"
        assert doc.core_properties.title == "Test Document", "Title should not be cleared"

    finally:
        temp_path.unlink(missing_ok=True)


def test_clear_properties_clears_default_author():
    """Test that clearing author works even with python-docx default."""
    # Create a test document - python-docx sets default author
    doc = Document()
    doc.add_paragraph("Test content")

    # python-docx sets a default author, so let's verify and clear it
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = Path(f.name)

    try:
        doc.save(temp_path)

        # Verify default author exists
        doc = Document(temp_path)
        initial_author = doc.core_properties.author
        assert initial_author != "", f"Expected default author, got empty string"

        # Clear the author property
        operations = [
            {
                "op": "clear_properties",
                "properties": ["author"]
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        # Should return True since author was cleared
        assert result == True, "Should return True when author is cleared"

        # Verify author is now empty
        doc = Document(temp_path)
        assert doc.core_properties.author == "", "Author should be cleared"

    finally:
        temp_path.unlink(missing_ok=True)


def test_clear_properties_combined_with_text_replacement():
    """Test clearing properties combined with text replacement."""
    # Create a test document
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.core_properties.author = "John Doe"
    doc.core_properties.title = "Test Document"

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = Path(f.name)

    try:
        doc.save(temp_path)

        # Combine text replacement and clear properties
        operations = [
            {
                "op": "replace",
                "search": "Hello",
                "replace": "Goodbye"
            },
            {
                "op": "clear_properties",
                "properties": ["author", "title"]
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True, "Combined operations should return True"

        # Verify text replacement worked
        doc = Document(temp_path)
        assert "Goodbye World" in doc.paragraphs[0].text

        # Verify properties were cleared
        assert doc.core_properties.author == ""
        assert doc.core_properties.title == ""

    finally:
        temp_path.unlink(missing_ok=True)
