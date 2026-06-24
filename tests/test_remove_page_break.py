"""
Unit tests for the remove_page_break operation.

remove_page_break strips <w:br w:type="page"/> runs from a paragraph located by
text, operating on the element tree (robust to XML whitespace), while leaving
<w:lastRenderedPageBreak/> render hints untouched.
"""
import tempfile

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml

from src.document_processor import DocxBulkUpdater


def _make_doc():
    doc = Document()
    doc.add_paragraph("BEFORE")
    # Paragraph with text run + a separate page-break run (the certs pattern).
    p = doc.add_paragraph()
    p._p.append(parse_xml(
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:lastRenderedPageBreak/><w:t>{% for img in cylinder_certs %}{{ img }}{% endfor %}</w:t></w:r>'))
    p._p.append(parse_xml(
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:br w:type="page"/></w:r>'))
    doc.add_paragraph("AFTER")
    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)
    return path


def _certs_para(doc):
    for el in doc.element.body:
        if el.tag == qn('w:p'):
            t = "".join(x.text or "" for x in el.iter(qn('w:t')))
            if 'cylinder_certs' in t:
                return el
    return None


class TestRemovePageBreak:

    def test_removes_page_break_run(self):
        path = _make_doc()
        proc = DocxBulkUpdater([{"op": "remove_page_break", "in_paragraph": "{% for img in cylinder_certs %}"}])
        assert proc.modify_docx(path) is True

        para = _certs_para(Document(path))
        page_brks = [b for b in para.iter(qn('w:br')) if b.get(qn('w:type')) == 'page']
        assert page_brks == []
        # The render hint and the text run are preserved.
        assert para.find(qn('w:lastRenderedPageBreak')) is not None or \
            any(r.find(qn('w:lastRenderedPageBreak')) is not None for r in para.findall(qn('w:r')))
        assert 'cylinder_certs' in "".join(x.text or "" for x in para.iter(qn('w:t')))

    def test_drops_emptied_run(self):
        path = _make_doc()
        before = len(_certs_para(Document(path)).findall(qn('w:r')))
        DocxBulkUpdater([{"op": "remove_page_break", "in_paragraph": "cylinder_certs"}]).modify_docx(path)
        after = len(_certs_para(Document(path)).findall(qn('w:r')))
        # The page-break-only run is removed entirely, leaving just the text run.
        assert after == before - 1

    def test_noop_when_no_page_break(self):
        path = _make_doc()
        op = {"op": "remove_page_break", "in_paragraph": "cylinder_certs"}
        assert DocxBulkUpdater([op]).modify_docx(path) is True
        # Second run finds no page break -> no-op.
        assert DocxBulkUpdater([op]).modify_docx(path) is False

    def test_missing_paragraph_is_noop(self):
        path = _make_doc()
        proc = DocxBulkUpdater([{"op": "remove_page_break", "in_paragraph": "NOPE"}])
        assert proc.modify_docx(path) is False
