"""Tests for document title removal functionality."""
import tempfile
import shutil
from pathlib import Path
from docx import Document

from src.document_processor import DocxBulkUpdater


def test_remove_title():
    """Test that document title can be removed."""
    # Create a test document with a title
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        temp_path = Path(tmp.name)

    try:
        # Create document with a title
        doc = Document()
        doc.core_properties.title = "Test Document Title"
        doc.add_paragraph("Some content")
        doc.save(temp_path)

        # Verify title exists
        doc = Document(temp_path)
        assert doc.core_properties.title == "Test Document Title"

        # Remove title using the operation
        operations = [{"op": "remove_title"}]
        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        # Verify modification was made
        assert result is True

        # Verify title was removed
        doc = Document(temp_path)
        assert doc.core_properties.title == ""

    finally:
        temp_path.unlink(missing_ok=True)


def test_remove_title_when_no_title():
    """Test that removing title when no title exists returns False."""
    # Create a test document without a title
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        temp_path = Path(tmp.name)

    try:
        # Create document without a title
        doc = Document()
        doc.add_paragraph("Some content")
        doc.save(temp_path)

        # Verify no title exists
        doc = Document(temp_path)
        assert doc.core_properties.title == ""

        # Try to remove title
        operations = [{"op": "remove_title"}]
        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        # Verify no modification was made
        assert result is False

        # Verify title is still empty
        doc = Document(temp_path)
        assert doc.core_properties.title == ""

    finally:
        temp_path.unlink(missing_ok=True)


def test_remove_title_preview():
    """Test that removing title shows up in preview."""
    # Create a test document with a title
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        temp_path = Path(tmp.name)

    try:
        # Create document with a title
        doc = Document()
        doc.core_properties.title = "Original Title"
        doc.add_paragraph("Some content")
        doc.save(temp_path)

        # Get preview of changes
        operations = [{"op": "remove_title"}]
        updater = DocxBulkUpdater(operations)
        changes = updater.get_document_changes_preview(temp_path)

        # Verify preview shows title removal
        assert "Formatting Operations" in changes
        _, operations_list = changes["Formatting Operations"]
        assert any("Removed document title: 'Original Title'" in op for op in operations_list)

        # Verify original document is unchanged
        doc = Document(temp_path)
        assert doc.core_properties.title == "Original Title"

    finally:
        temp_path.unlink(missing_ok=True)
