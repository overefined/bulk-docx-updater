"""Test set_comments functionality."""
import tempfile
from pathlib import Path
from docx import Document
from src.document_processor import DocxBulkUpdater


def test_set_comments_with_static_value():
    """Test setting comments with a static value."""
    doc = Document()
    doc.add_paragraph("Test content")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = Path(f.name)

    try:
        doc.save(temp_path)

        # Set comments to a static value
        operations = [
            {
                "op": "set_comments",
                "value": "This is a test template"
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True, "Set comments operation should return True"

        # Verify comments were set
        doc = Document(temp_path)
        assert doc.core_properties.comments == "This is a test template"

    finally:
        temp_path.unlink(missing_ok=True)


def test_set_comments_with_filename_placeholder():
    """Test setting comments with {{FILENAME}} placeholder."""
    doc = Document()
    doc.add_paragraph("Test content")

    import os
    temp_dir = tempfile.gettempdir()
    temp_path = Path(temp_dir) / "invoice_report.docx"

    try:
        doc.save(temp_path)

        # Set comments using FILENAME placeholder
        operations = [
            {
                "op": "set_comments",
                "value": "Template: {{FILENAME}}"
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True, "Set comments operation should return True"

        # Verify comments were set to the filename
        doc = Document(temp_path)
        assert doc.core_properties.comments == "Template: invoice_report.docx"

    finally:
        temp_path.unlink(missing_ok=True)


def test_set_comments_with_basename_placeholder():
    """Test setting comments with {{BASENAME}} placeholder."""
    doc = Document()
    doc.add_paragraph("Test content")

    import os
    temp_dir = tempfile.gettempdir()
    temp_path = Path(temp_dir) / "sales_template.docx"

    try:
        doc.save(temp_path)

        # Set comments using BASENAME placeholder
        operations = [
            {
                "op": "set_comments",
                "value": "Template: {{BASENAME}}"
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True

        # Verify comments were set to basename (without extension)
        doc = Document(temp_path)
        assert doc.core_properties.comments == "Template: sales_template"

    finally:
        temp_path.unlink(missing_ok=True)


def test_set_comments_with_extension_placeholder():
    """Test setting comments with {{EXTENSION}} placeholder."""
    doc = Document()
    doc.add_paragraph("Test content")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = Path(f.name)

    try:
        doc.save(temp_path)

        # Set comments using EXTENSION placeholder
        operations = [
            {
                "op": "set_comments",
                "value": "File type: {{EXTENSION}}"
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True

        # Verify comments were set to extension
        doc = Document(temp_path)
        assert doc.core_properties.comments == "File type: docx"

    finally:
        temp_path.unlink(missing_ok=True)


def test_set_comments_with_parent_dir_placeholder():
    """Test setting comments with {{PARENT_DIR}} placeholder."""
    doc = Document()
    doc.add_paragraph("Test content")

    import os
    # Create a temporary directory structure
    temp_dir = Path(tempfile.gettempdir()) / "test_templates"
    temp_dir.mkdir(exist_ok=True)
    temp_path = temp_dir / "test_doc.docx"

    try:
        doc.save(temp_path)

        # Set comments using PARENT_DIR placeholder
        operations = [
            {
                "op": "set_comments",
                "value": "Folder: {{PARENT_DIR}}"
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True

        # Verify comments were set to parent directory name
        doc = Document(temp_path)
        assert doc.core_properties.comments == "Folder: test_templates"

    finally:
        temp_path.unlink(missing_ok=True)
        temp_dir.rmdir()


def test_set_comments_with_multiple_placeholders():
    """Test setting comments with multiple placeholders."""
    doc = Document()
    doc.add_paragraph("Test content")

    temp_dir = Path(tempfile.gettempdir()) / "templates_dir"
    temp_dir.mkdir(exist_ok=True)
    temp_path = temp_dir / "invoice_template.docx"

    try:
        doc.save(temp_path)

        # Set comments using multiple placeholders
        operations = [
            {
                "op": "set_comments",
                "value": "{{PARENT_DIR}}/{{BASENAME}}.{{EXTENSION}}"
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True

        # Verify comments were set correctly
        doc = Document(temp_path)
        assert doc.core_properties.comments == "templates_dir/invoice_template.docx"

    finally:
        temp_path.unlink(missing_ok=True)
        temp_dir.rmdir()


def test_set_comments_overwrites_existing():
    """Test that set_comments overwrites existing comments."""
    doc = Document()
    doc.add_paragraph("Test content")
    doc.core_properties.comments = "Old comments"

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = Path(f.name)

    try:
        doc.save(temp_path)

        # Verify old comments exist
        doc = Document(temp_path)
        assert doc.core_properties.comments == "Old comments"

        # Set new comments
        operations = [
            {
                "op": "set_comments",
                "value": "New comments"
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True

        # Verify comments were overwritten
        doc = Document(temp_path)
        assert doc.core_properties.comments == "New comments"

    finally:
        temp_path.unlink(missing_ok=True)


def test_set_comments_with_empty_string():
    """Test setting comments to an empty string."""
    doc = Document()
    doc.add_paragraph("Test content")
    doc.core_properties.comments = "Existing comments"

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = Path(f.name)

    try:
        doc.save(temp_path)

        # Set comments to empty string
        operations = [
            {
                "op": "set_comments",
                "value": ""
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True

        # Verify comments were cleared
        doc = Document(temp_path)
        assert doc.core_properties.comments == ""

    finally:
        temp_path.unlink(missing_ok=True)


def test_set_comments_combined_with_clear_properties():
    """Test setting comments combined with clearing other properties."""
    doc = Document()
    doc.add_paragraph("Test content")
    doc.core_properties.author = "John Doe"
    doc.core_properties.title = "Test Document"
    doc.core_properties.comments = "Old comments"

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = Path(f.name)

    try:
        doc.save(temp_path)

        # Combine clear properties and set comments
        operations = [
            {
                "op": "clear_properties",
                "properties": ["author", "title"]
            },
            {
                "op": "set_comments",
                "value": "Template: {{FILENAME}}"
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True

        # Verify properties were cleared
        doc = Document(temp_path)
        assert doc.core_properties.author == ""
        assert doc.core_properties.title == ""

        # Verify comments were set (not cleared)
        assert "Template: " in doc.core_properties.comments
        assert temp_path.name in doc.core_properties.comments

    finally:
        temp_path.unlink(missing_ok=True)


def test_set_comments_combined_with_text_replacement():
    """Test setting comments combined with text replacement."""
    doc = Document()
    doc.add_paragraph("Hello World")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = Path(f.name)

    try:
        doc.save(temp_path)

        # Combine text replacement and set comments
        operations = [
            {
                "op": "replace",
                "search": "Hello",
                "replace": "Goodbye"
            },
            {
                "op": "set_comments",
                "value": "Modified template"
            }
        ]

        updater = DocxBulkUpdater(operations)
        result = updater.modify_docx(temp_path)

        assert result == True

        # Verify text replacement worked
        doc = Document(temp_path)
        assert "Goodbye World" in doc.paragraphs[0].text

        # Verify comments were set
        assert doc.core_properties.comments == "Modified template"

    finally:
        temp_path.unlink(missing_ok=True)
