"""
Unit tests for the FormattingProcessor class.

Tests parsing and application of formatting tokens including:
- Inline formatting blocks
- Global formatting tokens
- Font properties, alignment, spacing
"""
import pytest
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from formatting import FormattingProcessor


class TestFormattingProcessor:
    """Test cases for FormattingProcessor class."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.formatter = FormattingProcessor()
        self.doc = Document()
        self.paragraph = self.doc.add_paragraph()
    
    def test_process_formatting_tokens_simple_text(self):
        """Test processing simple text without formatting."""
        text = "Simple text without formatting"
        result = self.formatter.process_formatting_tokens(text, self.paragraph)
        
        assert len(result) == 1
        assert result[0][0] == "Simple text without formatting"
        assert result[0][1] == {
            'line_break_after': False,
            'paragraph_break_after': False,
            'alignment': None,
            'font_size': None,
            'bold': None,
            'italic': None,
            'underline': None,
            'center': False,
            'space_after': None,
            'space_before': None,
            'page_break_after': False
        }
    
    def test_process_formatting_tokens_inline_bold(self):
        """Test processing inline bold formatting."""
        text = "Normal {format:bold}bold text{/format} normal"
        result = self.formatter.process_formatting_tokens(text, self.paragraph)
        
        assert len(result) == 3
        assert result[0][0] == "Normal"
        assert result[0][1] == {
            'line_break_after': False,
            'paragraph_break_after': False,
            'alignment': None,
            'font_size': None,
            'bold': None,
            'italic': None,
            'underline': None,
            'center': False,
            'space_after': None,
            'space_before': None,
            'page_break_after': False
        }
        
        assert result[1][0] == "bold text"
        assert result[1][1]['bold'] is True
        assert result[1][1]['italic'] is False
        assert result[1][1]['underline'] is False
        
        assert result[2][0] == "normal"
        assert result[2][1] == {
            'line_break_after': False,
            'paragraph_break_after': False,
            'alignment': None,
            'font_size': None,
            'bold': None,
            'italic': None,
            'underline': None,
            'center': False,
            'space_after': None,
            'space_before': None,
            'page_break_after': False
        }
    
    def test_process_formatting_tokens_multiple_inline_options(self):
        """Test processing multiple inline formatting options."""
        text = "{format:bold,italic,size14,center}formatted text{/format}"
        result = self.formatter.process_formatting_tokens(text, self.paragraph)
        
        assert len(result) == 1
        assert result[0][0] == "formatted text"
        formatting = result[0][1]
        assert formatting['bold'] is True
        assert formatting['italic'] is True
        assert formatting['font_size'] == 14
        assert formatting['alignment'] == WD_ALIGN_PARAGRAPH.CENTER
    
    def test_process_formatting_tokens_spacing(self):
        """Test processing spacing formatting options."""
        text = "{format:spaceafter12,spacebefore6}spaced text{/format}"
        result = self.formatter.process_formatting_tokens(text, self.paragraph)
        
        assert len(result) == 1
        formatting = result[0][1]
        assert formatting['space_after'] == 12
        assert formatting['space_before'] == 6
    
    def test_process_formatting_tokens_pagebreak(self):
        """Test processing pagebreak token."""
        text = "Text before pagebreak after"
        result = self.formatter.process_formatting_tokens(text, self.paragraph)
        
        assert len(result) == 2
        assert result[0][0] == "Text before"
        assert result[0][1]['page_break_after'] is True
        assert result[1][0] == "after"
        assert result[1][1]['page_break_after'] is False
    
    def test_process_formatting_tokens_linebreak(self):
        """Test processing linebreak token - should split into separate segments."""
        text = "Text before linebreak after"
        result = self.formatter.process_formatting_tokens(text, self.paragraph)
        
        assert len(result) == 2
        assert result[0][0] == "Text before"
        assert result[0][1]['line_break_after'] is True
        assert result[1][0] == "after"
        assert result[1][1]['line_break_after'] is False
    
    def test_process_formatting_tokens_paragraphbreak(self):
        """Test processing paragraphbreak token - should split into separate segments."""
        text = "Text before paragraphbreak after"
        result = self.formatter.process_formatting_tokens(text, self.paragraph)
        
        assert len(result) == 2
        assert result[0][0] == "Text before"
        assert result[0][1]['paragraph_break_after'] is True
        assert result[1][0] == "after"
        assert result[1][1]['paragraph_break_after'] is False
    
    def test_process_formatting_tokens_multiple_paragraphbreaks(self):
        """Test processing multiple paragraphbreak tokens."""
        text = "Line1paragraphbreakLine2paragraphbreakLine3"
        result = self.formatter.process_formatting_tokens(text, self.paragraph)
        
        assert len(result) == 3
        assert result[0][0] == "Line1"
        assert result[0][1]['paragraph_break_after'] is True
        assert result[1][0] == "Line2"
        assert result[1][1]['paragraph_break_after'] is True
        assert result[2][0] == "Line3"
        assert result[2][1]['paragraph_break_after'] is False
    
    def test_process_formatting_tokens_paragraphbreak_with_inline_formatting(self):
        """Test paragraphbreak with inline formatting blocks."""
        text = "{format:center,size12}Photo1paragraphbreakPhoto2{/format}"
        result = self.formatter.process_formatting_tokens(text, self.paragraph)
        
        assert len(result) == 2
        assert result[0][0] == "Photo1"
        assert result[0][1]['paragraph_break_after'] is True
        assert result[0][1]['alignment'] == WD_ALIGN_PARAGRAPH.CENTER
        assert result[0][1]['font_size'] == 12
        assert result[1][0] == "Photo2"
        assert result[1][1]['paragraph_break_after'] is False
        assert result[1][1]['alignment'] == WD_ALIGN_PARAGRAPH.CENTER
        assert result[1][1]['font_size'] == 12
    
    def test_parse_format_options_all_alignments(self):
        """Test parsing all alignment options."""
        alignments = {
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        
        for option, expected_alignment in alignments.items():
            result = self.formatter._parse_format_options(option)
            assert result['alignment'] == expected_alignment
    
    def test_parse_format_options_font_formatting(self):
        """Test parsing font formatting options."""
        result = self.formatter._parse_format_options("bold,italic,underline")
        assert result['bold'] is True
        assert result['italic'] is True
        assert result['underline'] is True
    
    def test_parse_format_options_size_variations(self):
        """Test parsing various size format options."""
        for size in [8, 10, 12, 14, 16, 18, 24]:
            result = self.formatter._parse_format_options(f"size{size}")
            assert result['font_size'] == size
    
    def test_clean_formatting_tokens(self):
        """Test cleaning global formatting tokens from text."""
        test_cases = [
            ("Text pagebreak more", "Text more"),
            ("Text linebreak more", "Text more"),
            ("Text paragraphbreak more", "Text more"),
            ("PAGEBREAK", ""),
            ("Multiple  spaces   cleaned", "Multiple spaces cleaned"),
            ("  Leading and trailing  ", "Leading and trailing")
        ]
        
        for input_text, expected in test_cases:
            result = self.formatter._clean_formatting_tokens(input_text)
            assert result == expected
    
    def test_apply_formatting_to_run_font_properties(self):
        """Test applying font formatting to runs."""
        formatting = {
            'font_size': 14,
            'bold': True,
            'italic': False,
            'underline': True
        }
        
        run = self.paragraph.add_run("test text")
        self.formatter.apply_formatting_to_run(run, formatting, self.paragraph)
        
        assert run.font.size == Pt(14)
        assert run.font.bold is True
        assert run.font.italic is False
        assert run.font.underline is True
    
    def test_apply_formatting_to_run_breaks(self):
        """Test applying break formatting to runs."""
        run = self.paragraph.add_run("test text")
        initial_runs = len(self.paragraph.runs)
        
        # Test line break
        formatting = {'line_break_after': True}
        self.formatter.apply_formatting_to_run(run, formatting, self.paragraph)
        assert len(self.paragraph.runs) == initial_runs + 1  # New run added for break
        
        # Test paragraph break (should NOT create a run, handled at higher level)
        paragraph2 = self.doc.add_paragraph()
        run2 = paragraph2.add_run("test")
        initial_runs2 = len(paragraph2.runs)
        formatting = {'paragraph_break_after': True}
        self.formatter.apply_formatting_to_run(run2, formatting, paragraph2)
        assert len(paragraph2.runs) == initial_runs2  # No new run for paragraph break
        
        # Test page break
        paragraph3 = self.doc.add_paragraph()
        run3 = paragraph3.add_run("test")
        initial_runs3 = len(paragraph3.runs)
        formatting = {'page_break_after': True}
        self.formatter.apply_formatting_to_run(run3, formatting, paragraph3)
        assert len(paragraph3.runs) == initial_runs3 + 1  # New run added for break
    
    def test_apply_paragraph_formatting(self):
        """Test applying paragraph-level formatting."""
        formatting = {
            'alignment': WD_ALIGN_PARAGRAPH.CENTER,
            'space_after': 12,
            'space_before': 6
        }
        
        self.formatter.apply_paragraph_formatting(self.paragraph, formatting)
        
        assert self.paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert self.paragraph.paragraph_format.space_after == Pt(12)
        assert self.paragraph.paragraph_format.space_before == Pt(6)
    
    def test_complex_formatting_combination(self):
        """Test complex combination of formatting tokens."""
        text = "Start {format:bold,size16}Bold Large{/format} middle linebreak end"
        result = self.formatter.process_formatting_tokens(text, self.paragraph)
        
        # Should have segments for: "Start ", "Bold Large", " middle ", " end"
        assert len(result) >= 3
        
        # Find the bold segment
        bold_segment = None
        for text_part, formatting in result:
            if formatting.get('bold'):
                bold_segment = (text_part, formatting)
                break
        
        assert bold_segment is not None
        assert bold_segment[0] == "Bold Large"
        assert bold_segment[1]['bold'] is True
        assert bold_segment[1]['font_size'] == 16
    
    def test_empty_text_handling(self):
        """Test handling of empty text."""
        result = self.formatter.process_formatting_tokens("", self.paragraph)
        assert len(result) == 1
        assert result[0][0] == ""
        assert result[0][1] == {}
    
    def test_nested_formatting_not_supported(self):
        """Test that nested formatting blocks are not specially handled."""
        text = "{format:bold}outer {format:italic}inner{/format} outer{/format}"
        result = self.formatter.process_formatting_tokens(text, self.paragraph)
        
        # Should treat outer format block properly but inner format is literal text
        assert len(result) == 2
        assert result[0][0] == "outer {format:italic}inner"
        assert result[0][1]['bold'] is True
        assert result[1][0] == "outer{/format}"
        assert result[1][1]['bold'] is None