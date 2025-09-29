"""
Text replacement logic for DOCX documents.

Handles complex text replacement across DOCX runs while preserving formatting,
including alignment changes that require creating new paragraphs.
"""
from __future__ import annotations
import re
import logging
import xml.etree.ElementTree as ET
from typing import List, Dict, Optional, Tuple
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls
from docx.text.paragraph import Paragraph

from formatting import FormattingProcessor
from font_utils import FontFormatter


class TextReplacer:
    """Handles text replacement operations in DOCX paragraphs."""
    
    def __init__(self, operations: List[Dict[str, str]], formatting_processor: FormattingProcessor):
        # Keep only replace-like operations for this component
        self.operations = [op for op in operations if op.get('op') in ('replace', 'xml_replace')]
        self.text_ops = [op for op in self.operations if op.get('op') == 'replace']
        self.xml_ops = [op for op in self.operations if op.get('op') == 'xml_replace']
        self.formatter = formatting_processor
        # Cache compiled regex patterns for performance (for text ops only)
        self._compiled_patterns = {}
        self._precompile_patterns()
        # Cache page break information per paragraph to avoid repeated expensive checks
        self._page_break_cache = {}
        # Cache search patterns for quick lookup (only text ops)
        self._search_patterns = self._extract_search_patterns()
        self._search_patterns_set = set(self._search_patterns) if self._search_patterns else set()
        
        # Additional caches to reduce repeated extractions
        self._text_cache = {}  # Cache paragraph texts
        self._xml_cache = {}   # Cache paragraph XML to reduce xpath calls
        # Cache paragraph page break information
        self._paragraph_has_page_breaks_cache = {}
        
    def _extract_search_patterns(self) -> List[str]:
        """Extract all search patterns for quick lookup."""
        patterns = []
        for op in self.text_ops:
            if 'search' in op and 'replace' in op:
                patterns.append(op['search'])
        return patterns
        self._paragraph_has_page_breaks_cache = {}
    
    def _precompile_patterns(self):
        """Pre-compile regex patterns for all replacements to improve performance."""
        for i, op in enumerate(self.text_ops):
            if 'search' not in op:
                continue
            search_text = op['search']
            use_regex = bool(op.get('regex'))
            pattern = re.compile(search_text if use_regex else re.escape(search_text))
            self._compiled_patterns[i] = pattern
    
    def clear_caches(self):
        """Clear all caches to free memory."""
        self._page_break_cache.clear()
        self._text_cache.clear()
        self._xml_cache.clear()
    
    
    def replace_text_across_paragraphs(self, paragraphs: List[Paragraph]) -> bool:
        """Handle text replacement across multiple consecutive paragraphs."""
        if not paragraphs:
            return False
        
        # Find which paragraphs actually contain parts of the search patterns (text ops only)
        for op in self.text_ops:
            if 'search' not in op or 'replace' not in op:
                continue
            search_text = op['search']
            
            # Quick check: does the pattern exist across paragraphs?
            if not self._pattern_spans_paragraphs(paragraphs, search_text):
                continue
                
            # Find the paragraphs involved in this cross-paragraph pattern
            affected_paragraphs = self._find_affected_paragraphs(paragraphs, search_text)
            if not affected_paragraphs:
                continue
            
            # Apply replacement to the combined text
            if self._apply_cross_paragraph_replacement(paragraphs, affected_paragraphs, op):
                return True
        
        return False
    
    def _is_valid_replacement(self, replacement: Dict) -> bool:
        """Deprecated: kept for compatibility; not used in operations mode."""
        return ('search' in replacement and 'replace' in replacement)
    
    def _pattern_spans_paragraphs(self, paragraphs: List[Paragraph], search_text: str) -> bool:
        """Check if pattern spans across multiple paragraphs."""
        combined_text = "".join(para.text for para in paragraphs)
        if search_text not in combined_text:
            return False
            
        # Check if this pattern actually spans paragraphs
        for para in paragraphs:
            if search_text in para.text:
                return False  # Found in single paragraph
        return True
    
    def _find_affected_paragraphs(self, paragraphs: List[Paragraph], search_text: str) -> List[int]:
        """Find paragraph indices that contain parts of the search pattern."""
        # Find where the pattern starts
        start_idx = self._find_pattern_start(paragraphs, search_text)
        if start_idx is None:
            return []
        
        # Find consecutive paragraphs until we have the complete pattern
        affected_paragraphs = []
        accumulated_text = ""
        
        for i in range(start_idx, len(paragraphs)):
            accumulated_text += paragraphs[i].text
            affected_paragraphs.append(i)
            
            # Check if we now have the complete search pattern
            if search_text in accumulated_text:
                break
                
        return affected_paragraphs
    
    def _find_pattern_start(self, paragraphs: List[Paragraph], search_text: str) -> Optional[int]:
        """Find the paragraph where the cross-paragraph pattern starts."""
        for i, para in enumerate(paragraphs):
            para_text = para.text
            if not para_text:
                continue
                
            # Check if this paragraph contains the beginning of our search text
            if search_text.startswith(para_text[:50]):  # Check first 50 chars
                return i
                
            # Also check for partial matches at the end of paragraph
            for j in range(1, min(len(para_text), len(search_text)) + 1):
                if search_text.startswith(para_text[-j:]):
                    return i
        return None
    
    def _apply_cross_paragraph_replacement(self, paragraphs: List[Paragraph], 
                                         affected_indices: List[int], 
                                         replacement: Dict) -> bool:
        """Apply replacement to cross-paragraph text and update paragraphs."""
        # Combine text only from affected paragraphs
        combined_text = "".join(paragraphs[i].text for i in affected_indices)
        
        # Apply the replacement
        new_text, modified = self.apply_text_replacements(combined_text)
        
        if not modified:
            return False
        
        # Put the new text in the first affected paragraph
        first_paragraph = paragraphs[affected_indices[0]]
        self._rebuild_paragraph_with_text(first_paragraph, new_text)
        
        # Clear the remaining affected paragraphs
        for para_idx in affected_indices[1:]:
            self._clear_paragraph(paragraphs[para_idx])
        
        return True
    
    def _rebuild_paragraph_with_text(self, paragraph: Paragraph, new_text: str, preserve_advanced_formatting: bool = False):
        """Rebuild paragraph with new text while preserving formatting.
        
        Args:
            paragraph: The paragraph to rebuild
            new_text: The new text content
            preserve_advanced_formatting: If True, uses advanced formatting preservation (for single-paragraph replacements)
        """
        if preserve_advanced_formatting:
            self._rebuild_paragraph_advanced(paragraph, new_text)
        else:
            self._rebuild_paragraph_basic(paragraph, new_text)
    
    def _rebuild_paragraph_basic(self, paragraph: Paragraph, new_text: str):
        """Basic paragraph rebuilding for cross-paragraph replacements."""
        # Store original formatting from first run if available
        original_font_formatting = FontFormatter.get_base_font_formatting(paragraph.runs)
        
        # Clear all runs
        self._clear_paragraph(paragraph)
        
        # Process formatting tokens in the new text
        text_segments = self.formatter.process_formatting_tokens(new_text, paragraph)
        
        # Add runs with the new text and formatting
        for text, formatting in text_segments:
            if text:  # Only create runs for non-empty text
                run = paragraph.add_run(text)
                
                # Apply original formatting as base
                FontFormatter.apply_font_properties(run, original_font_formatting)
                
                # Apply new formatting from tokens
                self.formatter.apply_formatting_to_run(run, formatting, paragraph)
    
    def _rebuild_paragraph_advanced(self, paragraph: Paragraph, new_text: str):
        """Advanced paragraph rebuilding with sophisticated formatting preservation."""
        # Extract formatting information before modifying paragraph
        formatting_context = self._extract_formatting_context(paragraph)
        
        # Clear paragraph content while preserving structure
        self._clear_paragraph_preserving_structure(paragraph)
        
        # Process the new text for formatting tokens
        text_segments = self.formatter.process_formatting_tokens(new_text, paragraph)
        
        # Apply the segments based on their formatting requirements
        if self._requires_special_handling(text_segments):
            self._handle_alignment_segments(paragraph, text_segments, 
                                          formatting_context['first_run'], 
                                          formatting_context['leading_whitespace'],
                                          formatting_context['run_formats'])
        else:
            self._apply_text_segments_to_paragraph(paragraph, text_segments, 
                                                 formatting_context['run_formats'], 
                                                 formatting_context['leading_whitespace'])
    
    def _extract_formatting_context(self, paragraph: Paragraph) -> Dict:
        """Extract formatting context from paragraph before modification."""
        original_runs = list(paragraph.runs)
        original_formatting = []
        leading_whitespace_runs = []

        # Extract run formatting
        for run in original_runs:
            formatting = FontFormatter.extract_font_properties(run)
            original_formatting.append(formatting)

        # Find leading whitespace runs
        for run in original_runs:
            if run.text and all(c in '\n \t' for c in run.text):
                leading_whitespace_runs.append(run.text)
            else:
                break

        return {
            'original_runs': original_runs,
            'run_formats': original_formatting,
            'leading_whitespace': leading_whitespace_runs,
            'first_run': original_runs[0] if original_runs else None
        }


    def _clear_paragraph_preserving_structure(self, paragraph: Paragraph):
        """Clear paragraph content while preserving the first run structure."""
        # Clear all run text
        for run in paragraph.runs:
            run.text = ''

        # Remove all but the first run
        while len(paragraph.runs) > 1:
            last_run = paragraph.runs[-1]
            last_run._element.getparent().remove(last_run._element)
    
    def _requires_special_handling(self, text_segments: List[Tuple[str, Dict]]) -> bool:
        """Check if text segments require special alignment or paragraph break handling."""
        return any(seg_formatting.get('alignment') is not None or seg_formatting.get('paragraph_break_after')
                  for _, seg_formatting in text_segments)
    
    def _apply_text_segments_to_paragraph(self, paragraph: Paragraph, text_segments, original_formatting, leading_whitespace_runs):
        """Apply text segments to paragraph with sophisticated formatting preservation."""
        # Get the first run to work with
        first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

        # First, add back any leading whitespace runs
        current_run = first_run
        for i, whitespace_text in enumerate(leading_whitespace_runs):
            if i == 0:
                # Use the first run for the first whitespace
                current_run.text = whitespace_text
            else:
                # Create new runs for additional whitespace
                current_run = paragraph.add_run(whitespace_text)

            # Apply original formatting if available
            if i < len(original_formatting):
                FontFormatter.apply_font_properties(current_run, original_formatting[i])
        
        # Apply the text segments after the whitespace runs
        if text_segments:
            # Determine which run to use for the first text segment
            if leading_whitespace_runs:
                # If we have whitespace runs, create a new run for the first text segment
                first_text_run = paragraph.add_run()
                current_run = first_text_run
            else:
                # No whitespace runs, use the first run
                first_text_run = first_run
                current_run = first_text_run
            
            first_text, first_formatting = text_segments[0]
            first_text_run.text = first_text

            # Apply original formatting as base (use formatting from after whitespace)
            base_formatting_idx = len(leading_whitespace_runs)
            if base_formatting_idx < len(original_formatting):
                FontFormatter.apply_font_properties(first_text_run, original_formatting[base_formatting_idx])

            # Apply segment-specific formatting
            if first_formatting:
                self.formatter.apply_formatting_to_run(first_text_run, first_formatting, paragraph)
            
            # Create additional runs for remaining segments
            for i, (segment_text, segment_formatting) in enumerate(text_segments[1:], 1):
                if segment_text:
                    new_run = paragraph.add_run(segment_text)
                    
                    # Apply base formatting from original runs if available
                    base_idx = min(base_formatting_idx + i, len(original_formatting) - 1)
                    if base_idx < len(original_formatting):
                        FontFormatter.apply_font_properties(new_run, original_formatting[base_idx])
                    
                    # Apply segment-specific formatting
                    if segment_formatting:
                        self.formatter.apply_formatting_to_run(new_run, segment_formatting, paragraph)
    
    def _clear_paragraph(self, paragraph: Paragraph):
        """Clear all runs from a paragraph."""
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)

    
    def _has_page_break_in_run(self, run) -> bool:
        """Check if a run contains a page break."""
        # Check for hard page breaks (manual page breaks) in the XML
        return 'w:br' in run._element.xml and 'type="page"' in run._element.xml
    
    def _detect_page_breaks_after_text(self, paragraph, search_text: str) -> bool:
        """Check if there's a page break after specific text in a paragraph (optimized with caching)."""
        # Create cache key for this paragraph
        para_id = id(paragraph)
        
        # First check if this paragraph has any page breaks at all (cached)
        if para_id not in self._paragraph_has_page_breaks_cache:
            has_any_breaks = any(self._has_page_break_in_run(run) for run in paragraph.runs)
            self._paragraph_has_page_breaks_cache[para_id] = has_any_breaks
        
        # Early exit if no page breaks in entire paragraph
        if not self._paragraph_has_page_breaks_cache[para_id]:
            return False
            
        # Cache key for this specific text search
        cache_key = (para_id, search_text)
        if cache_key in self._page_break_cache:
            return self._page_break_cache[cache_key]
        
        # Get paragraph text once
        full_text = paragraph.text
        if search_text not in full_text:
            self._page_break_cache[cache_key] = False
            return False
        
        # Find the position of the search text (optimized)
        search_pos = full_text.find(search_text)
        search_end = search_pos + len(search_text)
        
        # Build run position map once per paragraph (cached)
        run_positions_key = (para_id, 'run_positions')
        if run_positions_key not in self._page_break_cache:
            run_positions = []
            char_pos = 0
            for i, run in enumerate(paragraph.runs):
                run_start = char_pos
                run_end = char_pos + len(run.text)
                run_positions.append((i, run_start, run_end))
                char_pos = run_end
            self._page_break_cache[run_positions_key] = run_positions
        
        run_positions = self._page_break_cache[run_positions_key]
        
        # Find which runs come after the search text
        result = False
        for i, run_start, run_end in run_positions:
            if run_start <= search_end <= run_end:
                # Check this run and subsequent runs for page breaks
                for j in range(i, len(paragraph.runs)):
                    if self._has_page_break_in_run(paragraph.runs[j]):
                        result = True
                        break
                break
        
        # Cache the result
        self._page_break_cache[cache_key] = result
        return result
    

    def _is_text_in_hyperlink(self, paragraph, search_text: str) -> bool:
        """Check if the search text is within a hyperlink in the paragraph (coarse check)."""
        if paragraph is None:
            return False
        if 'hyperlink' not in paragraph._p.xml.lower():
            return False
        try:
            xml_str = paragraph._p.xml
            root = ET.fromstring(xml_str)
            hyperlinks = root.findall('.//w:hyperlink', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            for hyperlink in hyperlinks:
                text_elements = hyperlink.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                hyperlink_text = ''.join(elem.text or '' for elem in text_elements)
                if search_text in hyperlink_text:
                    return True
        except Exception:
            pass
        return False

    def _compute_hyperlink_ranges(self, paragraph: Paragraph) -> List[Tuple[int, int]]:
        """Compute [start, end) character ranges of paragraph.text that are inside hyperlinks."""
        ranges: List[Tuple[int, int]] = []
        try:
            p_el = paragraph._p
            if 'hyperlink' not in p_el.xml.lower():
                return ranges
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            def sum_text_len(el) -> int:
                total_len = 0
                # XPath over lxml element to gather all w:t descendants
                for t in el.xpath('.//w:t', namespaces=ns):
                    total_len += len(t.text or '')
                return total_len

            current_index = 0
            for child in p_el.iterchildren():
                tag = str(child.tag)
                if tag.endswith('}hyperlink'):
                    hl_len = sum_text_len(child)
                    if hl_len > 0:
                        ranges.append((current_index, current_index + hl_len))
                    current_index += hl_len
                else:
                    current_index += sum_text_len(child)
        except Exception:
            return []
        return ranges

    def _match_overlaps_hyperlink(self, paragraph: Optional[Paragraph], pattern_text: str, match_start: int, match_end: int) -> bool:
        """Check if a specific match span overlaps a hyperlink span in paragraph."""
        if paragraph is None:
            return False
        ranges = self._compute_hyperlink_ranges(paragraph)
        if not ranges:
            return False
        for hs, he in ranges:
            if match_start < he and match_end > hs:
                return True
        return False

    def _replace_text_in_hyperlinks(self, paragraph, xml_str: str = None) -> bool:
        """Replace text within hyperlink elements while preserving XML structure (tabs, formatting, etc)."""
        modified = False
        if xml_str is None:
            xml_str = paragraph._p.xml

        # Only proceed if we have text operations
        if not self.text_ops:
            return False

        try:
            root = ET.fromstring(xml_str)

            # Find hyperlink elements
            hyperlinks = root.findall('.//w:hyperlink', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

            for hyperlink in hyperlinks:
                # Reconstruct text from all text elements to get full hyperlink content
                text_elements = hyperlink.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                full_text = ''.join(elem.text or '' for elem in text_elements)

                # Apply replacements to the full text using text ops
                new_full_text, text_modified = self.apply_text_replacements(full_text, None)

                if text_modified:
                    # Replace text content while preserving XML structure
                    # Find which text elements need to be updated
                    old_text_parts = [elem.text or '' for elem in text_elements]

                    # Apply the same replacements to each text part individually
                    new_text_parts = []
                    for part in old_text_parts:
                        if part:  # Only process non-empty parts
                            new_part, _ = self.apply_text_replacements(part, None)
                            new_text_parts.append(new_part)
                        else:
                            new_text_parts.append(part)

                    # Update the text elements with the new text parts
                    for elem, new_part in zip(text_elements, new_text_parts):
                        elem.text = new_part

                    modified = True

            if modified:
                # Replace paragraph XML with updated XML
                new_xml_str = ET.tostring(root, encoding='unicode')
                # Parse and replace the paragraph element
                new_p_element = parse_xml(new_xml_str)
                old_p_element = paragraph._p
                old_p_element.getparent().replace(old_p_element, new_p_element)

        except Exception as e:
            # If XML processing fails, fall back to normal text replacement
            pass

        return modified

    def _replace_xml_in_paragraph(self, paragraph) -> bool:
        """Replace raw XML patterns in paragraph while preserving document structure."""
        # Get all XML mode operations
        if not self.xml_ops:
            return False

        modified = False
        paragraph_xml = paragraph._p.xml

        # Skip if paragraph XML is None or empty
        if not paragraph_xml:
            return False

        new_xml = paragraph_xml

        try:
            for op in self.xml_ops:
                search_pattern = op.get('search')
                replace_pattern = op.get('replace', '')

                if not search_pattern:
                    continue

                # XML mode: only literal replacements are supported for safety
                new_xml_temp = new_xml.replace(search_pattern, replace_pattern)

                if new_xml_temp != new_xml:
                    new_xml = new_xml_temp
                    modified = True

            if modified:
                # Validate the new XML is well-formed by parsing it
                try:
                    # Use parse_xml from docx library to handle namespaces properly
                    new_p_element = parse_xml(new_xml)
                except Exception as e:
                    # If the resulting XML is malformed, don't apply the change
                    raise ValueError(f"XML replacement would create malformed XML: {e}")

                # Replace the paragraph with the new XML
                old_p_element = paragraph._p
                parent = old_p_element.getparent()
                if parent is not None:
                    parent.replace(old_p_element, new_p_element)
                else:
                    # If no parent, try to update the element in place
                    logging.getLogger(__name__).warning(f"Paragraph has no parent element, skipping XML replacement")
                    return False

        except Exception as e:
            # Log the error but don't fail the entire operation
            logging.getLogger(__name__).warning(f"XML replacement failed: {e}")
            return False

        return modified

    def apply_text_replacements(self, text: str, paragraph=None) -> tuple[str, bool]:
        """Apply text replacements to a string. Returns (new_text, modified)."""
        new_text = text
        modified = False
        
        # Apply all text operations to the full text
        for i, op in enumerate(self.text_ops):
            search_text = op['search']
            # Use pre-compiled pattern if available, otherwise compile on-the-fly (fallback)
            pattern = self._compiled_patterns.get(i)
            if pattern is None:
                use_regex = bool(op.get('regex'))
                pattern = re.compile(search_text if use_regex else re.escape(search_text))
            
            # Handle regular replace operation
            replace_text = op['replace']
            
            # Check if there's an existing page break after this search text (if paragraph provided)
            if paragraph is not None:
                has_existing_pagebreak = self._detect_page_breaks_after_text(paragraph, search_text)
                
                # If there's an existing page break and the replacement doesn't already include "pagebreak",
                # automatically add it to preserve the existing break
                if has_existing_pagebreak and 'pagebreak' not in replace_text.lower():
                    replace_text = replace_text + 'pagebreak'
            
            # Replace all non-hyperlink matches
            matches = list(pattern.finditer(new_text))
            replacements_made = 0
            
            # Process matches in reverse order to avoid offset issues
            for match in reversed(matches):
                start, end = match.start(), match.end()
                # Check if this match is in a hyperlink (if paragraph available)
                if paragraph and self._match_overlaps_hyperlink(paragraph, search_text, start, end):
                    continue  # Skip this match, try the next one
                
                # Replace this match
                new_text = new_text[:start] + replace_text + new_text[end:]
                replacements_made += 1
            
            if replacements_made > 0:
                modified = True
                
        return new_text, modified


    def replace_text_in_paragraph(self, paragraph) -> bool:
        """Replace text in a paragraph, handling splits across runs while preserving formatting."""
        # Use cached text to avoid repeated .text property calls
        para_id = id(paragraph)
        if para_id in self._text_cache:
            full_text = self._text_cache[para_id]
        else:
            full_text = paragraph.text
            self._text_cache[para_id] = full_text

        # Check for XML replacements first (they take precedence)
        xml_modified = self._replace_xml_in_paragraph(paragraph)
        if xml_modified:
            # Clear cache since paragraph was modified
            self._text_cache.pop(para_id, None)
            self._xml_cache.pop(para_id, None)
            return True

        # Early exit if paragraph is empty
        if not full_text.strip():
            return False

        # Quick check: does this paragraph contain any of our search patterns?
        has_any_matches = any(pattern in full_text for pattern in self._search_patterns_set)
        if not has_any_matches:
            return False

        # Check if paragraph contains hyperlinks that need text replacement (optimized)
        if para_id in self._xml_cache:
            paragraph_xml = self._xml_cache[para_id]
        else:
            paragraph_xml = paragraph._p.xml
            self._xml_cache[para_id] = paragraph_xml

        has_hyperlinks = 'hyperlink' in paragraph_xml.lower()
        if has_hyperlinks:
            # Handle hyperlink text replacement directly in XML
            modified = self._replace_text_in_hyperlinks(paragraph, paragraph_xml)
            if modified:
                # Clear cache since paragraph was modified
                self._text_cache.pop(para_id, None)
                self._xml_cache.pop(para_id, None)
                return True

        # Apply text replacements using the extracted method
        new_text, modified = self.apply_text_replacements(full_text, paragraph)

        if not modified:
            return False

        # Now we need to rebuild the paragraph with the new text
        # but preserve formatting where possible

        # Use unified paragraph rebuilding with advanced formatting preservation
        self._rebuild_paragraph_with_text(paragraph, new_text, preserve_advanced_formatting=True)

        # Clear cache since paragraph was modified
        self._text_cache.pop(para_id, None)
        self._xml_cache.pop(para_id, None)

        return True

    def _handle_alignment_segments(self, paragraph, new_run_data: List[Tuple[str, Dict]],
                                 original_run, leading_whitespace_runs: Optional[List[str]] = None,
                                 original_formatting: Optional[List[Dict]] = None):
        """Handle segments with different alignments or paragraph breaks by creating separate paragraphs if needed."""
        # Group segments by alignment and paragraph breaks
        segment_groups = []
        current_group = []
        current_alignment = None
        
        for text_segment, segment_formatting in new_run_data:
            seg_alignment = segment_formatting.get('alignment')

            # Add segment to current group
            current_group.append((text_segment, segment_formatting))

            # Check if this segment forces a new paragraph (alignment change or paragraph break)
            alignment_changed = seg_alignment != current_alignment and seg_alignment is not None
            has_paragraph_break = segment_formatting.get('paragraph_break_after', False)

            if alignment_changed or has_paragraph_break:
                # Finalize current group - use segment alignment if current is None
                group_alignment = seg_alignment if current_alignment is None else current_alignment
                segment_groups.append((group_alignment, current_group))
                current_group = []
                current_alignment = seg_alignment
            elif seg_alignment is not None:
                # Update current alignment without breaking
                current_alignment = seg_alignment
        
        if current_group:
            segment_groups.append((current_alignment, current_group))
        
        # Handle the first group in the existing paragraph
        if segment_groups:
            first_alignment, first_group = segment_groups[0]
            
            # First, add any leading whitespace runs
            if leading_whitespace_runs and original_formatting:
                for i, whitespace_text in enumerate(leading_whitespace_runs):
                    if i == 0 and original_run is not None:
                        # Use the original run for the first whitespace
                        original_run.text = whitespace_text
                        # Apply original formatting
                        if i < len(original_formatting):
                            FontFormatter.apply_font_properties(original_run, original_formatting[i])
                    else:
                        # Create new runs for additional whitespace
                        ws_run = paragraph.add_run(whitespace_text)
                        if i < len(original_formatting):
                            FontFormatter.apply_font_properties(ws_run, original_formatting[i])
            else:
                # No leading whitespace, clear the original run
                if original_run is not None:
                    original_run.text = ''
            
            # Add the text segments
            for i, (text_segment, segment_formatting) in enumerate(first_group):
                if text_segment:  # Only add non-empty text
                    if not leading_whitespace_runs and i == 0 and original_run is not None:
                        # No whitespace runs, use original run for first text segment
                        original_run.text = text_segment
                        if segment_formatting:
                            self.formatter.apply_formatting_to_run(original_run, segment_formatting, paragraph)
                    else:
                        # Create new run for text segment (either because we have whitespace runs, not first segment, or original_run is None)
                        new_run = paragraph.add_run(text_segment) 
                        # Apply base formatting
                        base_idx = len(leading_whitespace_runs or []) + i
                        if original_formatting and base_idx < len(original_formatting):
                            FontFormatter.apply_font_properties(new_run, original_formatting[base_idx])
                        
                        if segment_formatting:
                            self.formatter.apply_formatting_to_run(new_run, segment_formatting, paragraph)
            
            # Apply alignment to first paragraph
            if first_alignment is not None:
                # Use formatter's method to handle table cell alignment properly
                self.formatter.apply_paragraph_formatting(paragraph, {'alignment': first_alignment})
            
            # Create new paragraphs for remaining groups
            for alignment, group in segment_groups[1:]:
                # Create new paragraph after current one
                new_p_el = OxmlElement("w:p")
                paragraph._p.addnext(new_p_el)
                new_paragraph = Paragraph(new_p_el, paragraph._parent)
                
                # Set alignment for new paragraph
                if alignment is not None:
                    # Use formatter's method to handle table cell alignment properly
                    self.formatter.apply_paragraph_formatting(new_paragraph, {'alignment': alignment})
                
                # Add runs to new paragraph
                for text_segment, segment_formatting in group:
                    if text_segment:
                        new_run = new_paragraph.add_run(text_segment)
                        # Copy base formatting if original_run exists
                        if original_run is not None:
                            FontFormatter.copy_font_formatting(original_run, new_run)
                        
                        if segment_formatting:
                            self.formatter.apply_formatting_to_run(new_run, segment_formatting, new_paragraph)
