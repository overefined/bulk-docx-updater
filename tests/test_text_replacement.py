"""
Unit tests for the TextReplacer class.

Tests initialization, text replacement functionality, and paragraph break handling.
Complex DOCX document tests are in test_real_templates.py and test_integration.py.
"""
import pytest
from docx import Document
from docx.oxml import OxmlElement, ns
from src.text_replacement import TextReplacer
from src.formatting import FormattingProcessor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class TestTextReplacer:
    """Test cases for TextReplacer class."""

    def setup_method(self):
        """Set up test fixtures."""
        self.formatter = FormattingProcessor()
        self.operations = [
            {"op": "replace", "search": "old text", "replace": "new text"},
            {"op": "replace", "search": "TESTER QUALIFICATIONS", "replace": "INSPECTOR QUALIFICATIONS"}
        ]

    def test_init(self):
        """Test TextReplacer initialization."""
        text_replacer = TextReplacer(self.operations, self.formatter)

        # TextReplacer filters to only keep replace/xml_replace operations
        assert len(text_replacer.operations) == 2
        assert text_replacer.formatter == self.formatter

    def test_init_with_empty_replacements(self):
        """Test TextReplacer initialization with empty operations."""
        text_replacer = TextReplacer([], self.formatter)

        assert text_replacer.operations == []
        assert text_replacer.formatter == self.formatter


class TestTextReplacementFunctionality:
    """Test text replacement operations."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.formatter = FormattingProcessor()
    
    def test_apply_text_replacements_simple(self):
        """Test simple text replacement."""
        operations = [{"op": "replace", "search": "old", "replace": "new"}]
        replacer = TextReplacer(operations, self.formatter)
        
        result, modified = replacer.apply_text_replacements("This is old text")
        assert modified is True
        assert result == "This is new text"
    
    def test_apply_text_replacements_no_match(self):
        """Test text replacement with no matches."""
        operations = [{"op": "replace", "search": "missing", "replace": "new"}]
        replacer = TextReplacer(operations, self.formatter)
        
        result, modified = replacer.apply_text_replacements("This is old text")
        assert modified is False
        assert result == "This is old text"
    
    def test_apply_text_replacements_with_paragraphbreak(self):
        """Test text replacement with paragraph breaks."""
        operations = [{"op": "replace", "search": "PHOTOS", "replace": "Photo1paragraphbreakPhoto2"}]
        replacer = TextReplacer(operations, self.formatter)
        
        result, modified = replacer.apply_text_replacements("SITE PHOTOS section")
        assert modified is True
        assert result == "SITE Photo1paragraphbreakPhoto2 section"
    
    def test_apply_text_replacements_append_after(self):
        """Test appending content using replace."""
        operations = [{"op": "replace", "search": "PHOTOS", "replace": "PHOTOS\nPhoto1paragraphbreakPhoto2"}]
        replacer = TextReplacer(operations, self.formatter)
        
        result, modified = replacer.apply_text_replacements("SITE PHOTOS section")
        assert modified is True
        assert result == "SITE PHOTOS\nPhoto1paragraphbreakPhoto2 section"
    
    def test_apply_text_replacements_multiple_replacements(self):
        """Test multiple replacements in order."""
        operations = [
            {"op": "replace", "search": "old", "replace": "new"},
            {"op": "replace", "search": "bad", "replace": "good"}
        ]
        replacer = TextReplacer(operations, self.formatter)
        
        result, modified = replacer.apply_text_replacements("This old thing is bad")
        assert modified is True
        assert result == "This new thing is good"
    
    def test_apply_text_replacements_all_occurrences(self):
        """Test that all occurrences are replaced."""
        operations = [{"op": "replace", "search": "test", "replace": "example"}]
        replacer = TextReplacer(operations, self.formatter)
        
        result, modified = replacer.apply_text_replacements("This test has test twice")
        assert modified is True
        assert result == "This example has example twice"
    
    def test_apply_text_replacements_skip_cleanup_operations(self):
        """Test that cleanup operations (cleanup_empty_after) are filtered out."""
        operations = [
            {"op": "replace", "search": "old", "replace": "new"},
            {"op": "cleanup_empty_after", "pattern": "SITE PHOTOS"}  # This should be filtered
        ]
        replacer = TextReplacer(operations, self.formatter)
        
        result, modified = replacer.apply_text_replacements("This is old text")
        assert modified is True
        assert result == "This is new text"
        # The cleanup operation should not affect text replacement
    
    def test_apply_text_replacements_skip_replacements_without_search(self):
        """Test that operations without op=replace are filtered out."""
        operations = [
            {"op": "replace", "search": "old", "replace": "new"},
            {"op": "table_header_repeat", "pattern": "Header"}  # Should be filtered
        ]
        replacer = TextReplacer(operations, self.formatter)
        
        result, modified = replacer.apply_text_replacements("This is old text")
        assert modified is True
        assert result == "This is new text"


class TestHyperlinkDetection:
    """Test hyperlink detection functionality."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.formatter = FormattingProcessor()
        self.operations = [{"op": "replace", "search": "SITE PHOTOS", "replace": "Updated Photos"}]
        self.replacer = TextReplacer(self.operations, self.formatter)
        self.doc = Document()
    
    def create_paragraph_with_hyperlink(self, text: str) -> object:
        """Create a paragraph containing a hyperlink with the given text."""
        paragraph = self.doc.add_paragraph()
        
        # Create hyperlink XML element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(ns.qn('r:id'), 'rId1')
        
        # Create run within hyperlink
        run = OxmlElement('w:r')
        run_text = OxmlElement('w:t')
        run_text.text = text
        run.append(run_text)
        hyperlink.append(run)
        
        # Add hyperlink to paragraph
        paragraph._p.append(hyperlink)
        return paragraph
    
    def test_is_text_in_hyperlink_with_hyperlinks(self):
        """Test hyperlink detection correctly identifies text within hyperlinks."""
        paragraph = self.create_paragraph_with_hyperlink("APPENDIX H    SITE PHOTOS")
        
        result = self.replacer._is_text_in_hyperlink(paragraph, "SITE PHOTOS")
        assert result is True
    
    def test_is_text_in_hyperlink_without_hyperlinks(self):
        """Test hyperlink detection when paragraph has no hyperlinks."""
        paragraph = self.doc.add_paragraph("SITE PHOTOS")
        
        result = self.replacer._is_text_in_hyperlink(paragraph, "SITE PHOTOS")
        assert result is False
    
    def test_is_text_in_hyperlink_with_none_paragraph(self):
        """Test hyperlink detection with None paragraph."""
        result = self.replacer._is_text_in_hyperlink(None, "SITE PHOTOS")
        assert result is False
    
    def test_is_text_in_hyperlink_text_not_in_paragraph(self):
        """Test hyperlink detection when search text is not in paragraph."""
        paragraph = self.create_paragraph_with_hyperlink("Different content")
        
        result = self.replacer._is_text_in_hyperlink(paragraph, "SITE PHOTOS")
        assert result is False


class TestTextReplacementWithHyperlinkSkipping:
    """Test text replacement with hyperlink detection."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.formatter = FormattingProcessor()
    
    def test_apply_text_replacements_works_without_paragraph_context(self):
        """Test that text replacement works when no paragraph context is provided."""
        operations = [{"op": "replace", "search": "SITE PHOTOS", "replace": "Updated Photos"}]
        replacer = TextReplacer(operations, self.formatter)
        
        # No paragraph context means no hyperlink detection
        result, modified = replacer.apply_text_replacements("SITE PHOTOS", None)
        assert modified is True
        assert result == "Updated Photos"


class TestTextReplacementWithDocx:
    """Test text replacement with real DOCX objects."""
    
    def setup_method(self):
        """Set up test fixtures with real DOCX objects."""
        self.formatter = FormattingProcessor()
        self.doc = Document()
        self.paragraph = self.doc.add_paragraph("Original text content")
    
    def test_append_after_with_paragraphbreaks(self):
        """Test append-after via replace with paragraph breaks using string processing."""
        # Set up operations with paragraph breaks
        operations = [
            {"op": "replace", "search": "PHOTOS", "replace": "PHOTOSPhoto1paragraphbreakPhoto2paragraphbreakPhoto3"}
        ]
        replacer = TextReplacer(operations, self.formatter)
        
        # Test the core text replacement logic
        test_text = "SITE PHOTOS content"
        
        # Execute the replacement at text level
        result, modified = replacer.apply_text_replacements(test_text)
        
        # Verify the replacement logic works correctly
        assert modified is True
        assert "SITE PHOTOS" in result  # Original text preserved
        assert "Photo1" in result
        assert "paragraphbreak" in result  # Formatting token preserved
        assert "Photo2" in result
        assert "Photo3" in result
    
    def test_append_after_no_match(self):
        """Test append-after via replace with no matching text."""
        operations = [{"op": "replace", "search": "MISSING", "replace": "MISSINGnew content"}]
        replacer = TextReplacer(operations, self.formatter)
        
        result, modified = replacer.apply_text_replacements("SITE PHOTOS content")
        assert modified is False
        assert result == "SITE PHOTOS content"
    
    def create_paragraph_with_hyperlink(self, text: str) -> object:
        """Create a paragraph containing a hyperlink with the given text."""
        paragraph = self.doc.add_paragraph()
        
        # Create hyperlink XML element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(ns.qn('r:id'), 'rId1')
        
        # Create run within hyperlink
        run = OxmlElement('w:r')
        run_text = OxmlElement('w:t')
        run_text.text = text
        run.append(run_text)
        hyperlink.append(run)
        
        # Add hyperlink to paragraph
        paragraph._p.append(hyperlink)
        return paragraph
    
    def test_hyperlink_detection_helper(self):
        """Test hyperlink detection helper returns True when match is in hyperlink."""
        operations = [{"op": "replace", "search": "PHOTOS", "replace": "PHOTOS new content"}]
        replacer = TextReplacer(operations, self.formatter)
        
        # Create paragraph with hyperlink using our helper
        paragraph = self.create_paragraph_with_hyperlink("APPENDIX H    SITE PHOTOS")
        
        # Test that hyperlink detection correctly identifies hyperlinked text
        is_hyperlinked = replacer._is_text_in_hyperlink(paragraph, "PHOTOS")
        assert is_hyperlinked is True


class TestWhitespaceFlexibleMatching:
    """Test that searches match Unicode whitespace variants (en-space, non-breaking space, etc.)."""

    def setup_method(self):
        self.formatter = FormattingProcessor()

    def test_tab_matches_en_space(self):
        """Search with tab matches en-space (\u2002) in document text."""
        operations = [{"op": "replace", "search": "hello\tworld", "replace": "replaced"}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("hello\u2002world")
        assert modified is True
        assert result == "replaced"

    def test_space_matches_non_breaking_space(self):
        """Search with regular space matches non-breaking space (\u00a0)."""
        operations = [{"op": "replace", "search": "hello world", "replace": "replaced"}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("hello\u00a0world")
        assert modified is True
        assert result == "replaced"

    def test_space_matches_em_space(self):
        """Search with regular space matches em-space (\u2003)."""
        operations = [{"op": "replace", "search": "A B", "replace": "X"}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("A\u2003B")
        assert modified is True
        assert result == "X"

    def test_tab_matches_tab(self):
        """Normal tab still matches tab."""
        operations = [{"op": "replace", "search": "col1\tcol2", "replace": "replaced"}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("col1\tcol2")
        assert modified is True
        assert result == "replaced"

    def test_multiple_whitespace_chars(self):
        """Search with mixed whitespace matches different Unicode whitespace."""
        operations = [{"op": "replace", "search": "a \t b", "replace": "X"}]
        replacer = TextReplacer(operations, self.formatter)

        # Document has en-space + non-breaking space + em-space between tokens
        result, modified = replacer.apply_text_replacements("a\u2002\u00a0\u2003b")
        assert modified is True
        assert result == "X"

    def test_no_whitespace_still_exact(self):
        """Search without whitespace still requires exact match."""
        operations = [{"op": "replace", "search": "helloworld", "replace": "X"}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("hello world")
        assert modified is False
        assert result == "hello world"

    def test_regex_mode_unaffected(self):
        """Regex mode is not altered by whitespace normalization."""
        operations = [{"op": "replace", "search": r"hello\s+world", "replace": "X", "regex": True}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("hello\u2002world")
        assert modified is True
        assert result == "X"

    def test_real_world_permit_to_operate(self):
        """Reproduce the OH template scenario: tab + en-spaces in document text."""
        operations = [{"op": "replace", "search": "permit-to-operate?             \t", "replace": "permit-to-operate? 500BHP"}]
        replacer = TextReplacer(operations, self.formatter)

        # Simulate document text with en-spaces and tab
        doc_text = "permit-to-operate?\u2002\u2002\u2002\u2002\u2002\u2002\u2002\u2002\u2002\u2002\u2002\u2002\u2002\t"
        result, modified = replacer.apply_text_replacements(doc_text)
        assert modified is True
        assert result == "permit-to-operate? 500BHP"

    def test_compile_whitespace_flexible_pattern(self):
        """Test the static pattern compiler directly."""
        pattern = TextReplacer._compile_whitespace_flexible_pattern("a\tb")
        assert pattern.search("a\tb") is not None
        assert pattern.search("a\u2002b") is not None
        assert pattern.search("a  b") is not None
        assert pattern.search("ab") is None


class TestCountAndOccurrence:
    """Test count and occurrence options for text replacement."""

    def setup_method(self):
        """Set up test fixtures."""
        self.formatter = FormattingProcessor()

    def test_count_limits_replacements(self):
        """Test that count limits the number of replacements."""
        operations = [{"op": "replace", "search": "foo", "replace": "bar", "count": 2}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("foo foo foo foo")
        assert modified is True
        assert result == "bar bar foo foo"

    def test_count_zero_replaces_all(self):
        """Test that count=0 (default) replaces all occurrences."""
        operations = [{"op": "replace", "search": "foo", "replace": "bar", "count": 0}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("foo foo foo")
        assert modified is True
        assert result == "bar bar bar"

    def test_count_one(self):
        """Test that count=1 replaces only the first occurrence."""
        operations = [{"op": "replace", "search": "Address", "replace": "123 Main St", "count": 1}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("Address\nAddress")
        assert modified is True
        assert result == "123 Main St\nAddress"

    def test_count_exceeding_matches(self):
        """Test that count greater than total matches replaces all."""
        operations = [{"op": "replace", "search": "x", "replace": "y", "count": 10}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("x x x")
        assert modified is True
        assert result == "y y y"

    def test_occurrence_first(self):
        """Test targeting the first occurrence."""
        operations = [{"op": "replace", "search": "Address", "replace": "123 Main St", "occurrence": 1}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("Address\nCity\nAddress")
        assert modified is True
        assert result == "123 Main St\nCity\nAddress"

    def test_occurrence_second(self):
        """Test targeting the second occurrence."""
        operations = [{"op": "replace", "search": "Address", "replace": "City, ST 12345", "occurrence": 2}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("Address\nCity\nAddress")
        assert modified is True
        assert result == "Address\nCity\nCity, ST 12345"

    def test_occurrence_beyond_matches(self):
        """Test targeting an occurrence that doesn't exist."""
        operations = [{"op": "replace", "search": "foo", "replace": "bar", "occurrence": 5}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("foo foo")
        assert modified is False
        assert result == "foo foo"

    def test_occurrence_with_no_match(self):
        """Test occurrence on text with no matches at all."""
        operations = [{"op": "replace", "search": "missing", "replace": "bar", "occurrence": 1}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("foo foo")
        assert modified is False
        assert result == "foo foo"

    def test_default_replaces_all(self):
        """Test that without count/occurrence, all occurrences are replaced."""
        operations = [{"op": "replace", "search": "a", "replace": "b"}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("a a a a")
        assert modified is True
        assert result == "b b b b"

    def test_sequential_occurrence_replacements(self):
        """Test using occurrence to replace identical text sequentially with different values."""
        ops = [
            {"op": "replace", "search": "Address", "replace": "123 Main St", "occurrence": 1},
            {"op": "replace", "search": "Address", "replace": "Springfield, IL 62701", "occurrence": 1},
        ]
        replacer = TextReplacer(ops, self.formatter)

        result, modified = replacer.apply_text_replacements("Address\nAddress\nAddress")
        assert modified is True
        assert result == "123 Main St\nSpringfield, IL 62701\nAddress"

    def test_count_with_regex(self):
        """Test count option combined with regex."""
        operations = [{"op": "replace", "search": r"\d+", "replace": "X", "regex": True, "count": 2}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("Item 1, Item 2, Item 3")
        assert modified is True
        assert result == "Item X, Item X, Item 3"

    def test_occurrence_with_regex(self):
        """Test occurrence option combined with regex."""
        operations = [{"op": "replace", "search": r"\d+", "replace": "99", "regex": True, "occurrence": 2}]
        replacer = TextReplacer(operations, self.formatter)

        result, modified = replacer.apply_text_replacements("Item 1, Item 2, Item 3")
        assert modified is True
        assert result == "Item 1, Item 99, Item 3"


class TestParagraphBreakIntegration:
    """Integration tests for paragraph break functionality."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.formatter = FormattingProcessor()
    
    def test_paragraphbreak_processing_pipeline(self):
        """Test complete pipeline from text replacement to formatting segments."""
        operations = [
            {"op": "replace", "search": "PHOTOS", "replace": "Photo1paragraphbreakPhoto2paragraphbreakPhoto3"}
        ]
        replacer = TextReplacer(operations, self.formatter)
        
        # Apply replacement
        text = "SITE PHOTOS section"
        new_text, modified = replacer.apply_text_replacements(text)
        assert modified is True
        assert new_text == "SITE Photo1paragraphbreakPhoto2paragraphbreakPhoto3 section"
        
        # Process formatting tokens on the replaced text
        segments = self.formatter.process_formatting_tokens(new_text, None)
        
        # Should find segments with paragraph breaks
        photo_segments = []
        for text_part, formatting in segments:
            if 'Photo' in text_part:
                photo_segments.append((text_part, formatting))
        
        # Should have Photo1 and Photo2 with paragraph breaks, Photo3 without
        assert len(photo_segments) >= 2
        
        # Find Photo1 and Photo2 segments
        photo1_segment = next((s for s in photo_segments if 'Photo1' in s[0]), None)
        photo2_segment = next((s for s in photo_segments if 'Photo2' in s[0]), None)
        
        if photo1_segment:
            assert photo1_segment[1].get('paragraph_break_after') is True
        if photo2_segment:
            assert photo2_segment[1].get('paragraph_break_after') is True
    
    def test_paragraphbreak_with_inline_formatting_integration(self):
        """Test paragraph breaks combined with inline formatting."""
        operations = [
            {"op": "replace", "search": "PHOTOS", "replace": "{format:center,size12}Photo1paragraphbreakPhoto2{/format}"}
        ]
        replacer = TextReplacer(operations, self.formatter)
        
        # Apply replacement
        text = "SITE PHOTOS section"
        new_text, modified = replacer.apply_text_replacements(text)
        assert modified is True
        
        # Process formatting
        segments = self.formatter.process_formatting_tokens(new_text, None)
        
        # Find photo segments
        photo_segments = []
        for text_part, formatting in segments:
            if 'Photo' in text_part:
                photo_segments.append((text_part, formatting))
        
        # Both photos should have the inline formatting AND paragraph break handling
        for text_part, formatting in photo_segments:
            if 'Photo1' in text_part:
                assert formatting.get('alignment') == WD_ALIGN_PARAGRAPH.CENTER
                assert formatting.get('font_size') == 12
                assert formatting.get('paragraph_break_after') is True
            elif 'Photo2' in text_part:
                assert formatting.get('alignment') == WD_ALIGN_PARAGRAPH.CENTER
                assert formatting.get('font_size') == 12
                assert formatting.get('paragraph_break_after') is False
