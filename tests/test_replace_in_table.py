"""
Unit tests for replace_in_table functionality.

Tests the replace_in_table operation that allows targeting specific
tables by heading for scoped text replacement.
"""
import pytest
import json
from pathlib import Path
from docx import Document

from src.document_processor import DocxBulkUpdater
from src.config import load_operations_from_json


class TestReplaceInTable:
    """Test cases for replace_in_table functionality."""

    def create_test_document_with_two_tables(self):
        """Create a test document with two tables containing the same placeholder."""
        doc = Document()

        # First table: O2 VALUES
        table1 = doc.add_table(rows=4, cols=4)

        # Header row
        table1.rows[0].cells[0].text = "O2 VALUES"
        table1.rows[0].cells[1].text = ""
        table1.rows[0].cells[2].text = ""
        table1.rows[0].cells[3].text = ""

        # Subheader row
        table1.rows[1].cells[0].text = "Bias Selection Values"
        table1.rows[1].cells[1].text = "Run"
        table1.rows[1].cells[2].text = "% O2 Measured"
        table1.rows[1].cells[3].text = "% O2 Adjusted"

        # Data rows with placeholder
        table1.rows[2].cells[0].text = "{{ o2_qaqc.O2Low_cal_span }}"
        table1.rows[2].cells[1].text = "Run 1"
        table1.rows[2].cells[2].text = "{{ o2_qaqc.run1_O2 }}"
        table1.rows[2].cells[3].text = "{{ o2_qaqc.run1_O2_adjusted }}"

        table1.rows[3].cells[0].text = "{{ o2_qaqc.O2Low_cal_span }}"
        table1.rows[3].cells[1].text = "Run 2"
        table1.rows[3].cells[2].text = "{{ o2_qaqc.run2_O2 }}"
        table1.rows[3].cells[3].text = "{{ o2_qaqc.run2_O2_adjusted }}"

        # Second table: OXYGEN CALIBRATION ERROR (also contains the same placeholder)
        table2 = doc.add_table(rows=6, cols=4)

        # Header row
        table2.rows[0].cells[0].text = "OXYGEN CALIBRATION ERROR"
        table2.rows[0].cells[1].text = ""
        table2.rows[0].cells[2].text = ""
        table2.rows[0].cells[3].text = ""

        # Subheader row
        table2.rows[1].cells[0].text = ""
        table2.rows[1].cells[1].text = "Low"
        table2.rows[1].cells[2].text = "Mid"
        table2.rows[1].cells[3].text = "High"

        # Data rows
        table2.rows[2].cells[0].text = "Cylinder Reference"
        table2.rows[2].cells[1].text = "{{ o2_qaqc.O2LowCylNumber }}"
        table2.rows[2].cells[2].text = "{{ o2_qaqc.O2MidCylNumber }}"
        table2.rows[2].cells[3].text = "{{ o2_qaqc.O2HighCylNumber }}"

        table2.rows[3].cells[0].text = "Calibration Gas Value (Cv)"
        table2.rows[3].cells[1].text = "{{ o2_qaqc.O2LowCylValue }}"
        table2.rows[3].cells[2].text = "{{ o2_qaqc.O2MidCylValue }}"
        table2.rows[3].cells[3].text = "{{ o2_qaqc.O2HighCylValue }}"

        table2.rows[4].cells[0].text = "Calibration Span (CS)"
        table2.rows[4].cells[1].text = "{{ o2_qaqc.O2Low_cal_span }}"  # Same placeholder as table1!
        table2.rows[4].cells[2].text = "{{ o2_qaqc.O2Mid_cal_span }}"
        table2.rows[4].cells[3].text = "{{ o2_qaqc.O2High_cal_span }}"

        table2.rows[5].cells[0].text = "Calibration Error (ACE)"
        table2.rows[5].cells[1].text = "{{ o2_qaqc.O2Low_cal_error }}"
        table2.rows[5].cells[2].text = "{{ o2_qaqc.O2Mid_cal_error }}"
        table2.rows[5].cells[3].text = "{{ o2_qaqc.O2High_cal_error }}"

        return doc

    def test_replace_in_specific_table_only(self, tmp_path):
        """Test that replacement only happens in the specified table, not in other tables."""
        # Create test document
        doc = self.create_test_document_with_two_tables()
        test_file = tmp_path / "test_replace_in_table.docx"
        doc.save(test_file)

        # Configure replacement to target only the "O2 VALUES" table
        operations = [{
            'op': 'replace_in_table',
            'table_heading': 'O2 VALUES',
            'search': '{{ o2_qaqc.O2Low_cal_span }}',
            'replace': '20.95'
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)

        # First table (O2 VALUES) should have the replacement
        table1 = modified_doc.tables[0]
        assert table1.rows[2].cells[0].text == "20.95", "First table row 2 should be replaced"
        assert table1.rows[3].cells[0].text == "20.95", "First table row 3 should be replaced"

        # Second table (OXYGEN CALIBRATION ERROR) should NOT have the replacement
        table2 = modified_doc.tables[1]
        assert table2.rows[4].cells[1].text == "{{ o2_qaqc.O2Low_cal_span }}", "Second table should be unchanged"

    def test_replace_in_second_table_only(self, tmp_path):
        """Test replacing in the second table while leaving the first unchanged."""
        # Create test document
        doc = self.create_test_document_with_two_tables()
        test_file = tmp_path / "test_replace_in_table2.docx"
        doc.save(test_file)

        # Configure replacement to target only the "OXYGEN CALIBRATION ERROR" table
        operations = [{
            'op': 'replace_in_table',
            'table_heading': 'OXYGEN CALIBRATION ERROR',
            'search': '{{ o2_qaqc.O2Low_cal_span }}',
            'replace': '15.0'
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)

        # First table (O2 VALUES) should NOT have the replacement
        table1 = modified_doc.tables[0]
        assert table1.rows[2].cells[0].text == "{{ o2_qaqc.O2Low_cal_span }}", "First table should be unchanged"
        assert table1.rows[3].cells[0].text == "{{ o2_qaqc.O2Low_cal_span }}", "First table should be unchanged"

        # Second table (OXYGEN CALIBRATION ERROR) should have the replacement
        table2 = modified_doc.tables[1]
        assert table2.rows[4].cells[1].text == "15.0", "Second table should be replaced"

    def test_table_heading_not_found(self, tmp_path):
        """Test handling when table heading is not found."""
        # Create test document
        doc = self.create_test_document_with_two_tables()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement with non-existent heading
        operations = [{
            'op': 'replace_in_table',
            'table_heading': 'NonExistent Table',
            'search': 'test',
            'replace': 'replacement'
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify no changes were made
        assert result is False

    def test_multiple_occurrences_in_table(self, tmp_path):
        """Test that all occurrences in the target table are replaced."""
        # Create test document
        doc = self.create_test_document_with_two_tables()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement
        operations = [{
            'op': 'replace_in_table',
            'table_heading': 'O2 VALUES',
            'search': 'Run',
            'replace': 'Test'
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document - should replace multiple occurrences
        modified_doc = Document(test_file)
        table1 = modified_doc.tables[0]
        assert "Test 1" in table1.rows[2].cells[1].text
        assert "Test 2" in table1.rows[3].cells[1].text

    def test_regex_in_table(self, tmp_path):
        """Test regex replacement in specific table."""
        # Create test document
        doc = self.create_test_document_with_two_tables()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement with regex
        operations = [{
            'op': 'replace_in_table',
            'table_heading': 'O2 VALUES',
            'search': r'Run \d+',
            'replace': 'Test Run',
            'regex': True
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)
        table1 = modified_doc.tables[0]
        assert table1.rows[2].cells[1].text == "Test Run"
        assert table1.rows[3].cells[1].text == "Test Run"

    def test_partial_table_heading_match(self, tmp_path):
        """Test that partial heading matches work."""
        # Create test document
        doc = self.create_test_document_with_two_tables()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement using partial heading
        operations = [{
            'op': 'replace_in_table',
            'table_heading': 'O2 VALUES',  # Should match "O2 VALUES" in first cell
            'search': '{{ o2_qaqc.O2Low_cal_span }}',
            'replace': 'MATCHED'
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)
        table1 = modified_doc.tables[0]
        assert table1.rows[2].cells[0].text == "MATCHED"

    def test_config_integration(self, tmp_path):
        """Test replace_in_table through config file integration."""
        # Create test document
        doc = self.create_test_document_with_two_tables()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Create config file
        config_data = {
            "replace_in_table": [{
                "table_heading": "O2 VALUES",
                "search": "{{ o2_qaqc.O2Low_cal_span }}",
                "replace": "20.95"
            }]
        }

        config_file = tmp_path / "test_config.json"
        config_file.write_text(json.dumps(config_data))

        # Load config and apply replacements
        operations, _ = load_operations_from_json(config_file)
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacements were applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)
        table1 = modified_doc.tables[0]
        assert table1.rows[2].cells[0].text == "20.95"

        # Second table should be unchanged
        table2 = modified_doc.tables[1]
        assert table2.rows[4].cells[1].text == "{{ o2_qaqc.O2Low_cal_span }}"

    def test_no_matches_in_table(self, tmp_path):
        """Test when search text doesn't exist in the target table."""
        # Create test document
        doc = self.create_test_document_with_two_tables()
        test_file = tmp_path / "test_table.docx"
        doc.save(test_file)

        # Configure replacement with non-existent search text
        operations = [{
            'op': 'replace_in_table',
            'table_heading': 'O2 VALUES',
            'search': 'NonExistentText',
            'replace': 'Replacement'
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify no changes were made
        assert result is False

    def test_table_index_disambiguation(self, tmp_path):
        """Test using table_index when multiple tables have similar headings."""
        # Create document with two tables having similar headings
        doc = Document()

        # First table: O2 VALUES
        table1 = doc.add_table(rows=2, cols=2)
        table1.rows[0].cells[0].text = "O2 VALUES"
        table1.rows[1].cells[0].text = "Data1"

        # Second table: also O2 VALUES
        table2 = doc.add_table(rows=2, cols=2)
        table2.rows[0].cells[0].text = "O2 VALUES"
        table2.rows[1].cells[0].text = "Data2"

        test_file = tmp_path / "test_disambiguation.docx"
        doc.save(test_file)

        # Configure replacement targeting the second table specifically
        operations = [{
            'op': 'replace_in_table',
            'table_heading': 'O2 VALUES',
            'table_index': 1,  # Target the second table
            'search': 'Data2',
            'replace': 'Modified'
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check the modified document
        modified_doc = Document(test_file)

        # First table should be unchanged
        assert modified_doc.tables[0].rows[1].cells[0].text == "Data1"

        # Second table should be changed
        assert modified_doc.tables[1].rows[1].cells[0].text == "Modified"
