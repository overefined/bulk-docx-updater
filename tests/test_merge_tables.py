"""
Unit tests for merge_tables.

Documents rendered from split templates often repeat the same table (identical
title + header block) once per page. merge_tables folds those continuation
tables back into the first one, dropping the duplicated leading header rows so
the result is a single continuous table.
"""
import pytest
from docx import Document

from src.document_processor import DocxBulkUpdater
from src.config import load_operations_from_json


def _header_rows(table):
    """The two shared leading rows: a title row and a 'Station:' row."""
    table.rows[0].cells[0].text = "NMNEHC Test Results"
    table.rows[1].cells[0].text = "Station:"
    table.rows[1].cells[1].text = "{{ station }}"


def _make_split_doc(tmp_path, spacer_variants=("", "\xa0", "\xa0")):
    """Three tables sharing the same 2-row header, each with one unique data row.

    spacer_variants injects a differing spacer cell (empty vs non-breaking
    space) into the header row of each table to mimic the trivial whitespace
    drift seen between real continuation copies.
    """
    doc = Document()
    analytes = ["Acetaldehyde", "Ethylene", "Formic Acid"]
    for analyte, spacer in zip(analytes, spacer_variants):
        t = doc.add_table(rows=3, cols=3)
        _header_rows(t)
        t.rows[1].cells[2].text = spacer
        t.rows[2].cells[0].text = f"{analyte} Emission Results"
        # A blank, page-break-ish separator paragraph between tables.
        doc.add_paragraph("")
    path = tmp_path / "split.docx"
    doc.save(path)
    return path


def _labels(path):
    d = Document(path)
    nm = [t for t in d.tables if t.rows and "NMNEHC" in t.rows[0].cells[0].text]
    return nm, [r.cells[0].text.strip() for t in nm for r in t.rows]


def test_merge_collapses_duplicate_headers(tmp_path):
    path = _make_split_doc(tmp_path)
    ops = [{"op": "merge_tables", "match": "NMNEHC Test Results"}]

    assert DocxBulkUpdater(ops).modify_docx(path) is True

    tables, labels = _labels(path)
    assert len(tables) == 1
    # Header block appears exactly once, then all three data rows follow.
    assert labels.count("NMNEHC Test Results") == 1
    assert labels.count("Station:") == 1
    assert labels == [
        "NMNEHC Test Results", "Station:",
        "Acetaldehyde Emission Results",
        "Ethylene Emission Results",
        "Formic Acid Emission Results",
    ]


def test_merge_is_idempotent(tmp_path):
    path = _make_split_doc(tmp_path)
    ops = [{"op": "merge_tables", "match": "NMNEHC Test Results"}]

    assert DocxBulkUpdater(ops).modify_docx(path) is True
    # Only one table matches now, so a re-run changes nothing.
    assert DocxBulkUpdater(ops).modify_docx(path) is False
    tables, _ = _labels(path)
    assert len(tables) == 1


def test_merge_skip_rows_override(tmp_path):
    path = _make_split_doc(tmp_path)
    ops = [{"op": "merge_tables", "match": "NMNEHC Test Results", "skip_rows": 1}]

    assert DocxBulkUpdater(ops).modify_docx(path) is True
    _, labels = _labels(path)
    # Only the title row is dropped from continuations, so 'Station:' repeats.
    assert labels.count("NMNEHC Test Results") == 1
    assert labels.count("Station:") == 3


def test_merge_by_table_header(tmp_path):
    path = _make_split_doc(tmp_path)
    ops = [{"op": "merge_tables", "table_header": "NMNEHC Test Results"}]

    assert DocxBulkUpdater(ops).modify_docx(path) is True
    tables, _ = _labels(path)
    assert len(tables) == 1


def test_merge_noop_on_single_table(tmp_path):
    doc = Document()
    t = doc.add_table(rows=3, cols=3)
    _header_rows(t)
    t.rows[2].cells[0].text = "Benzene Emission Results"
    path = tmp_path / "one.docx"
    doc.save(path)

    ops = [{"op": "merge_tables", "match": "NMNEHC Test Results"}]
    assert DocxBulkUpdater(ops).modify_docx(path) is False


def test_config_rejects_merge_without_locator(tmp_path):
    cfg = tmp_path / "bad.json"
    cfg.write_text('{"merge_tables": {"skip_rows": 1}}')
    with pytest.raises(SystemExit):
        load_operations_from_json(cfg)
