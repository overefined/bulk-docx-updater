#!/usr/bin/env python3
"""
Unit tests for table cell alignment functionality.

Tests both the complex formatting token system and the config-driven alignment
to ensure table cell alignment works correctly.
"""

import pytest
import tempfile
import shutil
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table, _Row, _Cell

from src.formatting import FormattingProcessor
from src.document_processor import DocxBulkUpdater


class TestTableCellAlignment:
    """Test suite for table cell alignment functionality."""

    def setup_method(self):
        """Set up test fixtures."""
        self.formatter = FormattingProcessor()

    def create_test_document_with_table(self):
        """Create a test document with a table containing target text."""
        doc = Document()

        # Add a table with test data
        table = doc.add_table(rows=3, cols=3)

        # Add some regular content
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(0, 2).text = "Header 3"

        # Add target content in specific cells
        table.cell(1, 0).text = "{{ o2.phase_fmtd() }}"
        table.cell(1, 1).text = "{{ o2.ReadingTimestamp }}"
        table.cell(1, 2).text = "Some other data"

        # Add more content
        table.cell(2, 0).text = "More data"
        table.cell(2, 1).text = "Even more data"
        table.cell(2, 2).text = "Final data"

        # Set initial alignment to CENTER for target cells
        table.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        return doc

    def test_table_cell_detection(self):
        """Test that we can correctly detect paragraphs inside table cells."""
        doc = self.create_test_document_with_table()

        # Get the table and test paragraphs
        table = doc.tables[0]
        target_paragraph = table.cell(1, 0).paragraphs[0]
        regular_paragraph = doc.add_paragraph("Regular paragraph")

        # Test table cell detection
        assert self.formatter._is_paragraph_in_table_cell(target_paragraph) == True
        assert self.formatter._is_paragraph_in_table_cell(regular_paragraph) == False

    def test_table_cell_alignment_direct_setting(self):
        """Test direct setting of table cell paragraph alignment."""
        doc = self.create_test_document_with_table()
        table = doc.tables[0]

        # Get target paragraphs
        target_para1 = table.cell(1, 0).paragraphs[0]  # {{ o2.phase_fmtd() }}
        target_para2 = table.cell(1, 1).paragraphs[0]  # {{ o2.ReadingTimestamp }}

        # Verify initial alignment (should be CENTER)
        assert target_para1.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert target_para2.alignment == WD_ALIGN_PARAGRAPH.CENTER

        # Test direct alignment setting
        target_para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        target_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Verify alignment changed
        assert target_para1.alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert target_para2.alignment == WD_ALIGN_PARAGRAPH.LEFT

    def test_formatter_apply_paragraph_formatting_regular_paragraphs(self):
        """Test FormattingProcessor.apply_paragraph_formatting on regular paragraphs.

        Note: We've found that the complex formatter approach doesn't work reliably
        for table cells, which is why we use the simple direct approach for those.
        This test verifies the formatter works for regular paragraphs.
        """
        doc = Document()
        regular_para1 = doc.add_paragraph("Regular paragraph 1")
        regular_para2 = doc.add_paragraph("Regular paragraph 2")

        # Apply alignment using formatter
        self.formatter.apply_paragraph_formatting(regular_para1, {'alignment': WD_ALIGN_PARAGRAPH.LEFT})
        self.formatter.apply_paragraph_formatting(regular_para2, {'alignment': WD_ALIGN_PARAGRAPH.RIGHT})

        # Verify alignment changed
        assert regular_para1.alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert regular_para2.alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def test_format_token_parsing_left_alignment(self):
        """Test that {format:left} tokens are correctly parsed."""
        # Test format option parsing
        result = self.formatter._parse_format_options('left')
        assert result['alignment'] == WD_ALIGN_PARAGRAPH.LEFT

        # Test format token processing
        test_text = '{format:left}{{ o2.ReadingTimestamp }}{/format}'
        segments = self.formatter.process_formatting_tokens(test_text, None)

        assert len(segments) == 1
        text, formatting = segments[0]
        assert text == '{{ o2.ReadingTimestamp }}'
        assert formatting['alignment'] == WD_ALIGN_PARAGRAPH.LEFT

    def test_format_token_parsing_all_alignments(self):
        """Test parsing of all alignment format tokens."""
        alignment_tests = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }

        for token, expected_alignment in alignment_tests.items():
            result = self.formatter._parse_format_options(token)
            assert result['alignment'] == expected_alignment, f"Failed for {token}"

    def test_config_driven_table_alignment(self):
        """Test the config-driven table cell alignment operation."""
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            tmp_path = tmp_file.name

        try:
            # Create and save test document
            doc = self.create_test_document_with_table()
            doc.save(tmp_path)

            # Create updater with align_table_cells operation
            operations = [{
                'op': 'align_table_cells',
                'patterns': ["o2.phase_fmtd", "o2.ReadingTimestamp"],
                'alignment': 'left'
            }]
            updater = DocxBulkUpdater(operations)

            # Process the document
            result = updater.modify_docx(Path(tmp_path))
            assert result is True

            # Verify the changes by re-reading the document
            doc_after = Document(tmp_path)
            table = doc_after.tables[0]

            # Check that target cells are now LEFT aligned
            target_para1 = table.cell(1, 0).paragraphs[0]  # {{ o2.phase_fmtd() }}
            target_para2 = table.cell(1, 1).paragraphs[0]  # {{ o2.ReadingTimestamp }}

            assert target_para1.alignment == WD_ALIGN_PARAGRAPH.LEFT
            assert target_para2.alignment == WD_ALIGN_PARAGRAPH.LEFT

            # Check that non-target cells are unchanged
            header_para = table.cell(0, 0).paragraphs[0]  # "Header 1"
            other_para = table.cell(1, 2).paragraphs[0]   # "Some other data"

            # These should not have been modified (alignment should be None/default)
            assert header_para.alignment != WD_ALIGN_PARAGRAPH.LEFT or header_para.alignment is None
            assert other_para.alignment != WD_ALIGN_PARAGRAPH.LEFT or other_para.alignment is None

        finally:
            # Clean up temporary file
            Path(tmp_path).unlink(missing_ok=True)

    def test_requires_special_handling_with_alignment(self):
        """Test that _requires_special_handling correctly detects alignment formatting."""
        from src.text_replacement import TextReplacer

        # Create a mock replacer (we only need the method)
        replacer = TextReplacer([], self.formatter)

        # Test segments with alignment
        segments_with_alignment = [
            ("{{ o2.ReadingTimestamp }}", {'alignment': WD_ALIGN_PARAGRAPH.LEFT})
        ]

        segments_without_alignment = [
            ("{{ o2.ReadingTimestamp }}", {'bold': True})
        ]

        # Should return True for segments with alignment
        assert replacer._requires_special_handling(segments_with_alignment) == True

        # Should return False for segments without alignment
        assert replacer._requires_special_handling(segments_without_alignment) == False

    def test_alignment_enum_value_edge_cases(self):
        """Test edge cases with WD_PARAGRAPH_ALIGNMENT enum values."""
        # Test that LEFT (value 0) doesn't evaluate to False when checked properly
        left_alignment = WD_ALIGN_PARAGRAPH.LEFT

        # These should fail (demonstrating the bug we fixed)
        assert not bool(left_alignment)  # LEFT has value 0, so bool() is False
        assert not (left_alignment)      # Same issue in if statements

        # These should work (the correct way to check)
        assert left_alignment is not None
        assert left_alignment == WD_ALIGN_PARAGRAPH.LEFT

        # Test all alignment values
        alignments = [
            WD_ALIGN_PARAGRAPH.LEFT,    # 0
            WD_ALIGN_PARAGRAPH.CENTER,  # 1
            WD_ALIGN_PARAGRAPH.RIGHT,   # 2
            WD_ALIGN_PARAGRAPH.JUSTIFY  # 3
        ]

        for alignment in alignments:
            assert alignment is not None
            assert isinstance(alignment, type(WD_ALIGN_PARAGRAPH.LEFT))

    def test_table_cell_alignment_with_document_save_load(self):
        """Test that table cell alignment persists through document save/load cycles."""
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp1, \
             tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp2:
            tmp_path1 = tmp1.name
            tmp_path2 = tmp2.name

        try:
            # Create document and save
            doc = self.create_test_document_with_table()
            table = doc.tables[0]

            # Set alignment and save
            table.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc.save(tmp_path1)

            # Load document and verify alignment persisted
            doc2 = Document(tmp_path1)
            table2 = doc2.tables[0]

            assert table2.cell(1, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
            assert table2.cell(1, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

            # Modify alignment and save again
            table2.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc2.save(tmp_path2)

            # Load again and verify new alignment
            doc3 = Document(tmp_path2)
            table3 = doc3.tables[0]

            assert table3.cell(1, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
            assert table3.cell(1, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

        finally:
            # Clean up
            Path(tmp_path1).unlink(missing_ok=True)
            Path(tmp_path2).unlink(missing_ok=True)


if __name__ == "__main__":
    # Run tests if called directly
    pytest.main([__file__, "-v"])