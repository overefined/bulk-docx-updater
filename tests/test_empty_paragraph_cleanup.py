"""
Unit tests for empty paragraph cleanup functionality.

Tests the remove_empty_paragraphs_after_pattern method and its integration
with insert_after operations in the DocxBulkUpdater class.
"""
import pytest
from pathlib import Path
from docx import Document

from document_processor import DocxBulkUpdater


class TestEmptyParagraphCleanup:
    """Test cases for empty paragraph cleanup functionality."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.replacements = [
            {"search": "SITE PHOTOS", "insert_after": "New content", "remove_empty_paragraphs_after": True},
            {"remove_empty_paragraphs_after": "{{ technician_resume }}"}
        ]
        self.updater = DocxBulkUpdater(self.replacements)
    
    def create_test_document(self, paragraphs_content):
        """Helper method to create a test document with specified paragraph content."""
        doc = Document()
        for content in paragraphs_content:
            if content is None or content == "":
                # Create empty paragraph
                doc.add_paragraph("")
            else:
                doc.add_paragraph(content)
        return doc
    
    def test_remove_empty_paragraphs_after_pattern_basic(self):
        """Test removal of empty runs and formatting artifacts from the next paragraph after pattern."""
        doc = Document()
        doc.add_paragraph("Some content")
        doc.add_paragraph("SITE PHOTOS")
        
        # Create next paragraph with empty runs and text content
        next_para = doc.add_paragraph("")
        next_para.add_run("")  # Empty run - should be removed
        next_para.add_run("")  # Empty run - should be removed
        next_para.add_run("More content")
        
        modified = self.updater.remove_empty_paragraphs_after_pattern(doc, "SITE PHOTOS")
        
        assert modified is True
        # The empty runs should be removed, leaving only "More content"
        assert next_para.text == "More content"
    
    def test_remove_empty_paragraphs_after_pattern_no_empty(self):
        """Test when there are no leading whitespace runs after the pattern."""
        doc = Document()
        doc.add_paragraph("Some content")
        doc.add_paragraph("SITE PHOTOS") 
        doc.add_paragraph("More content")  # No leading whitespace
        
        modified = self.updater.remove_empty_paragraphs_after_pattern(doc, "SITE PHOTOS")
        
        assert modified is False
    
    def test_remove_empty_paragraphs_after_pattern_multiple_occurrences(self):
        """Test cleanup after multiple occurrences of the pattern."""
        doc = Document()
        doc.add_paragraph("First content")
        doc.add_paragraph("SITE PHOTOS")
        
        # Next paragraph with empty run and text
        next_para1 = doc.add_paragraph("")
        next_para1.add_run("")  # empty run - should be removed
        next_para1.add_run("Second content")
        
        doc.add_paragraph("SITE PHOTOS") 
        
        # Another paragraph with empty run and text
        next_para2 = doc.add_paragraph("")
        next_para2.add_run("")  # empty run - should be removed  
        next_para2.add_run("Third content")
        
        modified = self.updater.remove_empty_paragraphs_after_pattern(doc, "SITE PHOTOS")
        
        assert modified is True
        # Both paragraphs should have their empty runs cleaned
        assert next_para1.text == "Second content"
        assert next_para2.text == "Third content"
    
    def test_remove_empty_paragraphs_after_pattern_pattern_not_found(self):
        """Test when the pattern is not found in the document."""
        doc = Document()
        doc.add_paragraph("Some content")
        next_para = doc.add_paragraph("")
        next_para.add_run("\n")
        next_para.add_run("Other content")
        
        modified = self.updater.remove_empty_paragraphs_after_pattern(doc, "SITE PHOTOS")
        
        assert modified is False
        # No changes should be made since pattern not found
        assert next_para.text == "\nOther content"
    
    def test_remove_empty_paragraphs_after_pattern_whitespace_only(self):
        """Test that whitespace runs are preserved, only empty runs are removed."""
        doc = Document()
        doc.add_paragraph("Some content")
        doc.add_paragraph("SITE PHOTOS")
        
        # Next paragraph with whitespace runs and empty runs
        next_para = doc.add_paragraph("")
        next_para.add_run("")     # empty run - should be removed
        next_para.add_run("   ")  # spaces - preserved (intentional formatting)
        next_para.add_run("More content")
        
        modified = self.updater.remove_empty_paragraphs_after_pattern(doc, "SITE PHOTOS")
        
        assert modified is True
        # Whitespace and content should remain, empty run removed
        assert next_para.text == "   More content"
    
    def test_remove_empty_paragraphs_after_pattern_stops_at_content(self):
        """Test that cleanup only affects the immediate next paragraph."""
        doc = Document()
        doc.add_paragraph("Some content")
        doc.add_paragraph("SITE PHOTOS")
        
        # Next paragraph with empty run - this should be cleaned  
        next_para = doc.add_paragraph("")
        next_para.add_run("")  # empty run - should be removed
        next_para.add_run("Keep this content")
        
        # Another paragraph - this should NOT be affected
        further_para = doc.add_paragraph("")
        further_para.add_run(" ")
        further_para.add_run("More content")
        
        modified = self.updater.remove_empty_paragraphs_after_pattern(doc, "SITE PHOTOS")
        
        assert modified is True
        # Only the immediate next paragraph should be cleaned
        assert next_para.text == "Keep this content"
        assert further_para.text == " More content"  # This remains unchanged


class TestCleanupIntegration:
    """Integration tests for cleanup functionality."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.replacements = [
            {
                "search": "SITE PHOTOS",
                "insert_after": "pagebreak{format:center,size12}Inserted content{/format}",
                "remove_empty_paragraphs_after": True
            },
            {
                "remove_empty_paragraphs_after": "{{ technician_resume }}"
            }
        ]
        self.updater = DocxBulkUpdater(self.replacements)
    
    def test_cleanup_with_boolean_flag(self):
        """Test cleanup when remove_empty_paragraphs_after is set to True."""
        replacements = [
            {
                "search": "SITE PHOTOS",
                "insert_after": "New content",
                "remove_empty_paragraphs_after": True
            }
        ]
        updater = DocxBulkUpdater(replacements)
        
        doc = Document()
        doc.add_paragraph("Some content")
        doc.add_paragraph("SITE PHOTOS")
        
        # Next paragraph with empty run
        next_para = doc.add_paragraph("")
        next_para.add_run("")  # empty run - should be removed
        next_para.add_run("More content")
        
        # Test the cleanup method directly
        modified = updater.remove_empty_paragraphs_after_pattern(doc, "SITE PHOTOS")
        
        assert modified is True
        assert next_para.text == "More content"
    
    def test_cleanup_with_string_pattern(self):
        """Test cleanup when remove_empty_paragraphs_after is a string pattern."""
        replacements = [
            {
                "remove_empty_paragraphs_after": "{{ technician_resume }}"
            }
        ]
        updater = DocxBulkUpdater(replacements)
        
        doc = Document()
        doc.add_paragraph("Some content")
        doc.add_paragraph("{{ technician_resume }}")
        
        # Next paragraph with empty run
        next_para = doc.add_paragraph("")
        next_para.add_run("")  # empty run - should be removed
        next_para.add_run("More content")
        
        # Test the cleanup method directly
        modified = updater.remove_empty_paragraphs_after_pattern(doc, "{{ technician_resume }}")
        
        assert modified is True
        assert next_para.text == "More content"
