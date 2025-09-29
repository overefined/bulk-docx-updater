"""
Main document processing class for DOCX bulk updates.

Contains the DocxBulkUpdater class and methods for document-level operations
like margin standardization, paragraph cleanup, and change preview.
"""
from __future__ import annotations
import re
import difflib
import tempfile
import shutil
from pathlib import Path
from typing import List, Dict, Optional, Tuple
import logging
from docx import Document
from docx.shared import Inches

from formatting import FormattingProcessor
from text_replacement import TextReplacer
from font_utils import FontFormatter


class DocxBulkUpdater:
    """Main class for bulk DOCX document processing and text replacement."""
    
    def __init__(self, replacements: List[Dict[str, str]], preserve_formatting: bool = True, 
                 standardize_margins: bool = False, margins: Optional[Dict[str, float]] = None,
                 diff_context: int = 3):
        self.replacements = replacements
        self.preserve_formatting = preserve_formatting
        self.standardize_margins = standardize_margins
        self.margins = margins or {
            'top': 1.0,
            'bottom': 1.0,
            'left': 1.0,
            'right': 1.0
        }
        self.diff_context = diff_context
        self._logger = logging.getLogger(__name__)
        
        # Initialize component processors
        self.formatter = FormattingProcessor()
        self.text_replacer = TextReplacer(replacements, self.formatter)
        
        # Pre-compute cross-paragraph patterns for optimization
        self._cross_paragraph_patterns = self._get_cross_paragraph_search_patterns()
        self._has_cross_paragraph_patterns = len(self._cross_paragraph_patterns) > 0
        
        # Performance optimization caches
        self._paragraph_cache = {}
        self._text_cache = {}
        self._xml_cache = {}  # Cache XML strings to reduce xpath calls
        
        # Pre-compile search patterns for faster matching
        self._search_patterns_set = {repl.get('search', '') for repl in replacements if repl.get('search', '')}
        self._search_patterns_set.discard('')  # Remove empty patterns
        
    def clear_caches(self):
        """Clear performance caches to free memory."""
        self._paragraph_cache.clear()
        self._text_cache.clear()
        self._xml_cache.clear()
        self.text_replacer.clear_caches()
    
    
    def _iter_all_paragraphs(self, doc: Document):
        """Yield all paragraphs across body, tables, headers, and footers.
        
        Uses caching to avoid repeated paragraph extraction.
        """
        doc_id = id(doc)
        if doc_id in self._paragraph_cache:
            yield from self._paragraph_cache[doc_id]
            return
        
        paragraphs = []
        
        # Body paragraphs
        paragraphs.extend(doc.paragraphs)
        
        # Tables (all cells' paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        
        # Headers and footers for all sections
        for section in doc.sections:
            paragraphs.extend(section.header.paragraphs)
            paragraphs.extend(section.footer.paragraphs)
        
        # Cache the full list for subsequent calls
        self._paragraph_cache[doc_id] = paragraphs
        
        # Yield all at once after building complete list
        yield from paragraphs
    
    
    def _get_cross_paragraph_search_patterns(self) -> List[str]:
        """Pre-compute list of search patterns that might span paragraphs."""
        patterns = []
        for replacement in self.replacements:
            if 'search' not in replacement:
                continue
            if 'replace' not in replacement:
                continue
            patterns.append(replacement['search'])
        return patterns
    
    def _chunk_has_cross_paragraph_potential(self, paragraphs: List, paragraph_texts: List[str] = None) -> bool:
        """Check if a chunk of paragraphs might contain cross-paragraph patterns (optimized)."""
        if len(paragraphs) < 2 or not self._has_cross_paragraph_patterns:
            return False
        
        # Use cached paragraph texts if provided
        if paragraph_texts is None:
            paragraph_texts = [para.text for para in paragraphs]
        
        # Early exit: if all paragraphs are empty, no potential
        if not any(text.strip() for text in paragraph_texts):
            return False
        
        # Combine text from all paragraphs (using pre-computed texts)
        combined_text = "".join(paragraph_texts)
        if not combined_text.strip():
            return False
        
        # Check pre-computed patterns only
        for search_text in self._cross_paragraph_patterns:
            if search_text in combined_text:
                # Check if this pattern actually spans paragraphs
                # (i.e., it doesn't appear complete in any single paragraph)
                spans_paragraphs = True
                for text in paragraph_texts:
                    if search_text in text:
                        spans_paragraphs = False
                        break
                
                if spans_paragraphs:
                    return True
        
        return False
    
    def _process_all_text_replacements(self, doc: Document) -> bool:
        """Efficiently process both cross-paragraph and single-paragraph replacements."""
        modified = False
        processed_paragraphs = set()  # Track which paragraphs were already processed
        
        # First, handle cross-paragraph replacements if any exist
        if self._has_cross_paragraph_patterns:
            # Process body paragraphs in chunks
            body_paragraphs = list(doc.paragraphs)
            if body_paragraphs:
                if self._process_paragraph_chunks_tracked(body_paragraphs, processed_paragraphs):
                    modified = True
            
            # Process table cell paragraphs in chunks
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_paragraphs = list(cell.paragraphs)
                        if cell_paragraphs:
                            if self._process_paragraph_chunks_tracked(cell_paragraphs, processed_paragraphs):
                                modified = True
            
            # Process header/footer paragraphs in chunks
            for section in doc.sections:
                if hasattr(section, 'header') and section.header:
                    header_paragraphs = list(section.header.paragraphs)
                    if header_paragraphs:
                        if self._process_paragraph_chunks_tracked(header_paragraphs, processed_paragraphs):
                            modified = True
                            
                if hasattr(section, 'footer') and section.footer:
                    footer_paragraphs = list(section.footer.paragraphs)
                    if footer_paragraphs:
                        if self._process_paragraph_chunks_tracked(footer_paragraphs, processed_paragraphs):
                            modified = True
        
        # Then process remaining single-paragraph replacements
        for paragraph in self._iter_all_paragraphs(doc):
            # Skip paragraphs that were already processed in cross-paragraph chunks
            if id(paragraph) not in processed_paragraphs:
                if self.text_replacer.replace_text_in_paragraph(paragraph):
                    modified = True
        
        return modified
    
    def _process_paragraph_chunks_tracked(self, paragraphs: List, processed_paragraphs: set) -> bool:
        """Process paragraph chunks and track which paragraphs were processed."""
        if len(paragraphs) < 2:
            return False
        
        # Pre-compute all paragraph texts once to avoid repeated .text calls
        # Use caching to avoid repeated text extraction
        paragraph_texts = []
        for para in paragraphs:
            para_id = id(para)
            if para_id in self._text_cache:
                paragraph_texts.append(self._text_cache[para_id])
            else:
                text = para.text
                self._text_cache[para_id] = text
                paragraph_texts.append(text)
        
        # Early exit if no cross-paragraph patterns could possibly exist
        combined_text = "".join(paragraph_texts)
        # Use pre-compiled set for faster pattern matching
        if not any(pattern in combined_text for pattern in self._cross_paragraph_patterns):
            return False
        
        modified = False
        max_chunk_size = 5  # Maximum paragraphs to consider at once
        
        # Try different chunk sizes, starting with larger chunks
        for chunk_size in range(min(max_chunk_size, len(paragraphs)), 1, -1):
            i = 0
            while i <= len(paragraphs) - chunk_size:
                chunk = paragraphs[i:i + chunk_size]
                chunk_texts = paragraph_texts[i:i + chunk_size]
                
                # Check if this chunk has any potential cross-paragraph patterns
                if self._chunk_has_cross_paragraph_potential(chunk, chunk_texts):
                    if self.text_replacer.replace_text_across_paragraphs(chunk):
                        modified = True
                        # Mark all paragraphs in this chunk as processed
                        for para in chunk:
                            processed_paragraphs.add(id(para))
                        # Skip ahead since we processed this chunk
                        i += chunk_size
                        continue
                
                i += 1
        
        return modified
    
    def standardize_document_margins(self, doc: Document) -> bool:
        """Standardize margins for all sections in the document."""
        if not self.standardize_margins:
            return False
            
        modified = False
        for section in doc.sections:
            # Set margins using Inches for consistency
            section.top_margin = Inches(self.margins['top'])
            section.bottom_margin = Inches(self.margins['bottom'])
            section.left_margin = Inches(self.margins['left'])
            section.right_margin = Inches(self.margins['right'])
            modified = True
            
        return modified
    
    def _has_column_break_in_run(self, run) -> bool:
        """Check if a run contains a column break."""
        # Check for column breaks in the XML
        return 'w:br' in run._element.xml and 'type="column"' in run._element.xml
    
    def _has_page_break_in_run(self, run) -> bool:
        """Check if a run contains a page break."""
        # Check for page breaks in the XML
        return 'w:br' in run._element.xml and 'type="page"' in run._element.xml
    
    def remove_empty_paragraphs_after_pattern(self, doc: Document, pattern: str) -> bool:
        """Remove column breaks and specific formatting artifacts from the next paragraph after the pattern."""
        modified = False
        
        # Find paragraphs containing the pattern
        for i, para in enumerate(doc.paragraphs):
            if pattern in para.text:
                # Check the very next paragraph
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    
                    # Only remove specific formatting artifacts, preserve intentional spacing
                    runs_to_remove = []
                    for run in next_para.runs:
                        # Check if this run contains specific breaks we want to remove
                        if self._has_column_break_in_run(run):
                            runs_to_remove.append(run)
                        elif self._has_page_break_in_run(run):
                            runs_to_remove.append(run)
                        elif not run.text:  # Completely empty runs only
                            runs_to_remove.append(run)
                        else:
                            # Stop at first run with any text content (including spaces/newlines)
                            break
                    
                    for run in runs_to_remove:
                        run._element.getparent().remove(run._element)
                        modified = True
        
        return modified

    def modify_docx(self, file_path: Path) -> bool:
        """Modify a DOCX file with the specified replacements."""
        try:
            # Load the document
            doc = Document(file_path)
            modified = False

            # Standardize margins if enabled
            if self.standardize_margins:
                if self.standardize_document_margins(doc):
                    modified = True

            # Set or unset table header repeat as configured
            for replacement in self.replacements:
                if "set_table_header_repeat" in replacement:
                    header_config = replacement["set_table_header_repeat"]
                    if isinstance(header_config, bool) and header_config:
                        # Use the search pattern to identify header rows
                        header_pattern = replacement.get("search")
                        if self.set_table_header_repeat(doc, header_pattern, enable=True):
                            modified = True
                    elif isinstance(header_config, str):
                        # Use the provided pattern string
                        if self.set_table_header_repeat(doc, header_config, enable=True):
                            modified = True
                    elif isinstance(header_config, dict):
                        # Use pattern from dict config
                        header_pattern = header_config.get("pattern")
                        enable = True if header_config.get("enabled", True) else False
                        if self.set_table_header_repeat(doc, header_pattern, enable=enable):
                            modified = True

            # Change font sizes if enabled
            for replacement in self.replacements:
                if "change_font_size" in replacement:
                    font_config = replacement["change_font_size"]
                    if isinstance(font_config, dict):
                        from_size = font_config.get("from")
                        to_size = font_config.get("to")
                        if from_size and to_size:
                            if self.change_font_sizes(doc, from_size, to_size):
                                modified = True

            # Remove empty paragraphs after patterns if cleanup is enabled
            for replacement in self.replacements:
                if "remove_empty_paragraphs_after" in replacement:
                    cleanup_value = replacement["remove_empty_paragraphs_after"]
                    if isinstance(cleanup_value, bool) and cleanup_value:
                        # Use the search pattern from the same replacement
                        if "search" in replacement:
                            pattern = replacement["search"]
                            if self.remove_empty_paragraphs_after_pattern(doc, pattern):
                                modified = True
                    elif isinstance(cleanup_value, str):
                        # Use the provided pattern string
                        if self.remove_empty_paragraphs_after_pattern(doc, cleanup_value):
                            modified = True

            # Handle table cell replacements
            for replacement in self.replacements:
                if "replace_table_cell" in replacement:
                    cell_config = replacement["replace_table_cell"]
                    if self.replace_table_cell(doc, cell_config):
                        modified = True

            # Then do the text replacements and inserts
            has_search_ops = any(("search" in r) and ("replace" in r) for r in self.replacements)

            if has_search_ops:
                # Process both cross-paragraph and single-paragraph replacements efficiently
                if self._process_all_text_replacements(doc):
                    modified = True

            # Save changes if any modifications were made
            if modified:
                doc.save(file_path)
                return True
            
            return False
            
        except Exception as e:
            logging.getLogger(__name__).error("Error processing %s: %s", file_path, e)
            return False
    
    
    def get_document_changes_preview(self, file_path: Path) -> Dict[str, str]:
        """Get a preview of changes by running actual operations on a temporary copy."""

        try:
            # Create temporary copy
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_path = Path(temp_file.name)

            shutil.copy2(file_path, temp_path)

            try:
                # Track formatting operations and get content changes
                return self._compare_document_contents_with_formatting(file_path, temp_path)

            finally:
                # Clean up temporary file
                temp_path.unlink(missing_ok=True)

        except Exception as e:
            logging.getLogger(__name__).error("Error previewing changes for %s: %s", file_path, e)
            return {}
    
    def _compare_document_contents_with_formatting(self, original_path: Path, temp_path: Path) -> Dict[str, str]:
        """Compare original and modified document contents and track formatting operations."""
        # Clear cache to ensure fresh content extraction
        self._text_cache.clear()

        # Get original content and properties
        original_doc = Document(original_path)
        original_content = self._extract_all_content(original_doc, extract_xml=False)
        original_properties = self._extract_document_properties(original_doc)

        # Track operations performed during modification
        operation_results = []

        # Apply modifications to temporary copy while tracking operations
        modified_doc = Document(temp_path)

        # Track formatting operations
        for replacement in self.replacements:
            if "set_table_header_repeat" in replacement:
                header_config = replacement["set_table_header_repeat"]
                if isinstance(header_config, str):
                    count = self.set_table_header_repeat(modified_doc, header_config, enable=True)
                    if count > 0:
                        operation_results.append(f"Set {count} table header row(s) to repeat: '{header_config}'")
                elif isinstance(header_config, dict):
                    pat = header_config.get("pattern")
                    enable = True if header_config.get("enabled", True) else False
                    count = self.set_table_header_repeat(modified_doc, pat, enable=enable)
                    if count > 0:
                        action = "Set" if enable else "Unset"
                        pattxt = f" '{pat}'" if pat else " (first row)"
                        operation_results.append(f"{action} table header repeat on {count} row(s){pattxt}")

            if "change_font_size" in replacement:
                font_config = replacement["change_font_size"]
                if isinstance(font_config, dict):
                    from_size = font_config.get("from")
                    to_size = font_config.get("to")
                    if from_size and to_size:
                        count = self.change_font_sizes(modified_doc, from_size, to_size)
                        if count > 0:
                            operation_results.append(f"Changed font size from {from_size}pt to {to_size}pt in {count} text run(s)")

        # Apply standard text replacements
        self._process_all_text_replacements(modified_doc)

        # Save the modified document
        modified_doc.save(temp_path)

        # Clear cache again to avoid reusing original content for modified doc
        self._text_cache.clear()

        # Get modified content
        modified_doc = Document(temp_path)
        modified_content = self._extract_all_content(modified_doc, extract_xml=False)

        # Find text content differences
        changes = {}
        for section_name in original_content.keys():
            if section_name in modified_content:
                orig_lines = original_content[section_name]
                mod_lines = modified_content[section_name]

                if orig_lines != mod_lines:
                    changes[section_name] = (orig_lines, mod_lines)

        # Add formatting operation results as a special section
        if operation_results:
            changes["Formatting Operations"] = ([], operation_results)

        return changes

    def _extract_document_properties(self, doc: Document) -> Dict[str, any]:
        """Extract document properties for comparison."""
        properties = {}

        # Extract table header repeat properties
        table_headers = []
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                try:
                    tr_element = row._tr
                    from docx.oxml.ns import qn
                    tr_pr = tr_element.find(qn('w:trPr'))
                    if tr_pr is not None:
                        header_elem = tr_pr.find(qn('w:tblHeader'))
                        if header_elem is not None:
                            row_text = ' '.join(cell.text.strip() for cell in row.cells)
                            table_headers.append(f"Table {i+1}, Row {j+1}: {row_text}")
                except:
                    continue
        properties['table_headers'] = table_headers

        # Extract font sizes
        font_sizes = {}
        for paragraph in self._iter_all_paragraphs(doc):
            for run in paragraph.runs:
                if run.font.size is not None:
                    size_pt = int(run.font.size.pt)
                    if size_pt not in font_sizes:
                        font_sizes[size_pt] = 0
                    font_sizes[size_pt] += 1
        properties['font_sizes'] = font_sizes

        return properties

    def _extract_all_content(self, doc, extract_xml: bool = False) -> Dict[str, List[str]]:
        """Extract all content from a document organized by section.
        
        Args:
            doc: Document object
            extract_xml: If True, extract XML content; if False, extract text content
        """
        # Check cache first
        doc_id = id(doc)
        cache_key = f"{doc_id}_{extract_xml}"
        if cache_key in self._text_cache:
            return self._text_cache[cache_key]
        
        content = {}
        
        if extract_xml:
            # Body paragraphs XML
            body_xml_lines = []
            for para in doc.paragraphs:
                body_xml_lines.extend((para._p.xml or '').splitlines())
            if body_xml_lines:
                content["Body(XML)"] = body_xml_lines

            # Tables XML
            table_xml_lines = []
            for table in doc.tables:
                table_xml_lines.append("<table>")
                for row in table.rows:
                    table_xml_lines.append("  <row>")
                    for cell in row.cells:
                        table_xml_lines.append("    <cell>")
                        for para in cell.paragraphs:
                            table_xml_lines.extend((para._p.xml or '').splitlines())
                        table_xml_lines.append("    </cell>")
                    table_xml_lines.append("  </row>")
                table_xml_lines.append("</table>")
            if table_xml_lines:
                content["Tables(XML)"] = table_xml_lines

            # Headers/Footers XML
            header_footer_xml_lines = []
            for section in doc.sections:
                for para in section.header.paragraphs:
                    header_footer_xml_lines.extend((para._p.xml or '').splitlines())
                for para in section.footer.paragraphs:
                    header_footer_xml_lines.extend((para._p.xml or '').splitlines())
            if header_footer_xml_lines:
                content["Headers/Footers(XML)"] = header_footer_xml_lines
        else:
            # Body paragraphs
            content["Body"] = [para.text for para in doc.paragraphs]
            
            # Tables
            table_paragraphs = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        table_paragraphs.extend([para.text for para in cell.paragraphs])
            
            if table_paragraphs:
                content["Tables"] = table_paragraphs
            
            # Headers and footers
            header_footer_paragraphs = []
            for section in doc.sections:
                header_footer_paragraphs.extend([para.text for para in section.header.paragraphs])
                header_footer_paragraphs.extend([para.text for para in section.footer.paragraphs])
            
            if header_footer_paragraphs:
                content["Headers/Footers"] = header_footer_paragraphs
        
        # Cache the result
        self._text_cache[cache_key] = content
        return content
    
    def _extract_all_text_content(self, doc) -> Dict[str, List[str]]:
        """Extract all text content from a document organized by section."""
        return self._extract_all_content(doc, extract_xml=False)
    
    def set_table_header_repeat(self, doc: Document, header_pattern: str = None, enable: bool = True) -> int:
        """Enable or disable table header repeat on rows matching a pattern.

        Args:
            doc: Document object
            header_pattern: Text pattern to identify header rows. If None, target first row of each table.
            enable: True to set w:tblHeader, False to remove it.

        Returns:
            Number of rows modified
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        modified_count = 0

        for table in doc.tables:
            header_rows_found = []

            if header_pattern:
                # Search for rows containing the pattern
                for i, row in enumerate(table.rows):
                    row_text = ' '.join(cell.text.strip() for cell in row.cells)
                    if header_pattern in row_text:
                        header_rows_found.append(i)
            else:
                # Default: treat first row as header
                if table.rows:
                    header_rows_found.append(0)

            # Apply or remove repeat header property for found rows
            for row_idx in header_rows_found:
                try:
                    row = table.rows[row_idx]
                    tr_element = row._tr

                    # Check if trPr (table row properties) element exists
                    tr_pr = tr_element.find(qn('w:trPr'))
                    if tr_pr is None:
                        # Create trPr element if it doesn't exist
                        tr_pr = OxmlElement('w:trPr')
                        tr_element.insert(0, tr_pr)

                    # Check if tblHeader element exists
                    tbl_header = tr_pr.find(qn('w:tblHeader'))
                    if enable:
                        if tbl_header is None:
                            # Create and add tblHeader element
                            tbl_header = OxmlElement('w:tblHeader')
                            tr_pr.append(tbl_header)
                            modified_count += 1
                            self._logger.debug(f"Set repeat header for table row {row_idx}")
                    else:
                        if tbl_header is not None:
                            tr_pr.remove(tbl_header)
                            modified_count += 1
                            self._logger.debug(f"Removed repeat header for table row {row_idx}")

                except Exception as e:
                    self._logger.warning(f"Failed to set repeat header for row {row_idx}: {e}")

        return modified_count

    def change_font_sizes(self, doc: Document, from_size: int, to_size: int) -> int:
        """Change all text with a specific font size to a new font size.

        Args:
            doc: The Document object
            from_size: Original font size in points
            to_size: New font size in points

        Returns:
            Number of runs modified
        """
        from docx.shared import Pt

        modified_count = 0
        from_size_half_points = from_size * 2  # Word stores font sizes in half-points
        to_size_pt = Pt(to_size)

        # Process all paragraphs in body, tables, headers, and footers
        for paragraph in self._iter_all_paragraphs(doc):
            for run in paragraph.runs:
                # Check if run has font size property
                if run.font.size is not None:
                    # Convert to half-points for comparison (docx uses Emu internally)
                    current_size_half_points = int(run.font.size.pt * 2)
                    if current_size_half_points == from_size_half_points:
                        run.font.size = to_size_pt
                        modified_count += 1
                        self._logger.debug(f"Changed font size from {from_size}pt to {to_size}pt")

        return modified_count

    def format_diff(self, original_lines: List[str], modified_lines: List[str], section_name: str) -> str:
        """Format a unified diff for display."""
        diff = difflib.unified_diff(
            original_lines,
            modified_lines,
            fromfile=f"original/{section_name}",
            tofile=f"modified/{section_name}",
            n=max(0, int(self.diff_context)) if isinstance(self.diff_context, int) else 3,
            lineterm=''
        )
        return '\n'.join(diff)

    def _extract_all_xml_content(self, doc: Document) -> Dict[str, List[str]]:
        """Extract XML representations of document content organized by section."""
        return self._extract_all_content(doc, extract_xml=True)

    def get_document_xml_changes_preview(self, file_path: Path) -> Dict[str, Tuple[List[str], List[str]]]:
        """Get a preview of XML-level changes by running operations on a temporary copy.

        Returns a dict mapping section name to tuple of (original_xml_lines, modified_xml_lines).
        """

        try:
            # Create temporary copy
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_path = Path(temp_file.name)

            shutil.copy2(file_path, temp_path)

            try:
                # Get both original and modified XML content in one operation
                return self._compare_xml_contents(file_path, temp_path)

            finally:
                temp_path.unlink(missing_ok=True)

        except Exception as e:
            logging.getLogger(__name__).error("Error previewing XML changes for %s: %s", file_path, e)
            return {}
    
    def _compare_xml_contents(self, original_path: Path, temp_path: Path) -> Dict[str, Tuple]:
        """Compare original and modified XML contents efficiently."""
        # Get original XML content
        original_doc = Document(original_path)
        original_xml = self._extract_all_content(original_doc, extract_xml=True)

        # Apply modifications to temporary copy
        self.modify_docx(temp_path)

        # Get modified XML content
        modified_doc = Document(temp_path)
        modified_xml = self._extract_all_content(modified_doc, extract_xml=True)

        # Find differences
        changes: Dict[str, Tuple[List[str], List[str]]] = {}
        for section_name in original_xml.keys():
            if section_name in modified_xml:
                orig_lines = original_xml[section_name]
                mod_lines = modified_xml[section_name]
                if orig_lines != mod_lines:
                    changes[section_name] = (orig_lines, mod_lines)

        return changes

    def replace_table_cell(self, doc: Document, cell_config: Dict) -> bool:
        """Replace content in a specific table cell.

        Args:
            doc: The Document to modify
            cell_config: Configuration dict with keys:
                - row: Row index (0-based)
                - column: Column index (0-based)
                - replace: New content (supports formatting tokens)
                - table_index: Table index (0-based, optional)
                - table_header: Header text to match for finding table (optional)
                - search: Expected current content for validation (optional)

        Returns:
            True if replacement was made, False otherwise
        """
        row_index = cell_config['row']
        col_index = cell_config['column']
        new_content = cell_config['replace']
        expected_content = cell_config.get('search')
        table_index = cell_config.get('table_index')
        table_header = cell_config.get('table_header')

        try:
            if table_index is not None:
                # Use specified table index
                if table_index >= len(doc.tables):
                    self._logger.warning(f"Table index {table_index} not found (only {len(doc.tables)} tables exist)")
                    return False

                target_table = doc.tables[table_index]
                target_table_index = table_index
            elif table_header is not None:
                # Find table by header content
                target_table = None
                target_table_index = None

                for i, table in enumerate(doc.tables):
                    if len(table.rows) > 0:
                        # Check if header row matches the specified header pattern
                        header_row = table.rows[0]

                        # Try exact match first (tab-separated or comma-separated)
                        header_text_tab = '\t'.join(cell.text.strip() for cell in header_row.cells)
                        header_text_comma = ', '.join(cell.text.strip() for cell in header_row.cells)
                        header_text_space = ' '.join(cell.text.strip() for cell in header_row.cells)

                        if (table_header == header_text_tab or
                            table_header == header_text_comma or
                            table_header == header_text_space or
                            table_header in header_text_space):  # Fallback to contains for partial matches
                            target_table = table
                            target_table_index = i
                            break

                if target_table is None:
                    self._logger.warning(f"No table found with header matching '{table_header}'")
                    return False
            else:
                # Default to first table if no specification provided
                if not doc.tables:
                    self._logger.warning("No tables found in document")
                    return False

                target_table = doc.tables[0]
                target_table_index = 0

            # Check if row exists
            if row_index >= len(target_table.rows):
                self._logger.warning(f"Row index {row_index} not found in table {target_table_index} (only {len(target_table.rows)} rows exist)")
                return False

            row = target_table.rows[row_index]

            # Check if column exists
            if col_index >= len(row.cells):
                self._logger.warning(f"Column index {col_index} not found in table {target_table_index} row {row_index} (only {len(row.cells)} cells exist)")
                return False

            cell = row.cells[col_index]

            # Validate current content if search parameter provided
            if expected_content is not None:
                current_content = cell.text.strip()
                if current_content != expected_content:
                    self._logger.warning(f"Table {target_table_index}[{row_index},{col_index}] content '{current_content}' does not match expected '{expected_content}'")
                    return False

            # Replace cell content using the established text replacement system
            self._set_cell_content(cell, new_content)

            self._logger.info(f"Replaced content in table {target_table_index}[{row_index},{col_index}] with: {new_content}")
            return True

        except Exception as e:
            self._logger.error(f"Error replacing table cell: {e}")
            return False

    def _set_cell_content(self, cell, content: str) -> None:
        """Set cell content using the established pattern from _rebuild_paragraph_basic.

        This follows the exact same pattern used in text_replacement.py for consistent
        handling of formatting tokens and font property preservation.
        """
        # Get the first paragraph
        para = cell.paragraphs[0]

        # Store original font formatting from first run if available (same as _rebuild_paragraph_basic)
        original_font_formatting = FontFormatter.get_base_font_formatting(para.runs)

        # Clear all runs (same as _rebuild_paragraph_basic)
        self.text_replacer._clear_paragraph(para)

        if not content:
            return

        # Process formatting tokens in the new text (same as _rebuild_paragraph_basic)
        text_segments = self.formatter.process_formatting_tokens(content, para)

        # Add runs with the new text and formatting (same as _rebuild_paragraph_basic)
        for text, formatting in text_segments:
            if text:  # Only create runs for non-empty text
                run = para.add_run(text)

                # Apply original formatting as base (same as _rebuild_paragraph_basic)
                FontFormatter.apply_font_properties(run, original_font_formatting)

                # Apply new formatting from tokens, but only properties that were explicitly specified
                # Filter out default False values that weren't actually specified in the formatting token
                filtered_formatting = {}
                for key, value in formatting.items():
                    # Only include properties that were explicitly set (not default False values)
                    if key in ['alignment', 'font_size', 'font_name', 'space_after', 'space_before'] or value is True:
                        filtered_formatting[key] = value
                    elif key in ['line_break_after', 'paragraph_break_after', 'page_break_after'] and value is True:
                        filtered_formatting[key] = value

                self.formatter.apply_formatting_to_run(run, filtered_formatting, para)

                # Apply paragraph-level formatting (alignment, etc.)
                self.formatter.apply_paragraph_formatting(para, formatting)
