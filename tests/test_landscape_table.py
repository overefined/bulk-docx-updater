"""
Unit tests for the landscape_table operation.

landscape_table wraps a located table in its own landscape section, leaving
the surrounding content in its original (portrait) orientation.
"""
import json
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT

from src.document_processor import DocxBulkUpdater
from src.config import load_operations_from_json


def _make_doc_with_table():
    """Portrait doc: intro paragraph, a table, trailing paragraph."""
    doc = Document()
    doc.add_paragraph("Intro portrait text")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Spectrum"
    table.rows[0].cells[1].text = "Phase"
    table.rows[1].cells[0].text = "0013BKG"
    doc.add_paragraph("Trailing portrait text")
    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)
    return path


def _pgsz(section):
    pgSz = section._sectPr.find(qn('w:pgSz'))
    return pgSz.get(qn('w:w')), pgSz.get(qn('w:h'))


class TestLandscapeTable:

    def test_wraps_table_in_landscape_section(self):
        path = _make_doc_with_table()
        processor = DocxBulkUpdater([{
            "op": "landscape_table", "match": "Spectrum",
            "margins": "0.5,0.5,0.5,0.5",
        }])
        assert processor.modify_docx(path) is True

        doc = Document(path)
        # Original single section becomes three: portrait / landscape / portrait
        assert len(doc.sections) == 3
        assert doc.sections[0].orientation == WD_ORIENT.PORTRAIT
        assert doc.sections[1].orientation == WD_ORIENT.LANDSCAPE
        assert doc.sections[2].orientation == WD_ORIENT.PORTRAIT

        # Landscape section has swapped page dimensions
        w, h = _pgsz(doc.sections[1])
        assert (w, h) == ("15840", "12240")
        # ...and the requested 0.5" margins (0.5 * 914400 EMU)
        assert doc.sections[1].left_margin == 457200
        assert doc.sections[1].right_margin == 457200

        # Surrounding portrait sections keep the original page size
        assert _pgsz(doc.sections[0]) == ("12240", "15840")
        assert _pgsz(doc.sections[2]) == ("12240", "15840")

    def test_table_sits_in_the_landscape_section(self):
        path = _make_doc_with_table()
        DocxBulkUpdater([{"op": "landscape_table", "match": "Spectrum"}]).modify_docx(path)

        doc = Document(path)
        body = doc.element.body
        order = [el.tag.split('}')[1] for el in body]
        # intro p, section-break p, table, landscape-break p, trailing p, final sectPr
        assert order == ['p', 'p', 'tbl', 'p', 'p', 'sectPr']

        # The paragraph immediately after the table carries the landscape break
        tbl = doc.tables[0]._tbl
        after = tbl.getnext()
        sectPr = after.find(qn('w:pPr')).find(qn('w:sectPr'))
        assert sectPr.find(qn('w:pgSz')).get(qn('w:orient')) == 'landscape'

        # Content is preserved
        assert "Intro portrait text" in doc.paragraphs[0].text
        assert any("Trailing portrait text" in p.text for p in doc.paragraphs)

    def test_default_margins_half_inch(self):
        path = _make_doc_with_table()
        DocxBulkUpdater([{"op": "landscape_table", "match": "Spectrum"}]).modify_docx(path)
        doc = Document(path)
        sec = doc.sections[1]
        assert sec.top_margin == 457200
        assert sec.bottom_margin == 457200
        assert sec.left_margin == 457200
        assert sec.right_margin == 457200

    def test_idempotent(self):
        path = _make_doc_with_table()
        processor = DocxBulkUpdater([{"op": "landscape_table", "match": "Spectrum"}])
        assert processor.modify_docx(path) is True
        # Second run is a no-op (table already in a landscape section)
        assert processor.modify_docx(path) is False
        assert len(Document(path).sections) == 3

    def test_missing_table_returns_false(self):
        path = _make_doc_with_table()
        processor = DocxBulkUpdater([{"op": "landscape_table", "match": "NoSuchTable"}])
        assert processor.modify_docx(path) is False
        assert len(Document(path).sections) == 1

    def test_locate_by_table_index(self):
        path = _make_doc_with_table()
        processor = DocxBulkUpdater([{"op": "landscape_table", "table_index": 0}])
        assert processor.modify_docx(path) is True
        assert Document(path).sections[1].orientation == WD_ORIENT.LANDSCAPE

    def test_already_landscape_only_adjusts_margins(self):
        # Build a doc whose single section is already landscape with 1" margins
        doc = Document()
        sec = doc.sections[0]
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
        from docx.shared import Inches
        sec.left_margin = sec.right_margin = Inches(1.0)
        table = doc.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "Spectrum"
        path = tempfile.mktemp(suffix=".docx")
        doc.save(path)

        processor = DocxBulkUpdater([{
            "op": "landscape_table", "match": "Spectrum", "margins": "0.5,0.5,0.5,0.5"}])
        assert processor.modify_docx(path) is True

        result = Document(path)
        # No redundant section was added; margins were tightened in place
        assert len(result.sections) == 1
        assert result.sections[0].orientation == WD_ORIENT.LANDSCAPE
        assert result.sections[0].left_margin == 457200
        assert result.sections[0].right_margin == 457200
        # Second run is a no-op (already landscape with requested margins)
        assert processor.modify_docx(path) is False

    def test_config_loading_dict_form(self, tmp_path):
        config = tmp_path / "cfg.json"
        config.write_text(json.dumps({
            "landscape_table": {"match": "Spectrum", "margins": "0.5,0.5,0.5,0.5"}
        }))
        operations, _ = load_operations_from_json(config)
        assert len(operations) == 1
        assert operations[0]["op"] == "landscape_table"
        assert operations[0]["match"] == "Spectrum"

    def test_config_loading_list_form(self, tmp_path):
        config = tmp_path / "cfg.json"
        config.write_text(json.dumps({
            "landscape_table": [
                {"match": "run1"},
                {"match": "run2"},
            ]
        }))
        operations, _ = load_operations_from_json(config)
        assert [op["op"] for op in operations] == ["landscape_table", "landscape_table"]
