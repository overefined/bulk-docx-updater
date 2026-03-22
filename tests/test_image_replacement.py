"""
Unit tests for image replacement functionality.

Tests the replace_image operation that allows replacing images in a document
by name, alt text, or index.
"""
import pytest
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

from src.document_processor import DocxBulkUpdater
from src.config import load_operations_from_json


class TestImageReplacement:
    """Test cases for image replacement functionality."""

    TEMPLATES_DIR = Path("/mnt/c/Development/scripts/docx-templates/templates")
    REPLACEMENT_IMAGE = Path("/mnt/c/Development/scripts/docx-templates/Alliance-logo_LG_cropped.png")

    def find_source_doc(self) -> Path:
        """Find any available docx in the test templates directory."""
        matches = sorted(self.TEMPLATES_DIR.glob("*.docx"))
        if not matches:
            pytest.skip("No test document found in test_templates/")
        return matches[0]

    def get_image_names(self, doc_path: Path) -> list[str]:
        """Return the name attributes of all images in a document."""
        doc = Document(doc_path)
        names = []
        for para in doc.paragraphs:
            for run in para.runs:
                for drawing in run._element.iter(qn('wp:docPr')):
                    name = drawing.get('name')
                    if name:
                        names.append(name)
        return names

    def get_image_alt_texts(self, doc_path: Path) -> list[str]:
        """Return the alt text (descr) of all images in a document."""
        doc = Document(doc_path)
        alt_texts = []
        for para in doc.paragraphs:
            for run in para.runs:
                for drawing in run._element.iter(qn('wp:docPr')):
                    descr = drawing.get('descr')
                    if descr:
                        alt_texts.append(descr)
        return alt_texts

    def test_replace_first_image_by_default(self, tmp_path):
        """Test replacing the first image in a document (default behavior)."""
        source_doc = self.find_source_doc()
        if not source_doc.exists():
            pytest.skip("Test document not found")

        import shutil
        test_file = tmp_path / source_doc.name
        shutil.copy2(source_doc, test_file)

        # Get original image info
        original_doc = Document(test_file)
        original_image_part = None
        for para in original_doc.paragraphs:
            for run in para.runs:
                drawings = run._element.findall(qn('w:drawing'))
                if drawings:
                    drawing = drawings[0]
                    blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    if blip is not None:
                        rel_id = blip.get(qn('r:embed'))
                        original_image_part = original_doc.part.related_parts[rel_id]
                        break
            if original_image_part:
                break

        assert original_image_part is not None, "No image found in test document"
        original_size = len(original_image_part._blob)

        replacement_image = self.REPLACEMENT_IMAGE
        if not replacement_image.exists():
            pytest.skip("Replacement image not found")

        operations = [{
            'op': 'replace_image',
            'image_path': str(replacement_image)
        }]

        # Apply replacement
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        # Verify replacement was applied
        assert result is True

        # Check that the image data changed
        modified_doc = Document(test_file)
        modified_image_part = None
        for para in modified_doc.paragraphs:
            for run in para.runs:
                drawings = run._element.findall(qn('w:drawing'))
                if drawings:
                    drawing = drawings[0]
                    blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    if blip is not None:
                        rel_id = blip.get(qn('r:embed'))
                        modified_image_part = modified_doc.part.related_parts[rel_id]
                        break
            if modified_image_part:
                break

        assert modified_image_part is not None
        modified_size = len(modified_image_part._blob)

        # Image sizes should be different
        assert modified_size != original_size

        # Modified image size should match replacement image size
        replacement_size = replacement_image.stat().st_size
        assert modified_size == replacement_size

    def test_replace_image_by_name(self, tmp_path):
        """Test replacing an image by its name attribute."""
        source_doc = self.find_source_doc()
        import shutil
        test_file = tmp_path / source_doc.name
        shutil.copy2(source_doc, test_file)

        replacement_image = self.REPLACEMENT_IMAGE
        if not replacement_image.exists():
            pytest.skip("Replacement image not found")

        names = self.get_image_names(test_file)
        if not names:
            pytest.skip("No named images found in test document")
        operations = [{
            'op': 'replace_image',
            'image_path': str(replacement_image),
            'name': names[0]
        }]

        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        assert result is True

    def test_replace_image_by_index(self, tmp_path):
        """Test replacing an image by its index."""
        source_doc = self.find_source_doc()
        import shutil
        test_file = tmp_path / source_doc.name
        shutil.copy2(source_doc, test_file)

        replacement_image = self.REPLACEMENT_IMAGE
        if not replacement_image.exists():
            pytest.skip("Replacement image not found")

        # Replace image at index 0 (first image)
        operations = [{
            'op': 'replace_image',
            'image_path': str(replacement_image),
            'index': 0
        }]

        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        assert result is True

    def test_replace_image_by_alt_text(self, tmp_path):
        """Test replacing an image by its alt text."""
        source_doc = self.find_source_doc()
        import shutil
        test_file = tmp_path / source_doc.name
        shutil.copy2(source_doc, test_file)

        replacement_image = self.REPLACEMENT_IMAGE
        if not replacement_image.exists():
            pytest.skip("Replacement image not found")

        alt_texts = self.get_image_alt_texts(test_file)
        if not alt_texts:
            pytest.skip("No images with alt text found in test document")
        operations = [{
            'op': 'replace_image',
            'image_path': str(replacement_image),
            'alt_text': alt_texts[0]
        }]

        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        assert result is True

    def test_replace_nonexistent_image(self, tmp_path):
        """Test handling of replacement when image file doesn't exist."""
        source_doc = self.find_source_doc()
        import shutil
        test_file = tmp_path / source_doc.name
        shutil.copy2(source_doc, test_file)

        # Try to replace with non-existent image
        operations = [{
            'op': 'replace_image',
            'image_path': 'nonexistent_image.png'
        }]

        processor = DocxBulkUpdater(operations)
        doc = Document(test_file)
        result = processor.replace_image(doc, operations[0])

        assert result is False

    def test_replace_image_with_invalid_name(self, tmp_path):
        """Test handling when specified image name doesn't exist."""
        source_doc = self.find_source_doc()
        import shutil
        test_file = tmp_path / source_doc.name
        shutil.copy2(source_doc, test_file)

        replacement_image = self.REPLACEMENT_IMAGE
        if not replacement_image.exists():
            pytest.skip("Replacement image not found")

        # Try to replace image with non-existent name
        operations = [{
            'op': 'replace_image',
            'image_path': str(replacement_image),
            'name': 'NonExistentPicture'
        }]

        processor = DocxBulkUpdater(operations)
        doc = Document(test_file)
        result = processor.replace_image(doc, operations[0])

        assert result is False

    def test_replace_image_with_invalid_index(self, tmp_path):
        """Test handling when specified index is out of range."""
        source_doc = self.find_source_doc()
        import shutil
        test_file = tmp_path / source_doc.name
        shutil.copy2(source_doc, test_file)

        replacement_image = self.REPLACEMENT_IMAGE
        if not replacement_image.exists():
            pytest.skip("Replacement image not found")

        # Try to replace image at invalid index
        operations = [{
            'op': 'replace_image',
            'image_path': str(replacement_image),
            'index': 99
        }]

        processor = DocxBulkUpdater(operations)
        doc = Document(test_file)
        result = processor.replace_image(doc, operations[0])

        assert result is False

    def test_config_integration(self, tmp_path):
        """Test image replacement through config file integration."""
        source_doc = self.find_source_doc()
        import shutil
        test_file = tmp_path / source_doc.name
        shutil.copy2(source_doc, test_file)

        replacement_image = self.REPLACEMENT_IMAGE
        if not replacement_image.exists():
            pytest.skip("Replacement image not found")

        # Use absolute path in config to avoid path resolution issues
        absolute_image_path = replacement_image.absolute()

        # Create config file
        names = self.get_image_names(test_file)
        if not names:
            pytest.skip("No named images found in test document")
        config_data = {
            "replace_image": [{
                "image_path": str(absolute_image_path),
                "name": names[0]
            }]
        }

        config_file = tmp_path / "test_config.json"
        import json
        config_file.write_text(json.dumps(config_data))

        # Load config and apply replacement
        operations, _ = load_operations_from_json(config_file)
        processor = DocxBulkUpdater(operations)
        result = processor.modify_docx(test_file)

        assert result is True
