#!/usr/bin/env python3
"""
Simple, direct approach to table cell alignment after text replacement.
This bypasses the complex formatting token system and just applies alignment
directly to table cells containing specific text patterns.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_table_cell_left_alignment(docx_path: str, target_patterns: list):
    """
    Simple function to left-align table cells containing specific text patterns.

    Args:
        docx_path: Path to the DOCX file
        target_patterns: List of text patterns to search for in table cells
    """
    try:
        doc = Document(docx_path)

        cells_modified = 0

        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        para_text = paragraph.text.strip()

                        # Check if this paragraph contains any of our target patterns
                        if any(pattern in para_text for pattern in target_patterns):
                            # Set alignment to LEFT
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            cells_modified += 1
                            print(f"Set LEFT alignment: Table {table_idx+1}, Row {row_idx+1}, Cell {cell_idx+1}: '{para_text[:40]}...'")

        # Save the document
        doc.save(docx_path)
        print(f"Successfully modified {cells_modified} table cell paragraphs to LEFT alignment")
        return cells_modified

    except Exception as e:
        print(f"Error applying table cell alignment: {e}")
        return 0

if __name__ == "__main__":
    import sys

    if len(sys.argv) != 2:
        print("Usage: python simple_table_alignment.py <docx_file>")
        sys.exit(1)

    docx_file = sys.argv[1]
    target_patterns = ["o2.phase_fmtd", "o2.ReadingTimestamp"]

    apply_table_cell_left_alignment(docx_file, target_patterns)