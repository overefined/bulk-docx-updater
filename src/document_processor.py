"""
Main document processing class for DOCX bulk updates.

Contains the DocxBulkUpdater class and methods for document-level operations
like margin standardization, paragraph cleanup, and change preview.
"""
from __future__ import annotations
import re
import difflib
import tempfile
import shutil
from pathlib import Path
from typing import List, Dict, Optional, Tuple
import logging
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

from src.formatting import FormattingProcessor
from src.text_replacement import TextReplacer
from src.font_utils import FontFormatter


class DocxBulkUpdater:
    """Main class for bulk DOCX document processing and text replacement."""
    
    def __init__(self, operations: List[Dict[str, any]], preserve_formatting: bool = True, 
                 standardize_margins: bool = False, margins: Optional[Dict[str, float]] = None,
                 diff_context: int = 3):
        self.operations = operations
        self.preserve_formatting = preserve_formatting
        self.standardize_margins = standardize_margins
        self.margins = margins or {
            'top': 1.0,
            'bottom': 1.0,
            'left': 1.0,
            'right': 1.0
        }
        self.diff_context = diff_context
        self._logger = logging.getLogger(__name__)
        
        # Initialize component processors
        self.formatter = FormattingProcessor()
        self.text_replacer = TextReplacer(operations, self.formatter)
        
        # Pre-compute cross-paragraph patterns for optimization
        self._cross_paragraph_patterns = self._get_cross_paragraph_search_patterns()
        self._has_cross_paragraph_patterns = len(self._cross_paragraph_patterns) > 0
        
        # Performance optimization caches
        self._paragraph_cache = {}
        self._text_cache = {}
        self._xml_cache = {}  # Cache XML strings to reduce xpath calls
        
        # No separate global search set needed; TextReplacer handles patterns
        
    def clear_caches(self):
        """Clear performance caches to free memory."""
        self._paragraph_cache.clear()
        self._text_cache.clear()
        self._xml_cache.clear()
        self.text_replacer.clear_caches()
    
    
    def _iter_all_paragraphs(self, doc: Document):
        """Yield all paragraphs across body, tables, headers, and footers.
        
        Uses caching to avoid repeated paragraph extraction.
        """
        doc_id = id(doc)
        if doc_id in self._paragraph_cache:
            yield from self._paragraph_cache[doc_id]
            return
        
        paragraphs = []
        
        # Body paragraphs
        paragraphs.extend(doc.paragraphs)
        
        # Tables (all cells' paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        
        # Headers and footers for all sections
        for section in doc.sections:
            paragraphs.extend(section.header.paragraphs)
            paragraphs.extend(section.footer.paragraphs)
        
        # Cache the full list for subsequent calls
        self._paragraph_cache[doc_id] = paragraphs
        
        # Yield all at once after building complete list
        yield from paragraphs
    
    
    def _get_cross_paragraph_search_patterns(self) -> List[str]:
        """Pre-compute list of search patterns that might span paragraphs."""
        patterns = []
        for op in self.operations:
            if op.get('op') != 'replace':
                continue
            if 'search' not in op or 'replace' not in op:
                continue
            patterns.append(op['search'])
        return patterns
    
    def _chunk_has_cross_paragraph_potential(self, paragraphs: List, paragraph_texts: List[str] = None) -> bool:
        """Check if a chunk of paragraphs might contain cross-paragraph patterns (optimized)."""
        if len(paragraphs) < 2 or not self._has_cross_paragraph_patterns:
            return False
        
        # Use cached paragraph texts if provided
        if paragraph_texts is None:
            paragraph_texts = [para.text for para in paragraphs]
        
        # Early exit: if all paragraphs are empty, no potential
        if not any(text.strip() for text in paragraph_texts):
            return False
        
        # Combine text from all paragraphs (using pre-computed texts)
        combined_text = "".join(paragraph_texts)
        if not combined_text.strip():
            return False
        
        # Check pre-computed patterns only
        for search_text in self._cross_paragraph_patterns:
            if search_text in combined_text:
                # Check if this pattern actually spans paragraphs
                # (i.e., it doesn't appear complete in any single paragraph)
                spans_paragraphs = True
                for text in paragraph_texts:
                    if search_text in text:
                        spans_paragraphs = False
                        break
                
                if spans_paragraphs:
                    return True
        
        return False
    
    def _process_all_text_replacements(self, doc: Document) -> bool:
        """Efficiently process both cross-paragraph and single-paragraph replacements."""
        modified = False
        processed_paragraphs = set()  # Track which paragraphs were already processed
        
        # First, handle cross-paragraph replacements if any exist
        if self._has_cross_paragraph_patterns:
            # Process body paragraphs in chunks
            body_paragraphs = list(doc.paragraphs)
            if body_paragraphs:
                if self._process_paragraph_chunks_tracked(body_paragraphs, processed_paragraphs):
                    modified = True
            
            # Process table cell paragraphs in chunks
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_paragraphs = list(cell.paragraphs)
                        if cell_paragraphs:
                            if self._process_paragraph_chunks_tracked(cell_paragraphs, processed_paragraphs):
                                modified = True
            
            # Process header/footer paragraphs in chunks
            for section in doc.sections:
                if hasattr(section, 'header') and section.header:
                    header_paragraphs = list(section.header.paragraphs)
                    if header_paragraphs:
                        if self._process_paragraph_chunks_tracked(header_paragraphs, processed_paragraphs):
                            modified = True
                            
                if hasattr(section, 'footer') and section.footer:
                    footer_paragraphs = list(section.footer.paragraphs)
                    if footer_paragraphs:
                        if self._process_paragraph_chunks_tracked(footer_paragraphs, processed_paragraphs):
                            modified = True
        
        # Then process remaining single-paragraph replacements
        for paragraph in self._iter_all_paragraphs(doc):
            # Skip paragraphs that were already processed in cross-paragraph chunks
            if id(paragraph) not in processed_paragraphs:
                if self.text_replacer.replace_text_in_paragraph(paragraph):
                    modified = True
        
        return modified
    
    def _process_paragraph_chunks_tracked(self, paragraphs: List, processed_paragraphs: set) -> bool:
        """Process paragraph chunks and track which paragraphs were processed."""
        if len(paragraphs) < 2:
            return False
        
        # Pre-compute all paragraph texts once to avoid repeated .text calls
        # Use caching to avoid repeated text extraction
        paragraph_texts = []
        for para in paragraphs:
            para_id = id(para)
            if para_id in self._text_cache:
                paragraph_texts.append(self._text_cache[para_id])
            else:
                text = para.text
                self._text_cache[para_id] = text
                paragraph_texts.append(text)
        
        # Early exit if no cross-paragraph patterns could possibly exist
        combined_text = "".join(paragraph_texts)
        # Use pre-compiled set for faster pattern matching
        if not any(pattern in combined_text for pattern in self._cross_paragraph_patterns):
            return False
        
        modified = False
        max_chunk_size = 5  # Maximum paragraphs to consider at once
        
        # Try different chunk sizes, starting with larger chunks
        for chunk_size in range(min(max_chunk_size, len(paragraphs)), 1, -1):
            i = 0
            while i <= len(paragraphs) - chunk_size:
                chunk = paragraphs[i:i + chunk_size]
                chunk_texts = paragraph_texts[i:i + chunk_size]
                
                # Check if this chunk has any potential cross-paragraph patterns
                if self._chunk_has_cross_paragraph_potential(chunk, chunk_texts):
                    if self.text_replacer.replace_text_across_paragraphs(chunk):
                        modified = True
                        # Mark all paragraphs in this chunk as processed
                        for para in chunk:
                            processed_paragraphs.add(id(para))
                        # Skip ahead since we processed this chunk
                        i += chunk_size
                        continue
                
                i += 1
        
        return modified
    
    def standardize_document_margins(self, doc: Document) -> bool:
        """Standardize margins for all sections in the document."""
        if not self.standardize_margins:
            return False

        modified = False
        for section in doc.sections:
            # Set margins using Inches for consistency
            section.top_margin = Inches(self.margins['top'])
            section.bottom_margin = Inches(self.margins['bottom'])
            section.left_margin = Inches(self.margins['left'])
            section.right_margin = Inches(self.margins['right'])
            modified = True

        return modified

    def clear_core_properties(self, doc: Document, properties: List[str]) -> bool:
        """Clear specified core document properties.

        Args:
            doc: The Document object
            properties: List of property names to clear (e.g., ['author', 'company', 'title'])

        Returns:
            True if any properties were cleared, False otherwise

        Supported properties:
            - title: Document title
            - subject: Document subject
            - author: Document author/creator
            - keywords: Document keywords
            - comments: Document comments
            - last_modified_by: Last modified by
            - category: Document category
            - content_status: Content status
            - company: Company name (from app.xml extended properties)
        """
        from lxml import etree

        modified = False
        core_props = doc.core_properties

        for prop in properties:
            try:
                if prop == 'title':
                    if core_props.title:
                        self._logger.debug(f"Clearing title: '{core_props.title}'")
                        core_props.title = ''
                        modified = True
                elif prop == 'subject':
                    if core_props.subject:
                        self._logger.debug(f"Clearing subject: '{core_props.subject}'")
                        core_props.subject = ''
                        modified = True
                elif prop == 'author':
                    if core_props.author:
                        self._logger.debug(f"Clearing author: '{core_props.author}'")
                        core_props.author = ''
                        modified = True
                elif prop == 'keywords':
                    if core_props.keywords:
                        self._logger.debug(f"Clearing keywords: '{core_props.keywords}'")
                        core_props.keywords = ''
                        modified = True
                elif prop == 'comments':
                    if core_props.comments:
                        self._logger.debug(f"Clearing comments")
                        core_props.comments = ''
                        modified = True
                elif prop == 'last_modified_by':
                    if core_props.last_modified_by:
                        self._logger.debug(f"Clearing last_modified_by: '{core_props.last_modified_by}'")
                        core_props.last_modified_by = ''
                        modified = True
                elif prop == 'category':
                    if core_props.category:
                        self._logger.debug(f"Clearing category: '{core_props.category}'")
                        core_props.category = ''
                        modified = True
                elif prop == 'content_status':
                    if core_props.content_status:
                        self._logger.debug(f"Clearing content_status: '{core_props.content_status}'")
                        core_props.content_status = ''
                        modified = True
                elif prop == 'company':
                    # Company is in app.xml (extended properties)
                    # Access it directly via package parts
                    app_part = self._get_app_xml_part(doc)
                    if app_part:
                        tree = etree.fromstring(app_part.blob)
                        APP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
                        namespaces = {'ep': APP_NS}

                        company_elem = tree.find('.//ep:Company', namespaces=namespaces)
                        if company_elem is not None:
                            self._logger.debug(f"Clearing Company property: '{company_elem.text}'")
                            tree.remove(company_elem)
                            app_part._blob = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
                            modified = True
                else:
                    self._logger.warning(f"Unknown property: '{prop}'")

            except Exception as e:
                self._logger.warning(f"Error clearing property '{prop}': {e}")

        return modified

    def _get_app_xml_part(self, doc: Document):
        """Get the app.xml part from the document package.

        Args:
            doc: The Document object

        Returns:
            The app.xml Part object, or None if not found
        """
        try:
            package = doc.part.package
            for part in package.iter_parts():
                if str(part.partname) == '/docProps/app.xml':
                    return part
            return None
        except Exception as e:
            self._logger.debug(f"Error accessing app.xml part: {e}")
            return None

    def _replace_placeholders(self, value: str, file_path: Path) -> str:
        """Replace placeholders in a value with file-based information.

        Args:
            value: The value that may contain placeholders
            file_path: Path to the current document being processed

        Returns:
            Value with placeholders replaced

        Supported placeholders:
            {{FILENAME}} - Full filename with extension (e.g., "invoice.docx")
            {{BASENAME}} - Filename without extension (e.g., "invoice")
            {{BASENAME_DOTX}} - Filename with .dotx extension (e.g., "invoice.dotx")
            {{EXTENSION}} - Just the extension (e.g., "docx")
            {{PARENT_DIR}} - Parent directory name
        """
        if not isinstance(value, str):
            return value

        filename = file_path.name
        basename = file_path.stem
        extension = file_path.suffix.lstrip('.')
        parent_dir = file_path.parent.name

        replacements = {
            '{{FILENAME}}': filename,
            '{{BASENAME}}': basename,
            '{{BASENAME_DOTX}}': f"{basename}.dotx",
            '{{EXTENSION}}': extension,
            '{{PARENT_DIR}}': parent_dir,
        }

        result = value
        for placeholder, replacement in replacements.items():
            result = result.replace(placeholder, replacement)

        return result

    def _has_column_break_in_run(self, run) -> bool:
        """Check if a run contains a column break."""
        # Check for column breaks in the XML
        return 'w:br' in run._element.xml and 'type="column"' in run._element.xml
    
    def _has_page_break_in_run(self, run) -> bool:
        """Check if a run contains a page break."""
        # Check for page breaks in the XML
        return 'w:br' in run._element.xml and 'type="page"' in run._element.xml
    
    def remove_empty_paragraphs_after_pattern(self, doc: Document, pattern: str) -> bool:
        """Remove column breaks and specific formatting artifacts from the next paragraph after the pattern."""
        modified = False
        
        # Find paragraphs containing the pattern
        for i, para in enumerate(doc.paragraphs):
            if pattern in para.text:
                # Check the very next paragraph
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    
                    # Only remove specific formatting artifacts, preserve intentional spacing
                    runs_to_remove = []
                    for run in next_para.runs:
                        # Check if this run contains specific breaks we want to remove
                        if self._has_column_break_in_run(run):
                            runs_to_remove.append(run)
                        elif self._has_page_break_in_run(run):
                            runs_to_remove.append(run)
                        elif not run.text:  # Completely empty runs only
                            runs_to_remove.append(run)
                        else:
                            # Stop at first run with any text content (including spaces/newlines)
                            break
                    
                    for run in runs_to_remove:
                        run._element.getparent().remove(run._element)
                        modified = True
        
        return modified

    # Ordered op-dispatch pipeline: (op name, handler method name). Each handler
    # takes (doc, op) and returns True if it modified the document. Order is
    # significant and documented inline — e.g. insert_block must run before
    # landscape_table so a freshly-inserted table can be located and rotated, and
    # section_break_before must run before divider so the divider is already on
    # its own page. (replace / xml_replace text replacements and set_comments are
    # handled separately in modify_docx because they aren't simple (doc, op)
    # handlers.)
    _OP_PIPELINE = (
        ('clear_properties',        '_op_clear_properties'),
        ('table_header_repeat',     '_op_table_header_repeat'),
        ('font_size',               '_op_font_size'),
        ('set_table_column_widths', '_op_set_table_column_widths'),
        ('cleanup_empty_after',     '_op_cleanup_empty_after'),
        ('replace_table_cell',      'replace_table_cell'),
        ('align_table_cells',       'align_table_cells'),
        ('replace_image',           'replace_image'),
        ('replace_in_table',        'replace_text_in_table'),
        ('merge_tables',            'merge_tables'),
        ('replace_table',           'replace_table'),
        ('replace_block',           'replace_block'),
        ('insert_block',            'insert_block'),
        ('remove_page_break',       'remove_page_break'),
        ('landscape_table',         'landscape_table'),
        ('format_table',            'format_table'),
        ('section_break_before',    'section_break_before'),
        ('divider',                 'divider'),
    )

    # Thin adapters so ops with non-uniform signatures still expose the common
    # (doc, op) -> bool handler shape used by _OP_PIPELINE.
    def _op_clear_properties(self, doc: Document, op: Dict) -> bool:
        return self.clear_core_properties(doc, op.get('properties', []))

    def _op_table_header_repeat(self, doc: Document, op: Dict) -> bool:
        return bool(self.set_table_header_repeat(
            doc, op.get('pattern'), enable=bool(op.get('enabled', True))))

    def _op_font_size(self, doc: Document, op: Dict) -> bool:
        from_size, to_size = op.get('from'), op.get('to')
        if from_size is None or to_size is None:
            return False
        return bool(self.change_font_sizes(doc, from_size, to_size))

    def _op_set_table_column_widths(self, doc: Document, op: Dict) -> bool:
        return bool(self.set_table_column_widths(doc, op))

    def _op_cleanup_empty_after(self, doc: Document, op: Dict) -> bool:
        pattern = op.get('pattern')
        return bool(pattern) and self.remove_empty_paragraphs_after_pattern(doc, pattern)

    def modify_docx(self, file_path: Path) -> bool:
        """Modify a DOCX file with the specified replacements."""
        try:
            # Load the document
            doc = Document(file_path)
            modified = False

            # Standardize margins if enabled
            if self.standardize_margins:
                if self.standardize_document_margins(doc):
                    modified = True

            # Apply the ordered op pipeline. Each handler takes (doc, op) and
            # returns True if it changed the document; order is set by
            # _OP_PIPELINE (see its docstring for the ordering constraints).
            for op_name, method_name in self._OP_PIPELINE:
                handler = getattr(self, method_name)
                for op in self.operations:
                    if op.get('op') == op_name and handler(doc, op):
                        modified = True

            # Then do the text replacements and inserts
            has_search_ops = any(op.get('op') in ('replace', 'xml_replace') for op in self.operations)

            if has_search_ops:
                # Process both cross-paragraph and single-paragraph replacements efficiently
                if self._process_all_text_replacements(doc):
                    modified = True

            # Handle setting Comments field to store template filename
            for op in self.operations:
                if op.get('op') == 'set_comments':
                    value = op.get('value', '')
                    value = self._replace_placeholders(value, file_path)
                    doc.core_properties.comments = value
                    self._logger.debug(f"Set Comments to '{value}'")
                    modified = True

            # Save changes if any modifications were made
            if modified:
                doc.save(file_path)

            return modified
            
        except Exception as e:
            logging.getLogger(__name__).error("Error processing %s: %s", file_path, e)
            return False
    
    
    def get_document_changes_preview(self, file_path: Path) -> Dict[str, str]:
        """Get a preview of changes by running actual operations on a temporary copy."""

        try:
            # Create temporary copy
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_path = Path(temp_file.name)

            shutil.copy2(file_path, temp_path)

            try:
                # Track formatting operations and get content changes
                return self._compare_document_contents_with_formatting(file_path, temp_path)

            finally:
                # Clean up temporary file
                temp_path.unlink(missing_ok=True)

        except Exception as e:
            logging.getLogger(__name__).error("Error previewing changes for %s: %s", file_path, e)
            return {}
    
    def _compare_document_contents_with_formatting(self, original_path: Path, temp_path: Path) -> Dict[str, str]:
        """Compare original and modified document contents and track formatting operations."""
        # Clear cache to ensure fresh content extraction
        self._text_cache.clear()

        # Get original content and properties
        original_doc = Document(original_path)
        original_content = self._extract_all_content(original_doc, extract_xml=False)
        original_properties = self._extract_document_properties(original_doc)

        # Track operations performed during modification
        operation_results = []

        # Apply modifications to temporary copy while tracking operations
        modified_doc = Document(temp_path)

        # Track formatting operations
        for op in self.operations:
            if op.get('op') == 'table_header_repeat':
                pat = op.get('pattern')
                enable = True if op.get('enabled', True) else False
                count = self.set_table_header_repeat(modified_doc, pat, enable=enable)
                if count > 0:
                    action = "Set" if enable else "Unset"
                    pattxt = f" '{pat}'" if pat else " (first row)"
                    operation_results.append(f"{action} table header repeat on {count} row(s){pattxt}")

            if op.get('op') == 'font_size':
                from_size = op.get('from')
                to_size = op.get('to')
                if from_size is not None and to_size is not None:
                    count = self.change_font_sizes(modified_doc, from_size, to_size)
                    if count > 0:
                        operation_results.append(f"Changed font size from {from_size}pt to {to_size}pt in {count} text run(s)")

            if op.get('op') == 'set_table_column_widths':
                count = self.set_table_column_widths(modified_doc, op)
                if count > 0:
                    table_header = op.get('table_header', 'first table')
                    table_index = op.get('table_index')
                    if table_index is not None:
                        table_desc = f"table {table_index}"
                    else:
                        table_desc = f"table with header '{table_header}'" if table_header != 'first table' else table_header
                    column_widths = op.get('column_widths', [])
                    operation_results.append(f"Set column widths for {table_desc}: {column_widths}")

            if op.get('op') == 'replace_table':
                if self.replace_table(modified_doc, op):
                    if 'table_index' in op:
                        loc = f"table {op['table_index']}"
                    elif 'table_header' in op:
                        loc = f"table with header '{op['table_header']}'"
                    else:
                        loc = f"table containing '{op.get('match')}'"
                    operation_results.append(f"Replaced entire {loc} with new table XML")

            if op.get('op') == 'replace_block':
                if self.replace_block(modified_doc, op):
                    act = "Replaced" if op.get('replace') else "Removed"
                    operation_results.append(
                        f"{act} block from '{op.get('from')}' to '{op.get('to')}'")

        # Apply standard text replacements
        self._process_all_text_replacements(modified_doc)

        # Save the modified document
        modified_doc.save(temp_path)

        # Clear cache again to avoid reusing original content for modified doc
        self._text_cache.clear()

        # Get modified content
        modified_doc = Document(temp_path)
        modified_content = self._extract_all_content(modified_doc, extract_xml=False)

        # Find text content differences
        changes = {}
        for section_name in original_content.keys():
            if section_name in modified_content:
                orig_lines = original_content[section_name]
                mod_lines = modified_content[section_name]

                if orig_lines != mod_lines:
                    changes[section_name] = (orig_lines, mod_lines)

        # Add formatting operation results as a special section
        if operation_results:
            changes["Formatting Operations"] = ([], operation_results)

        return changes

    def _extract_document_properties(self, doc: Document) -> Dict[str, any]:
        """Extract document properties for comparison."""
        properties = {}

        # Extract core properties
        properties['title'] = doc.core_properties.title or ''

        # Extract table header repeat properties
        table_headers = []
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                try:
                    tr_element = row._tr
                    from docx.oxml.ns import qn
                    tr_pr = tr_element.find(qn('w:trPr'))
                    if tr_pr is not None:
                        header_elem = tr_pr.find(qn('w:tblHeader'))
                        if header_elem is not None:
                            row_text = ' '.join(cell.text.strip() for cell in row.cells)
                            table_headers.append(f"Table {i+1}, Row {j+1}: {row_text}")
                except:
                    continue
        properties['table_headers'] = table_headers

        # Extract font sizes
        font_sizes = {}
        for paragraph in self._iter_all_paragraphs(doc):
            for run in paragraph.runs:
                if run.font.size is not None:
                    size_pt = int(run.font.size.pt)
                    if size_pt not in font_sizes:
                        font_sizes[size_pt] = 0
                    font_sizes[size_pt] += 1
        properties['font_sizes'] = font_sizes

        return properties

    def _extract_all_content(self, doc, extract_xml: bool = False) -> Dict[str, List[str]]:
        """Extract all content from a document organized by section.
        
        Args:
            doc: Document object
            extract_xml: If True, extract XML content; if False, extract text content
        """
        # Check cache first
        doc_id = id(doc)
        cache_key = f"{doc_id}_{extract_xml}"
        if cache_key in self._text_cache:
            return self._text_cache[cache_key]
        
        content = {}
        
        if extract_xml:
            # Body paragraphs XML
            body_xml_lines = []
            for para in doc.paragraphs:
                body_xml_lines.extend((para._p.xml or '').splitlines())
            if body_xml_lines:
                content["Body(XML)"] = body_xml_lines

            # Tables XML
            table_xml_lines = []
            for table in doc.tables:
                table_xml_lines.append("<table>")
                for row in table.rows:
                    table_xml_lines.append("  <row>")
                    for cell in row.cells:
                        table_xml_lines.append("    <cell>")
                        for para in cell.paragraphs:
                            table_xml_lines.extend((para._p.xml or '').splitlines())
                        table_xml_lines.append("    </cell>")
                    table_xml_lines.append("  </row>")
                table_xml_lines.append("</table>")
            if table_xml_lines:
                content["Tables(XML)"] = table_xml_lines

            # Headers/Footers XML
            header_footer_xml_lines = []
            for section in doc.sections:
                for para in section.header.paragraphs:
                    header_footer_xml_lines.extend((para._p.xml or '').splitlines())
                for para in section.footer.paragraphs:
                    header_footer_xml_lines.extend((para._p.xml or '').splitlines())
            if header_footer_xml_lines:
                content["Headers/Footers(XML)"] = header_footer_xml_lines
        else:
            # Body paragraphs
            content["Body"] = [para.text for para in doc.paragraphs]
            
            # Tables
            table_paragraphs = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        table_paragraphs.extend([para.text for para in cell.paragraphs])
            
            if table_paragraphs:
                content["Tables"] = table_paragraphs
            
            # Headers and footers
            header_footer_paragraphs = []
            for section in doc.sections:
                header_footer_paragraphs.extend([para.text for para in section.header.paragraphs])
                header_footer_paragraphs.extend([para.text for para in section.footer.paragraphs])
            
            if header_footer_paragraphs:
                content["Headers/Footers"] = header_footer_paragraphs
        
        # Cache the result
        self._text_cache[cache_key] = content
        return content
    
    def _extract_all_text_content(self, doc) -> Dict[str, List[str]]:
        """Extract all text content from a document organized by section."""
        return self._extract_all_content(doc, extract_xml=False)
    
    def set_table_header_repeat(self, doc: Document, header_pattern: str = None, enable: bool = True) -> int:
        """Enable or disable table header repeat on rows matching a pattern.

        Args:
            doc: Document object
            header_pattern: Text pattern to identify header rows. If None, target first row of each table.
            enable: True to set w:tblHeader, False to remove it.

        Returns:
            Number of rows modified
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        modified_count = 0

        for table in doc.tables:
            header_rows_found = []

            if header_pattern:
                # Search for rows containing the pattern
                for i, row in enumerate(table.rows):
                    row_text = ' '.join(cell.text.strip() for cell in row.cells)
                    if header_pattern in row_text:
                        header_rows_found.append(i)
            else:
                # Default: treat first row as header
                if table.rows:
                    header_rows_found.append(0)

            # Apply or remove repeat header property for found rows
            for row_idx in header_rows_found:
                try:
                    row = table.rows[row_idx]
                    tr_element = row._tr

                    # Check if trPr (table row properties) element exists
                    tr_pr = tr_element.find(qn('w:trPr'))
                    if tr_pr is None:
                        # Create trPr element if it doesn't exist
                        tr_pr = OxmlElement('w:trPr')
                        tr_element.insert(0, tr_pr)

                    # Check if tblHeader element exists
                    tbl_header = tr_pr.find(qn('w:tblHeader'))
                    if enable:
                        if tbl_header is None:
                            # Create and add tblHeader element
                            tbl_header = OxmlElement('w:tblHeader')
                            tr_pr.append(tbl_header)
                            modified_count += 1
                            self._logger.debug(f"Set repeat header for table row {row_idx}")
                    else:
                        if tbl_header is not None:
                            tr_pr.remove(tbl_header)
                            modified_count += 1
                            self._logger.debug(f"Removed repeat header for table row {row_idx}")

                except Exception as e:
                    self._logger.warning(f"Failed to set repeat header for row {row_idx}: {e}")

        return modified_count

    def change_font_sizes(self, doc: Document, from_size: int, to_size: int) -> int:
        """Change all text with a specific font size to a new font size.

        Args:
            doc: The Document object
            from_size: Original font size in points
            to_size: New font size in points

        Returns:
            Number of runs modified
        """
        from docx.shared import Pt

        modified_count = 0
        from_size_half_points = from_size * 2  # Word stores font sizes in half-points
        to_size_pt = Pt(to_size)

        # Process all paragraphs in body, tables, headers, and footers
        for paragraph in self._iter_all_paragraphs(doc):
            for run in paragraph.runs:
                # Check if run has font size property
                if run.font.size is not None:
                    # Convert to half-points for comparison (docx uses Emu internally)
                    current_size_half_points = int(run.font.size.pt * 2)
                    if current_size_half_points == from_size_half_points:
                        run.font.size = to_size_pt
                        modified_count += 1
                        self._logger.debug(f"Changed font size from {from_size}pt to {to_size}pt")

        return modified_count

    def set_table_column_widths(self, doc: Document, table_config: Dict) -> int:
        """Set column widths for tables matching specified criteria.

        Args:
            doc: The Document object
            table_config: Configuration dictionary with:
                - table_header: Header text to match for finding table (optional)
                - table_index: Zero-based table index (optional, alternative to table_header)
                - column_widths: List of column widths in inches

        Returns:
            Number of tables modified
        """
        from docx.shared import Inches

        table_header = table_config.get('table_header')
        table_index = table_config.get('table_index')
        column_widths = table_config.get('column_widths', [])

        if not column_widths:
            self._logger.warning("No column widths specified")
            return 0

        modified_count = 0
        target_table = None

        # Find the target table
        if table_index is not None:
            # Use specific table index
            if 0 <= table_index < len(doc.tables):
                target_table = doc.tables[table_index]
            else:
                self._logger.warning(f"Table index {table_index} out of range (0-{len(doc.tables)-1})")
                return 0
        elif table_header is not None:
            # Find table by header text
            for table in doc.tables:
                if table.rows:
                    # Check first row for header match
                    header_text_tab = '\t'.join(cell.text.strip() for cell in table.rows[0].cells)
                    header_text_comma = ', '.join(cell.text.strip() for cell in table.rows[0].cells)
                    header_text_space = ' '.join(cell.text.strip() for cell in table.rows[0].cells)

                    if (table_header == header_text_tab or
                        table_header == header_text_comma or
                        table_header == header_text_space or
                        table_header in header_text_space):  # Fallback to contains for partial matches
                        target_table = table
                        break

            if target_table is None:
                self._logger.warning(f"No table found with header matching '{table_header}'")
                return 0
        else:
            # Default: use first table
            if doc.tables:
                target_table = doc.tables[0]
            else:
                self._logger.warning("No tables found in document")
                return 0

        # Apply column widths to the target table
        if target_table:
            try:
                # Get the number of columns from the first row
                if not target_table.rows:
                    self._logger.warning("Target table has no rows")
                    return 0

                num_columns = len(target_table.rows[0].cells)

                # Apply widths to each column
                for col_idx in range(min(num_columns, len(column_widths))):
                    width_inches = column_widths[col_idx]
                    if width_inches > 0:  # Only set positive widths
                        for row in target_table.rows:
                            if col_idx < len(row.cells):
                                row.cells[col_idx].width = Inches(width_inches)

                modified_count = 1
                self._logger.debug(f"Set column widths for table: {column_widths[:num_columns]}")

                if len(column_widths) > num_columns:
                    self._logger.warning(f"Table has {num_columns} columns but {len(column_widths)} widths specified")
                elif len(column_widths) < num_columns:
                    self._logger.info(f"Only set widths for first {len(column_widths)} of {num_columns} columns")

            except Exception as e:
                self._logger.warning(f"Failed to set column widths: {e}")

        return modified_count

    def format_diff(self, original_lines: List[str], modified_lines: List[str], section_name: str) -> str:
        """Format a unified diff for display."""
        diff = difflib.unified_diff(
            original_lines,
            modified_lines,
            fromfile=f"original/{section_name}",
            tofile=f"modified/{section_name}",
            n=max(0, int(self.diff_context)) if isinstance(self.diff_context, int) else 3,
            lineterm=''
        )
        return '\n'.join(diff)

    def _extract_all_xml_content(self, doc: Document) -> Dict[str, List[str]]:
        """Extract XML representations of document content organized by section."""
        return self._extract_all_content(doc, extract_xml=True)

    def get_document_xml_changes_preview(self, file_path: Path) -> Dict[str, Tuple[List[str], List[str]]]:
        """Get a preview of XML-level changes by running operations on a temporary copy.

        Returns a dict mapping section name to tuple of (original_xml_lines, modified_xml_lines).
        """

        try:
            # Create temporary copy
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_path = Path(temp_file.name)

            shutil.copy2(file_path, temp_path)

            try:
                # Get both original and modified XML content in one operation
                return self._compare_xml_contents(file_path, temp_path)

            finally:
                temp_path.unlink(missing_ok=True)

        except Exception as e:
            logging.getLogger(__name__).error("Error previewing XML changes for %s: %s", file_path, e)
            return {}
    
    def _compare_xml_contents(self, original_path: Path, temp_path: Path) -> Dict[str, Tuple]:
        """Compare original and modified XML contents efficiently."""
        # Get original XML content
        original_doc = Document(original_path)
        original_xml = self._extract_all_content(original_doc, extract_xml=True)

        # Apply modifications to temporary copy
        self.modify_docx(temp_path)

        # Get modified XML content
        modified_doc = Document(temp_path)
        modified_xml = self._extract_all_content(modified_doc, extract_xml=True)

        # Find differences
        changes: Dict[str, Tuple[List[str], List[str]]] = {}
        for section_name in original_xml.keys():
            if section_name in modified_xml:
                orig_lines = original_xml[section_name]
                mod_lines = modified_xml[section_name]
                if orig_lines != mod_lines:
                    changes[section_name] = (orig_lines, mod_lines)

        return changes

    def replace_table_cell(self, doc: Document, cell_config: Dict) -> bool:
        """Replace content in a specific table cell.

        Args:
            doc: The Document to modify
            cell_config: Configuration dict with keys:
                - row: Row index (0-based)
                - column: Column index (0-based)
                - replace: New content (supports formatting tokens)
                - table_index: Table index (0-based, optional)
                - table_header: Header text to match for finding table (optional)
                - search: Expected current content for validation (optional)

        Returns:
            True if replacement was made, False otherwise
        """
        row_index = cell_config['row']
        col_index = cell_config['column']
        new_content = cell_config['replace']
        expected_content = cell_config.get('search')
        table_index = cell_config.get('table_index')
        table_header = cell_config.get('table_header')
        header_row_index = cell_config.get('header_row', 0)

        try:
            if table_index is not None:
                # Use specified table index
                if table_index >= len(doc.tables):
                    self._logger.warning(f"Table index {table_index} not found (only {len(doc.tables)} tables exist)")
                    return False

                target_table = doc.tables[table_index]
                target_table_index = table_index
            elif table_header is not None:
                # Find table by header content
                target_table = None
                target_table_index = None

                for i, table in enumerate(doc.tables):
                    if len(table.rows) <= header_row_index:
                        continue
                    # Check if header row matches the specified header pattern
                    match_row = table.rows[header_row_index]

                    # Try exact match first (tab-separated or comma-separated)
                    header_text_tab = '\t'.join(cell.text.strip() for cell in match_row.cells)
                    header_text_comma = ', '.join(cell.text.strip() for cell in match_row.cells)
                    header_text_space = ' '.join(cell.text.strip() for cell in match_row.cells)

                    if (table_header == header_text_tab or
                        table_header == header_text_comma or
                        table_header == header_text_space or
                        table_header in header_text_space):  # Fallback to contains for partial matches
                        target_table = table
                        target_table_index = i
                        break

                if target_table is None:
                    self._logger.warning(f"No table found with header matching '{table_header}' in row {header_row_index}")
                    return False
            else:
                # Default to first table if no specification provided
                if not doc.tables:
                    self._logger.warning("No tables found in document")
                    return False

                target_table = doc.tables[0]
                target_table_index = 0

            # Check if row exists
            if row_index >= len(target_table.rows):
                self._logger.warning(f"Row index {row_index} not found in table {target_table_index} (only {len(target_table.rows)} rows exist)")
                return False

            row = target_table.rows[row_index]

            # Check if column exists
            if col_index >= len(row.cells):
                self._logger.warning(f"Column index {col_index} not found in table {target_table_index} row {row_index} (only {len(row.cells)} cells exist)")
                return False

            cell = row.cells[col_index]

            # Validate current content if search parameter provided
            if expected_content is not None:
                current_content = cell.text.strip()
                if current_content != expected_content:
                    self._logger.warning(f"Table {target_table_index}[{row_index},{col_index}] content '{current_content}' does not match expected '{expected_content}'")
                    return False

            # Replace cell content using the established text replacement system
            self._set_cell_content(cell, new_content)

            self._logger.info(f"Replaced content in table {target_table_index}[{row_index},{col_index}] with: {new_content}")
            return True

        except Exception as e:
            self._logger.error(f"Error replacing table cell: {e}")
            return False

    # Fallback namespace declarations for hand-written replacement table XML that
    # only uses a subset of prefixes. A <w:tbl> copied straight out of Word already
    # carries its own xmlns declarations, so this is only used when they're absent.
    _WORD_NSDECLS = (
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
        'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" '
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
    )

    def _find_table_for_replace(self, doc: Document, op: Dict):
        """Locate the target table for a replace_table op.

        Resolution order: table_index, then table_header (header-row text match,
        like replace_table_cell), then match (substring found anywhere in the
        table's text). Returns (table, index) or (None, None).
        """
        table_index = op.get('table_index')
        table_header = op.get('table_header')
        header_row_index = op.get('header_row', 0)
        match_text = op.get('match')

        if table_index is not None:
            if 0 <= table_index < len(doc.tables):
                return doc.tables[table_index], table_index
            self._logger.warning(f"Table index {table_index} not found (only {len(doc.tables)} tables exist)")
            return None, None

        if table_header is not None:
            for i, table in enumerate(doc.tables):
                if len(table.rows) <= header_row_index:
                    continue
                match_row = table.rows[header_row_index]
                header_text_tab = '\t'.join(cell.text.strip() for cell in match_row.cells)
                header_text_comma = ', '.join(cell.text.strip() for cell in match_row.cells)
                header_text_space = ' '.join(cell.text.strip() for cell in match_row.cells)
                if (table_header == header_text_tab or
                        table_header == header_text_comma or
                        table_header == header_text_space or
                        table_header in header_text_space):
                    return table, i
            self._logger.warning(f"No table found with header matching '{table_header}' in row {header_row_index}")
            return None, None

        if match_text is not None:
            for i, table in enumerate(doc.tables):
                full_text = '\n'.join(cell.text for row in table.rows for cell in row.cells)
                if match_text in full_text:
                    return table, i
            self._logger.warning(f"No table found containing text '{match_text}'")
            return None, None

        self._logger.warning("replace_table requires 'table_index', 'table_header', or 'match'")
        return None, None

    def replace_table(self, doc: Document, op: Dict) -> bool:
        """Replace an entire table's <w:tbl> element with new table XML.

        Unlike replace_table_cell (which only edits cell text), this swaps the
        whole table, so the replacement may have a completely different shape,
        orientation, or docxtpl loop tags.

        Config keys:
            - replace / replace_file: the new <w:tbl> XML (replace_file resolved
              by the config loader)
            - table_index / table_header / match: how to locate the table
            - header_row: header row index for table_header matching (default 0)
        """
        try:
            new_xml = op.get('replace')
            if not new_xml:
                self._logger.warning("replace_table: no replacement XML provided")
                return False

            target_table, target_index = self._find_table_for_replace(doc, op)
            if target_table is None:
                return False

            # Parse the replacement XML; if prefixes aren't declared, inject a
            # standard set and retry once.
            try:
                new_tbl = parse_xml(new_xml)
            except Exception:
                if 'xmlns:w' not in new_xml:
                    patched = re.sub(r'<w:tbl\b', f'<w:tbl {self._WORD_NSDECLS}', new_xml, count=1)
                    new_tbl = parse_xml(patched)
                else:
                    raise

            if not new_tbl.tag.endswith('}tbl'):
                self._logger.warning(f"replace_table: replacement XML root is <{new_tbl.tag}>, expected <w:tbl>")
                return False

            old_tbl = target_table._tbl
            parent = old_tbl.getparent()
            if parent is None:
                self._logger.warning(f"replace_table: table {target_index} has no parent element")
                return False

            parent.replace(old_tbl, new_tbl)
            self._logger.info(f"Replaced entire table {target_index} with new <w:tbl>")
            return True

        except Exception as e:
            self._logger.error(f"Error replacing table: {e}")
            return False

    @staticmethod
    def _row_text(tr) -> str:
        """Normalized visible text of a <w:tr> row, for comparing rows.

        Collapses all runs of whitespace (including non-breaking spaces) to a
        single space so trivial spacer-cell differences between an original
        table and its continuation copies don't defeat duplicate detection.
        """
        joined = ' '.join(t.text or '' for t in tr.iter(qn('w:t')))
        return ' '.join(joined.split())

    @staticmethod
    def _paragraph_is_empty(el) -> bool:
        """True if a <w:p> has no visible text (page breaks / empty runs only)."""
        return el.tag == qn('w:p') and not ''.join(
            t.text or '' for t in el.iter(qn('w:t'))).strip()

    def _find_all_tables_for_merge(self, doc: Document, op: Dict):
        """Return every table matching the locator, in body order.

        Uses the same header-row / substring matching as replace_table, but
        collects all matches instead of the first. table_index is rejected
        because merging needs two or more tables to combine.
        """
        table_header = op.get('table_header')
        header_row_index = op.get('header_row', 0)
        match_text = op.get('match')
        matches = []

        for table in doc.tables:
            if table_header is not None:
                if len(table.rows) <= header_row_index:
                    continue
                row = table.rows[header_row_index]
                header_tab = '\t'.join(c.text.strip() for c in row.cells)
                header_comma = ', '.join(c.text.strip() for c in row.cells)
                header_space = ' '.join(c.text.strip() for c in row.cells)
                if (table_header == header_tab or table_header == header_comma
                        or table_header == header_space or table_header in header_space):
                    matches.append(table)
            elif match_text is not None:
                full_text = '\n'.join(c.text for r in table.rows for c in r.cells)
                if match_text in full_text:
                    matches.append(table)

        return matches

    def merge_tables(self, doc: Document, op: Dict) -> bool:
        """Merge consecutive tables that share a header into a single table.

        Documents rendered from split templates often repeat the same table
        (identical title + header block) several times, one continuation per
        page. This op locates every table matching the locator, keeps the
        first, and appends each later table's data rows to it — dropping the
        duplicated leading header rows so the result is one continuous table
        with no repeated rows. The now-empty continuation tables and the blank
        (page-break) paragraphs that separated them are removed.

        Config keys:
            - table_header / match: locate the tables to merge (same matching
              as replace_table). table_index is not accepted.
            - header_row: header row index for table_header matching (default 0).
            - skip_rows: rows to drop from the front of each continuation table.
              If omitted, the identical leading rows (compared against the first
              table) are auto-detected and dropped.

        Idempotent: once merged, only one table matches, so a re-run is a no-op.
        """
        try:
            tables = self._find_all_tables_for_merge(doc, op)
            if len(tables) < 2:
                self._logger.debug("merge_tables: fewer than 2 matching tables; nothing to merge")
                return False

            target_tbl = tables[0]._tbl
            body = target_tbl.getparent()
            if body is None:
                self._logger.warning("merge_tables: target table has no parent")
                return False

            skip_rows = op.get('skip_rows')
            target_texts = [self._row_text(tr) for tr in target_tbl.findall(qn('w:tr'))]

            # Position of every relevant element up front (indices are stable
            # because we only remove afterwards).
            children = list(body)
            consumed = {id(t._tbl) for t in tables[1:]}
            first_pos = children.index(target_tbl)
            last_pos = children.index(tables[-1]._tbl)

            # Move data rows from each continuation table into the first table.
            for other in tables[1:]:
                other_tbl = other._tbl
                rows = other_tbl.findall(qn('w:tr'))
                if skip_rows is not None:
                    n_skip = skip_rows
                else:
                    n_skip = 0
                    for i, tr in enumerate(rows):
                        if i < len(target_texts) and self._row_text(tr) == target_texts[i]:
                            n_skip += 1
                        else:
                            break
                for tr in rows[n_skip:]:
                    target_tbl.append(tr)  # moves the element out of other_tbl

            # Remove the emptied continuation tables and the blank separator
            # paragraphs between them (leave any real content untouched).
            removed = 0
            for pos in range(first_pos + 1, last_pos + 1):
                el = children[pos]
                if id(el) in consumed or self._paragraph_is_empty(el):
                    el.getparent().remove(el)
                    removed += 1

            self._logger.info(
                f"Merged {len(tables)} tables into one; removed {removed} intervening element(s)")
            return True

        except Exception as e:
            self._logger.error(f"Error merging tables: {e}")
            return False

    def replace_block(self, doc: Document, op: Dict) -> bool:
        """Remove a contiguous range of body-level elements (paragraphs and/or
        tables) delimited by text anchors, optionally inserting new XML in their
        place.

        Unlike insert_block (which only adds) or replace_table (which swaps a
        single <w:tbl>), this deletes an arbitrary run of body-level siblings
        located by a 'from' and a 'to' anchor — so it can drop a whole
        sub-section (headings + equations + a table) or swap it for new content,
        while leaving each document's own numbering/layout outside the range
        untouched.

        Config keys:
            - from: text of the first anchor (a body paragraph; exact-stripped
              match preferred, else the first paragraph containing the text).
            - to: text identifying the last anchor at or after 'from' (paragraph
              OR table — the first body element whose text contains it).
            - keep_from / keep_to: if true, that anchor element is left in place
              and the removed range starts after / ends before it (default
              false: the anchor elements are themselves removed).
            - replace / replace_file: optional XML wrapped in a single root
              element (e.g. <block> ... </block>); its children are inserted
              where the removed range began, and the wrapper is discarded. Omit
              for a pure deletion. Standard Word namespace prefixes are injected
              if the root doesn't declare them.
            - skip_if_present: optional; if this text already appears anywhere in
              the body, the op is skipped. (A run is also naturally idempotent
              when the edit removes the 'to' text, so the anchor can't be found
              on a re-run.)

        Runs after replace_table (so a freshly-swapped table can bound a range)
        and before insert_block / landscape_table.
        """
        try:
            from_text = op.get('from', '')
            to_text = op.get('to', '')
            body = doc.element.body

            skip_text = op.get('skip_if_present')
            if skip_text:
                body_text = "".join(t.text or "" for t in body.iter(qn('w:t')))
                if skip_text in body_text:
                    self._logger.debug(
                        f"replace_block: '{skip_text}' already present; skipping")
                    return False

            children = list(body)

            # Locate the 'from' paragraph (index into children): prefer an exact
            # stripped text match, fall back to the first paragraph containing it.
            ft = from_text.strip()
            i_from = i_contains = None
            for k, el in enumerate(children):
                if el.tag != qn('w:p'):
                    continue
                t = "".join(x.text or "" for x in el.iter(qn('w:t'))).strip()
                if t == ft and i_from is None:
                    i_from = k
                elif ft in t and i_contains is None:
                    i_contains = k
            if i_from is None:
                i_from = i_contains
            if i_from is None:
                self._logger.warning(
                    f"replace_block: no paragraph matching 'from' = '{from_text}'")
                return False

            # Locate the 'to' element (paragraph or table) at or after 'from'.
            i_to = None
            for j in range(i_from, len(children)):
                el_text = "".join(t.text or "" for t in children[j].iter(qn('w:t')))
                if to_text in el_text:
                    i_to = j
                    break
            if i_to is None:
                self._logger.warning(
                    f"replace_block: no element containing 'to' = '{to_text}' "
                    f"at or after 'from'")
                return False

            keep_from = bool(op.get('keep_from'))
            keep_to = bool(op.get('keep_to'))
            start = i_from + (1 if keep_from else 0)
            end = i_to - (1 if keep_to else 0)
            to_remove = children[start:end + 1] if end >= start else []

            # Reference element the removed range began at (new content is
            # inserted immediately before it).
            if to_remove:
                ref = to_remove[0]
            elif end + 1 < len(children):
                ref = children[end + 1]
            else:
                ref = None

            new_children = []
            new_xml = op.get('replace')
            if new_xml:
                try:
                    root = parse_xml(new_xml)
                except Exception:
                    if 'xmlns:w' not in new_xml:
                        patched = re.sub(r'^\s*<(\w+)\b', rf'<\1 {self._WORD_NSDECLS}', new_xml, count=1)
                        root = parse_xml(patched)
                    else:
                        raise
                new_children = list(root)

            if not to_remove and not new_children:
                self._logger.debug(
                    "replace_block: empty range and no replacement; nothing to do")
                return False

            if new_children:
                if ref is not None:
                    for child in new_children:
                        ref.addprevious(child)
                else:
                    for child in new_children:
                        body.append(child)

            for el in to_remove:
                body.remove(el)

            self._logger.info(
                f"replace_block: removed {len(to_remove)} element(s) from "
                f"'{from_text}' to '{to_text}'"
                + (f", inserted {len(new_children)}" if new_children else ""))
            return True

        except Exception as e:
            self._logger.error(f"Error applying replace_block: {e}")
            return False

    def insert_block(self, doc: Document, op: Dict) -> bool:
        """Insert new body-level content (paragraphs and/or tables) at an anchor
        paragraph located by text.

        Unlike replace_table (which swaps an existing <w:tbl>), this adds brand
        new content, so it can introduce a section that didn't exist before —
        e.g. a new raw-data appendix.

        Config keys:
            - before / after: text of the anchor paragraph (exactly one). The
              block is inserted immediately before / after that paragraph.
            - replace / replace_file: the XML to insert. Several top-level
              elements (paragraphs, tables) must be wrapped in a single root
              element (e.g. <block> ... </block>); the root's children are
              inserted in order and the wrapper itself is discarded. Standard
              Word namespace prefixes are injected if the root doesn't declare
              them.
            - skip_if_present: optional text; if it already appears anywhere in
              the document body, the insert is skipped (idempotent re-runs).
        """
        try:
            new_xml = op.get('replace')
            if not new_xml:
                self._logger.warning("insert_block: no XML provided")
                return False

            insert_after = 'after' in op
            anchor_text = op.get('after') if insert_after else op.get('before')

            body = doc.element.body

            skip_text = op.get('skip_if_present')
            if skip_text:
                body_text = "".join(t.text or "" for t in body.iter(qn('w:t')))
                if skip_text in body_text:
                    self._logger.debug(
                        f"insert_block: '{skip_text}' already present; skipping")
                    return False

            anchor = self._find_paragraph_by_text(body, anchor_text)
            if anchor is None:
                self._logger.warning(
                    f"insert_block: no paragraph matching '{anchor_text}'")
                return False

            # Parse the wrapper; inject namespace declarations if the root
            # doesn't declare them (hand-written fragments using w:/w14:/... ).
            try:
                root = parse_xml(new_xml)
            except Exception:
                if 'xmlns:w' not in new_xml:
                    patched = re.sub(r'^\s*<(\w+)\b', rf'<\1 {self._WORD_NSDECLS}', new_xml, count=1)
                    root = parse_xml(patched)
                else:
                    raise

            children = list(root)
            if not children:
                self._logger.warning("insert_block: wrapper element has no children to insert")
                return False

            if insert_after:
                # addnext reverses order, so walk children back to front.
                for child in reversed(children):
                    anchor.addnext(child)
            else:
                for child in children:
                    anchor.addprevious(child)

            self._logger.info(
                f"Inserted {len(children)} element(s) "
                f"{'after' if insert_after else 'before'} '{anchor_text}'")
            return True

        except Exception as e:
            self._logger.error(f"Error applying insert_block: {e}")
            return False

    def remove_page_break(self, doc: Document, op: Dict) -> bool:
        """Remove page break(s) from the paragraph located by text.

        Strips every ``<w:br w:type="page"/>`` in the matched paragraph (and drops
        the run if that leaves it empty). Operates on the element tree, so it is
        robust to XML whitespace/serialization — unlike a literal xml_replace.

        ``<w:lastRenderedPageBreak/>`` (a render hint, not a real break) is left
        untouched.

        Config keys:
            - in_paragraph: text identifying the paragraph (exact stripped match
              preferred, falls back to the first paragraph containing the text).
        """
        try:
            text = op.get('in_paragraph')
            body = doc.element.body
            target = self._find_paragraph_by_text(body, text)
            if target is None:
                self._logger.warning(
                    f"remove_page_break: no paragraph matching '{text}'")
                return False

            removed = False
            for run in list(target.findall(qn('w:r'))):
                run_changed = False
                for br in run.findall(qn('w:br')):
                    if br.get(qn('w:type')) == 'page':
                        run.remove(br)
                        run_changed = True
                        removed = True
                # Drop the run if stripping its break left it empty.
                if run_changed and len(run) == 0 and not (run.text or '').strip():
                    target.remove(run)

            if removed:
                self._logger.info(f"Removed page break from paragraph '{text}'")
            else:
                self._logger.debug(f"remove_page_break: no page break in '{text}'")
            return removed

        except Exception as e:
            self._logger.error(f"Error applying remove_page_break: {e}")
            return False

    def landscape_table(self, doc: Document, op: Dict) -> bool:
        """Ensure the located table sits in a landscape section.

        If the table is already in a landscape section, only its section
        margins are adjusted (no redundant section is added). Otherwise the
        table is wrapped in its own landscape section: a section break is
        inserted just before it (cloning the governing section's properties,
        kept as-is) and a landscape section break just after it, so only the
        table's section is rotated and surrounding content is untouched.

        Config keys:
            - table_index / table_header / match: how to locate the table
            - header_row: header row index for table_header matching (default 0)
            - margins: optional margins for the landscape section, as
              "top,bottom,left,right" inches or a dict; default 0.5" all round
        """
        import copy

        try:
            target_table, target_index = self._find_table_for_replace(doc, op)
            if target_table is None:
                return False

            tbl = target_table._tbl
            body = doc.element.body
            if tbl.getparent() is not body:
                self._logger.warning(
                    f"landscape_table: table {target_index} is not a direct child of the "
                    "document body (nested tables are not supported)")
                return False

            margins = self._landscape_margins(op)
            governing = self._governing_sectPr(doc, tbl)
            pgSz = governing.find(qn('w:pgSz'))
            already_landscape = pgSz is not None and pgSz.get(qn('w:orient')) == 'landscape'

            if already_landscape:
                # The table already lives in a landscape section — don't add a
                # redundant section, just apply the requested margins there.
                changed = self._set_section_margins(governing, margins)
                if changed:
                    self._logger.info(
                        f"Table {target_index} already landscape; updated section margins")
                else:
                    self._logger.debug(
                        f"Table {target_index} already landscape with requested margins")
                return changed

            # Portrait: wrap the table in its own landscape island. The section
            # before keeps the governing (portrait) properties; the one after
            # carries the landscape properties for the table's section.
            before_sectPr = copy.deepcopy(governing)
            land_sectPr = copy.deepcopy(governing)
            self._make_sectPr_landscape(land_sectPr, margins)

            tbl.addprevious(self._wrap_sectPr_in_paragraph(before_sectPr))
            tbl.addnext(self._wrap_sectPr_in_paragraph(land_sectPr))

            self._logger.info(f"Wrapped table {target_index} in a landscape section")
            return True

        except Exception as e:
            self._logger.error(f"Error applying landscape_table: {e}")
            return False

    def format_table(self, doc: Document, op: Dict) -> bool:
        """Format a located table in place: tighten cell margins and/or set
        the alignment of every cell's text.

        Unlike align_table_cells (which matches individual cell text patterns
        anywhere in the document), this targets one table and applies to ALL
        its cells — useful for giving the small in-template tables (e.g. O2 /
        THC raw-data) the same tight, left-justified look as the swapped FTIR
        fragments.

        Config keys:
            - table_index / table_header / match: how to locate the table
            - header_row: header row index for table_header matching (default 0)
            - cell_margins: table cell margins in twips. An int (or one-value
              string) sets left=right=N with top=bottom=0; a "top,bottom,left,
              right" string sets all four. Omit to leave margins untouched.
            - align: 'left' | 'center' | 'right' | 'justify', applied to every
              cell paragraph. Omit to leave alignment untouched.
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        try:
            target_table, target_index = self._find_table_for_replace(doc, op)
            if target_table is None:
                return False

            changed = False

            margins = op.get('cell_margins')
            if margins is not None and self._set_table_cell_margins(target_table._tbl, margins):
                changed = True

            align = op.get('align')
            if align:
                amap = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                alignment = amap.get(str(align).lower())
                if alignment is None:
                    self._logger.warning(f"format_table: unknown align '{align}'")
                else:
                    for row in target_table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = alignment
                    changed = True

            if changed:
                self._logger.info(
                    f"Formatted table {target_index} (cell_margins={margins}, align={align})")
            return changed

        except Exception as e:
            self._logger.error(f"Error applying format_table: {e}")
            return False

    def _set_table_cell_margins(self, tbl, margins) -> bool:
        """Set the table-level default cell margins (<w:tblCellMar>) on a <w:tbl>.

        margins accepts an int / one-value string (left=right=N twips, top=
        bottom=0) or a "top,bottom,left,right" twips string. Returns True if
        applied. tblCellMar is (re)inserted in schema order — immediately
        before <w:tblLook> when present.
        """
        if isinstance(margins, str):
            parts = [p.strip() for p in margins.split(',') if p.strip() != '']
            if len(parts) == 1:
                left = right = int(parts[0]); top = bottom = 0
            elif len(parts) == 4:
                top, bottom, left, right = (int(p) for p in parts)
            else:
                self._logger.warning(
                    f"format_table: cell_margins '{margins}' must be 1 or 4 twip values")
                return False
        else:
            left = right = int(margins); top = bottom = 0

        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        existing = tblPr.find(qn('w:tblCellMar'))
        if existing is not None:
            tblPr.remove(existing)

        cellMar = OxmlElement('w:tblCellMar')
        for tag, val in (('w:top', top), ('w:left', left), ('w:bottom', bottom), ('w:right', right)):
            m = OxmlElement(tag)
            m.set(qn('w:w'), str(val))
            m.set(qn('w:type'), 'dxa')
            cellMar.append(m)

        tblLook = tblPr.find(qn('w:tblLook'))
        if tblLook is not None:
            tblLook.addprevious(cellMar)
        else:
            tblPr.append(cellMar)
        return True

    def _governing_sectPr(self, doc: Document, tbl):
        """Return the sectPr governing the table's position: the next in-body
        paragraph-level sectPr after the table, else the body's final sectPr."""
        el = tbl.getnext()
        while el is not None:
            if el.tag == qn('w:p'):
                pPr = el.find(qn('w:pPr'))
                if pPr is not None:
                    inner = pPr.find(qn('w:sectPr'))
                    if inner is not None:
                        return inner
            el = el.getnext()
        return doc.sections[-1]._sectPr

    @staticmethod
    def _wrap_sectPr_in_paragraph(sectPr):
        """Build an empty <w:p> whose pPr carries the given sectPr (a section
        break attaches to the last paragraph of the section it ends)."""
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pPr.append(sectPr)
        p.append(pPr)
        return p

    @staticmethod
    def _make_sectPr_landscape(sectPr, margins: Dict[str, float]) -> None:
        """Rotate a sectPr to landscape and apply margins (inches)."""
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is None:
            pgSz = OxmlElement('w:pgSz')
            sectPr.append(pgSz)
        cur_w, cur_h = pgSz.get(qn('w:w')), pgSz.get(qn('w:h'))
        if cur_w and cur_h:
            short, long = sorted((int(cur_w), int(cur_h)))
            pgSz.set(qn('w:w'), str(long))
            pgSz.set(qn('w:h'), str(short))
        pgSz.set(qn('w:orient'), 'landscape')
        DocxBulkUpdater._set_section_margins(sectPr, margins)

    @staticmethod
    def _set_section_margins(sectPr, margins: Dict[str, float]) -> bool:
        """Apply margins (inches) to a sectPr's pgMar. Returns True if changed."""
        pgMar = sectPr.find(qn('w:pgMar'))
        if pgMar is None:
            pgMar = OxmlElement('w:pgMar')
            sectPr.append(pgMar)
        changed = False
        for side, inches in margins.items():
            twips = str(int(round(inches * 1440)))
            if pgMar.get(qn('w:' + side)) != twips:
                pgMar.set(qn('w:' + side), twips)
                changed = True
        return changed

    @staticmethod
    def _landscape_margins(op: Dict) -> Dict[str, float]:
        """Resolve the landscape section margins (inches) from the op config."""
        margins = {'top': 0.5, 'bottom': 0.5, 'left': 0.5, 'right': 0.5}
        val = op.get('margins')
        if isinstance(val, dict):
            for side in margins:
                if side in val:
                    margins[side] = float(val[side])
        elif isinstance(val, str):
            parts = [p.strip() for p in val.split(',')]
            if len(parts) == 4:
                margins['top'], margins['bottom'], margins['left'], margins['right'] = (
                    float(parts[0]), float(parts[1]), float(parts[2]), float(parts[3]))
        return margins

    def section_break_before(self, doc: Document, op: Dict) -> bool:
        """Make the matched paragraph start its own section by moving the
        section break that currently follows it to immediately before it.

        Fixes templates where a heading is stranded at the tail of the previous
        section — e.g. an "O2 RAW DATA" heading left inside the landscape FTIR
        rawdata section, so it renders at the end of those pages instead of
        heading its own (portrait) page. The following section break is
        relocated before the heading, so the heading begins a new page in the
        next section's orientation and the preceding content keeps the moved
        break's orientation.

        Config keys:
            - match: text of the target paragraph (exact match preferred,
              falls back to substring)
        """
        try:
            match_text = op.get('match')
            if not match_text:
                self._logger.warning("section_break_before: 'match' is required")
                return False

            body = doc.element.body
            target = self._find_paragraph_by_text(body, match_text)
            if target is None:
                self._logger.warning(
                    f"section_break_before: no paragraph matching '{match_text}'")
                return False

            # Idempotency: already immediately preceded by a section break?
            prev = target.getprevious()
            if prev is not None and prev.tag == qn('w:p'):
                pPr = prev.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    self._logger.debug(
                        f"'{match_text}' already starts a section; skipping")
                    return False

            # The break that ends the target's current section is the next
            # paragraph-level sectPr after it.
            source_p, source_sectPr = self._next_paragraph_sectPr(target)
            if source_sectPr is None:
                self._logger.warning(
                    f"section_break_before: no following section break to move "
                    f"before '{match_text}'")
                return False

            source_sectPr.getparent().remove(source_sectPr)
            target.addprevious(self._wrap_sectPr_in_paragraph(source_sectPr))

            # Drop the paragraph that used to hold the break if it's now empty.
            if source_p is not None and not self._paragraph_has_content(source_p):
                parent = source_p.getparent()
                if parent is not None:
                    parent.remove(source_p)

            self._logger.info(f"Moved section break before '{match_text}'")
            return True

        except Exception as e:
            self._logger.error(f"Error in section_break_before: {e}")
            return False

    def divider(self, doc: Document, op: Dict) -> bool:
        """Isolate the matched paragraph on its own vertically-centered page,
        WITHOUT touching the paragraph itself.

        Used for the "APPENDIX F / O2 RAW DATA" divider. The paragraph already
        centers horizontally via its style (EMITAppendix1 has jc=center) plus
        the numbering's own indent — so this op deliberately leaves the
        paragraph's runs/pPr untouched (an earlier attempt that forced
        jc/indent on it actually *broke* the numbered-heading centering).

        It only does the "around it" structural work so the divider sits alone
        on one vertically-centered page with the real content on the next page:

          - The divider is followed by a block of docxtpl directives
            ({%p set ... %}, {%p if ... %}) that vanish on render. In the
            templates these directives live in their OWN section, bracketed by
            two section breaks — which renders as a BLANK page. So this op
            collapses every section break between the divider and the first
            real content (table / visible text) down to a single break, so the
            divider and the vanishing block share ONE section.
          - That single kept break gets <w:vAlign w:val="center"/> so the
            divider is vertically centered (vAlign is a section property).

        The result heals itself on re-run (extra breaks are removed again).
        Runs after section_break_before, which has already moved the divider
        onto its own page by relocating the preceding (landscape) break.

        Config keys:
            - match: text of the target paragraph (exact match preferred,
              falls back to substring)
        """
        import copy
        try:
            match_text = op.get('match')
            if not match_text:
                self._logger.warning("divider: 'match' is required")
                return False

            body = doc.element.body
            target = self._find_paragraph_by_text(body, match_text)
            if target is None:
                self._logger.warning(
                    f"divider: no paragraph matching '{match_text}'")
                return False

            # Scan forward from the divider, skipping "vanishing" paragraphs
            # (empty, or docxtpl directives {%...%} / {{...}} that disappear on
            # render), collecting section breaks, until the first real content
            # (a table or a paragraph with visible text).
            breaks = []
            content_el = None
            el = target.getnext()
            while el is not None:
                if el.tag == qn('w:tbl'):
                    content_el = el
                    break
                if el.tag == qn('w:p'):
                    pPr = el.find(qn('w:pPr'))
                    sp = pPr.find(qn('w:sectPr')) if pPr is not None else None
                    if sp is not None:
                        breaks.append(sp)
                    else:
                        txt = "".join(x.text or "" for x in el.iter(qn('w:t'))).strip()
                        if txt and not (txt.startswith('{%') or txt.startswith('{{')):
                            content_el = el
                            break
                el = el.getnext()

            changed = False

            if breaks:
                # Keep the break nearest the content; drop the rest so the
                # divider + vanishing block become one section (no blank page).
                keep = breaks[-1]
                for sp in breaks[:-1]:
                    sp.getparent().remove(sp)
                    changed = True
                if self._set_sectPr_valign(keep, 'center'):
                    changed = True
            else:
                # No section break before the content: create one (after the
                # vanishing block) so the divider is isolated on its own page.
                # Clone the governing section to keep page size / orientation.
                _, gov = self._next_paragraph_sectPr(target)
                if gov is None:
                    gov = body.find(qn('w:sectPr'))  # body-final sectPr
                if gov is None:
                    self._logger.warning(
                        "divider: no section found to clone for the break")
                else:
                    new_sectPr = copy.deepcopy(gov)
                    self._set_sectPr_valign(new_sectPr, 'center')
                    wrap = self._wrap_sectPr_in_paragraph(new_sectPr)
                    if content_el is not None:
                        content_el.addprevious(wrap)
                    else:
                        target.addnext(wrap)
                    changed = True

            if changed:
                self._logger.info(f"Applied divider to '{match_text}'")
            return changed

        except Exception as e:
            self._logger.error(f"Error in divider: {e}")
            return False

    @staticmethod
    def _set_sectPr_valign(sectPr, value: str) -> bool:
        """Set vertical text alignment (<w:vAlign>) on a sectPr. Returns True
        if changed."""
        vAlign = sectPr.find(qn('w:vAlign'))
        if vAlign is None:
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), value)
            # <w:vAlign> precedes <w:docGrid> in the sectPr schema order.
            docGrid = sectPr.find(qn('w:docGrid'))
            if docGrid is not None:
                docGrid.addprevious(vAlign)
            else:
                sectPr.append(vAlign)
            return True
        if vAlign.get(qn('w:val')) == value:
            return False
        vAlign.set(qn('w:val'), value)
        return True

    @staticmethod
    def _find_paragraph_by_text(body, text: str):
        """Find a body-level <w:p> by text. Prefers an exact (stripped) match,
        falls back to the first paragraph containing the text."""
        text = text.strip()
        exact = contains = None
        for el in body:
            if el.tag != qn('w:p'):
                continue
            t = "".join(x.text or "" for x in el.iter(qn('w:t'))).strip()
            if t == text and exact is None:
                exact = el
            elif text in t and contains is None:
                contains = el
        return exact if exact is not None else contains

    @staticmethod
    def _next_paragraph_sectPr(target):
        """Return (paragraph, sectPr) for the first paragraph-level section
        break after target, or (None, None)."""
        el = target.getnext()
        while el is not None:
            if el.tag == qn('w:p'):
                pPr = el.find(qn('w:pPr'))
                if pPr is not None:
                    sp = pPr.find(qn('w:sectPr'))
                    if sp is not None:
                        return el, sp
            el = el.getnext()
        return None, None

    @staticmethod
    def _paragraph_has_content(p) -> bool:
        """True if the paragraph has any run or non-whitespace text."""
        if p.find(qn('w:r')) is not None:
            return True
        return "".join(t.text or "" for t in p.iter(qn('w:t'))).strip() != ""

    def _set_cell_content(self, cell, content: str) -> None:
        """Set cell content using the established pattern from _rebuild_paragraph_basic.

        This follows the exact same pattern used in text_replacement.py for consistent
        handling of formatting tokens and font property preservation.
        """
        # Get the first paragraph
        para = cell.paragraphs[0]

        # Store original font formatting from first run if available (same as _rebuild_paragraph_basic)
        original_font_formatting = FontFormatter.get_base_font_formatting(para.runs)

        # Clear all runs (same as _rebuild_paragraph_basic)
        self.text_replacer._clear_paragraph(para)

        if not content:
            return

        # Process formatting tokens in the new text (same as _rebuild_paragraph_basic)
        text_segments = self.formatter.process_formatting_tokens(content, para)

        # Add runs with the new text and formatting (same as _rebuild_paragraph_basic)
        for text, formatting in text_segments:
            if text:  # Only create runs for non-empty text
                run = para.add_run(text)

                # Apply original formatting as base (same as _rebuild_paragraph_basic)
                FontFormatter.apply_font_properties(run, original_font_formatting)

                # Apply new formatting from tokens, but only properties that were explicitly specified
                # Filter out default False values that weren't actually specified in the formatting token
                filtered_formatting = {}
                for key, value in formatting.items():
                    # Only include properties that were explicitly set (not default False values)
                    if key in ['alignment', 'font_size', 'font_name', 'space_after', 'space_before'] or value is True:
                        filtered_formatting[key] = value
                    elif key in ['line_break_after', 'paragraph_break_after', 'page_break_after'] and value is True:
                        filtered_formatting[key] = value

                self.formatter.apply_formatting_to_run(run, filtered_formatting, para)

                # Apply paragraph-level formatting (alignment, etc.)
                self.formatter.apply_paragraph_formatting(para, formatting)

    def align_table_cells(self, doc: Document, align_config: Dict) -> bool:
        """Align table cells containing specific text patterns.

        Args:
            doc: The Document to modify
            align_config: Configuration dict with keys:
                - patterns: List of text patterns to search for in table cells
                - alignment: Alignment to apply ('left', 'center', 'right', 'justify')

        Returns:
            True if any cells were modified, False otherwise
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        patterns = align_config['patterns']
        alignment_str = align_config['alignment']

        # Map alignment string to enum
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        alignment = alignment_map[alignment_str]

        try:
            cells_modified = 0

            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            para_text = paragraph.text.strip()

                            # Check if this paragraph contains any of our target patterns
                            if any(pattern in para_text for pattern in patterns):
                                # Set alignment
                                paragraph.alignment = alignment
                                cells_modified += 1
                                self._logger.debug(f"Set {alignment_str.upper()} alignment: Table {table_idx+1}, Row {row_idx+1}, Cell {cell_idx+1}: '{para_text[:40]}...'")

            if cells_modified > 0:
                self._logger.info(f"Applied {alignment_str} alignment to {cells_modified} table cell paragraphs")
                return True
            else:
                self._logger.debug(f"No table cells found matching patterns: {patterns}")
                return False

        except Exception as e:
            self._logger.error(f"Error applying table cell alignment: {e}")
            return False

    def replace_image(self, doc: Document, image_config: Dict) -> bool:
        """Replace an image in the document with a new image file.

        Args:
            doc: The Document to modify
            image_config: Configuration dict with keys:
                - image_path: Path to the new image file (required)
                - name: Image name to match (e.g., "Picture 2") (optional)
                - alt_text: Alt text/description to match (optional)
                - index: Zero-based index of image to replace (optional, default: 0)
                - scale: Scale factor for the image (e.g., 0.5 for 50%, 2.0 for 200%) (optional)
                - center: Center the image horizontally on the page (optional, default: False)

        Returns:
            True if replacement was made, False otherwise
        """
        image_path = Path(image_config['image_path'])
        image_name = image_config.get('name')
        alt_text = image_config.get('alt_text')
        image_index = image_config.get('index', 0)
        scale = image_config.get('scale', 1.0)
        center = image_config.get('center', False)

        # Validate image file exists
        if not image_path.exists():
            self._logger.error(f"Image file not found: {image_path}")
            return False

        # Find all images in the document
        found_images = []
        for para in doc.paragraphs:
            for run in para.runs:
                drawings = run._element.findall(qn('w:drawing'))
                for drawing in drawings:
                    # Check for both inline and anchor (floating) images
                    inline = drawing.find(qn('wp:inline'))
                    anchor = drawing.find(qn('wp:anchor'))

                    image_element = inline if inline is not None else anchor
                    if image_element is not None:
                        docPr = image_element.find(qn('wp:docPr'))
                        if docPr is not None:
                            found_images.append((drawing, image_element, docPr, para, run))

        if not found_images:
            self._logger.warning("No images found in document")
            return False

        # Select target image based on criteria
        target_image = None

        if image_name:
            # Match by name
            for drawing, image_element, docPr, para, run in found_images:
                if docPr.get('name') == image_name:
                    target_image = (drawing, image_element, docPr, para, run)
                    break
            if not target_image:
                self._logger.warning(f"No image found with name '{image_name}'")
                return False

        elif alt_text:
            # Match by alt text/description
            for drawing, image_element, docPr, para, run in found_images:
                if docPr.get('descr') == alt_text:
                    target_image = (drawing, image_element, docPr, para, run)
                    break
            if not target_image:
                self._logger.warning(f"No image found with alt text '{alt_text}'")
                return False

        else:
            # Use index (default: first image)
            if image_index >= len(found_images):
                self._logger.warning(f"Image index {image_index} out of range (found {len(found_images)} images)")
                return False
            target_image = found_images[image_index]

        drawing, image_element, docPr, para, run = target_image

        # Get the blip element (contains the relationship ID to the image)
        blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        if blip is None:
            self._logger.error("Could not find blip element in image")
            return False

        # Get the relationship ID
        rel_id = blip.get(qn('r:embed'))
        if not rel_id:
            self._logger.error("Could not find relationship ID in image")
            return False

        # Get the document part and its relationships
        doc_part = doc.part

        # Get the related image part
        try:
            image_part = doc_part.related_parts[rel_id]
        except KeyError:
            self._logger.error(f"Relationship ID {rel_id} not found")
            return False

        # Read the new image
        with open(image_path, 'rb') as f:
            new_image_data = f.read()

        # Get the new image dimensions using PIL
        try:
            from PIL import Image as PILImage
            import io

            new_img = PILImage.open(io.BytesIO(new_image_data))
            new_width_px, new_height_px = new_img.size

            # Get current dimensions from the extent element
            # Extent is in EMUs (English Metric Units): 914400 EMU = 1 inch
            extent = image_element.find(qn('wp:extent'))
            if extent is not None:
                current_width_emu = int(extent.get('cx'))
                current_height_emu = int(extent.get('cy'))

                # Calculate new dimensions maintaining aspect ratio of new image
                # Use current width and calculate height based on new image's aspect ratio
                new_aspect_ratio = new_width_px / new_height_px
                new_height_emu = int(current_width_emu / new_aspect_ratio)

                # Apply scale factor if specified
                if scale != 1.0:
                    current_width_emu = int(current_width_emu * scale)
                    new_height_emu = int(new_height_emu * scale)
                    self._logger.debug(f"Applied scale factor: {scale}")

                # Update extent
                extent.set('cx', str(current_width_emu))
                extent.set('cy', str(new_height_emu))

                # Also update the extent in the graphic element if it exists
                graphic_extent = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm/{http://schemas.openxmlformats.org/drawingml/2006/main}ext')
                if graphic_extent is not None:
                    graphic_extent.set('cx', str(current_width_emu))
                    graphic_extent.set('cy', str(new_height_emu))

                self._logger.debug(f"Updated image dimensions to maintain aspect ratio: {current_width_emu} x {new_height_emu} EMU")
        except ImportError:
            self._logger.warning("PIL (Pillow) not available - cannot adjust image dimensions. Image may appear skewed.")
        except Exception as e:
            self._logger.warning(f"Could not adjust image dimensions: {e}")

        # Center the image horizontally if requested
        if center:
            anchor = drawing.find(qn('wp:anchor'))
            inline = drawing.find(qn('wp:inline'))

            if anchor is not None:
                # Floating (anchor) image - modify position in XML
                positionH = anchor.find(qn('wp:positionH'))
                if positionH is not None:
                    # Remove posOffset if it exists
                    posOffset = positionH.find(qn('wp:posOffset'))
                    if posOffset is not None:
                        positionH.remove(posOffset)

                    # Add or update align element
                    align = positionH.find(qn('wp:align'))
                    if align is None:
                        from lxml import etree
                        align = etree.SubElement(positionH, qn('wp:align'))
                        align.text = 'center'
                    else:
                        align.text = 'center'

                    # Set relativeFrom to page for centering
                    positionH.set('relativeFrom', 'page')

                    self._logger.debug("Centered floating image horizontally on page")

            elif inline is not None:
                # Inline image - convert to floating and center
                try:
                    from lxml import etree

                    # Get the inline element properties we need to preserve
                    inline_docPr = inline.find(qn('wp:docPr'))
                    inline_extent = inline.find(qn('wp:extent'))
                    inline_graphic = inline.find(qn('a:graphic'))

                    # Create a new anchor element to replace the inline
                    anchor = etree.Element(qn('wp:anchor'), nsmap=inline.nsmap)

                    # Copy basic attributes from inline
                    anchor.set('distT', '0')
                    anchor.set('distB', '0')
                    anchor.set('distL', '114300')
                    anchor.set('distR', '114300')
                    anchor.set('simplePos', '0')
                    anchor.set('relativeHeight', '251658240')
                    anchor.set('behindDoc', '0')
                    anchor.set('locked', '0')
                    anchor.set('layoutInCell', '1')
                    anchor.set('allowOverlap', '1')

                    # Add simplePos
                    simplePos = etree.SubElement(anchor, qn('wp:simplePos'))
                    simplePos.set('x', '0')
                    simplePos.set('y', '0')

                    # Add horizontal position (centered on page)
                    positionH = etree.SubElement(anchor, qn('wp:positionH'))
                    positionH.set('relativeFrom', 'page')
                    align_h = etree.SubElement(positionH, qn('wp:align'))
                    align_h.text = 'center'

                    # Add vertical position (relative to paragraph)
                    positionV = etree.SubElement(anchor, qn('wp:positionV'))
                    positionV.set('relativeFrom', 'paragraph')
                    posOffset_v = etree.SubElement(positionV, qn('wp:posOffset'))
                    posOffset_v.text = '0'

                    # Copy extent (size)
                    if inline_extent is not None:
                        anchor.append(etree.fromstring(etree.tostring(inline_extent)))

                    # Add effectExtent
                    effectExtent = etree.SubElement(anchor, qn('wp:effectExtent'))
                    effectExtent.set('l', '0')
                    effectExtent.set('t', '0')
                    effectExtent.set('r', '0')
                    effectExtent.set('b', '0')

                    # Add wrapSquare for text wrapping
                    wrapSquare = etree.SubElement(anchor, qn('wp:wrapSquare'))
                    wrapSquare.set('wrapText', 'bothSides')

                    # Copy docPr (document properties)
                    if inline_docPr is not None:
                        anchor.append(etree.fromstring(etree.tostring(inline_docPr)))

                    # Add cNvGraphicFramePr
                    cNvGraphicFramePr = etree.SubElement(anchor, qn('wp:cNvGraphicFramePr'))

                    # Copy graphic (contains the actual image reference)
                    if inline_graphic is not None:
                        anchor.append(etree.fromstring(etree.tostring(inline_graphic)))

                    # Replace inline with anchor in the drawing
                    drawing.remove(inline)
                    drawing.append(anchor)

                    # Add spacing after the paragraph to maintain layout
                    # This prevents the paragraph from collapsing when the inline image is removed
                    if para is not None:
                        from docx.shared import Pt
                        # Get the image height to use as spacing
                        if inline_extent is not None:
                            height_emu = int(inline_extent.get('cy', 0))
                            # Convert EMU to points (1 point = 12700 EMU)
                            height_pt = height_emu / 12700
                            para.paragraph_format.space_after = Pt(height_pt)
                            self._logger.debug(f"Added {height_pt:.1f}pt spacing after paragraph to maintain layout")

                    self._logger.debug("Converted inline image to floating and centered it horizontally")

                except Exception as e:
                    self._logger.warning(f"Could not convert inline image to floating: {e}. Falling back to paragraph centering.")
                    # Fallback: center the paragraph
                    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                    if para is not None:
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        self._logger.debug("Centered inline image by centering its paragraph")

        # Replace the image data in the existing image part
        # Note: We only replace the blob data, not the content type
        # The content type should match the original format for compatibility
        image_part._blob = new_image_data

        self._logger.info(f"Replaced image (name: {docPr.get('name')}) with {image_path}")
        return True

    def replace_text_in_table(self, doc: Document, replace_config: Dict) -> bool:
        """Replace text only within a specific table identified by its heading.

        Args:
            doc: The Document to modify
            replace_config: Configuration dict with keys:
                - table_heading: Header text to match for finding table (required)
                - search: Text to search for (required)
                - replace: Replacement text (required)
                - regex: Whether to use regex (optional, default: False)
                - table_index: Zero-based table index for disambiguation (optional)

        Returns:
            True if replacement was made, False otherwise
        """
        table_heading = replace_config['table_heading']
        search_text = replace_config['search']
        replace_text = replace_config['replace']
        use_regex = replace_config.get('regex', False)
        table_index_hint = replace_config.get('table_index')

        try:
            # Find the target table by heading
            target_table = None
            target_table_index = None

            matching_tables = []
            for i, table in enumerate(doc.tables):
                if len(table.rows) > 0:
                    # Check if header row matches the specified header pattern
                    header_row = table.rows[0]

                    # Try exact match first (tab-separated or comma-separated)
                    header_text_tab = '\t'.join(cell.text.strip() for cell in header_row.cells)
                    header_text_comma = ', '.join(cell.text.strip() for cell in header_row.cells)
                    header_text_space = ' '.join(cell.text.strip() for cell in header_row.cells)

                    if (table_heading == header_text_tab or
                        table_heading == header_text_comma or
                        table_heading == header_text_space or
                        table_heading in header_text_space):  # Fallback to contains for partial matches
                        matching_tables.append((i, table))

            if not matching_tables:
                self._logger.warning(f"No table found with heading matching '{table_heading}'")
                return False

            # If table_index is specified and matches one of our candidates, use it
            if table_index_hint is not None:
                for i, table in matching_tables:
                    if i == table_index_hint:
                        target_table = table
                        target_table_index = i
                        break

                if target_table is None:
                    self._logger.warning(f"Table index {table_index_hint} does not match any table with heading '{table_heading}'")
                    return False
            else:
                # Use the first matching table
                target_table_index, target_table = matching_tables[0]

                if len(matching_tables) > 1:
                    self._logger.info(f"Multiple tables found with heading '{table_heading}', using the first one (index {target_table_index})")

            # Now apply text replacement to all cells in this table
            modified = False
            cells_modified = 0

            for row_idx, row in enumerate(target_table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        # Check if this paragraph contains the search text
                        if (use_regex and re.search(search_text, paragraph.text)) or (not use_regex and search_text in paragraph.text):
                            # Create a temporary operation for this replacement
                            temp_op = {
                                'op': 'replace',
                                'search': search_text,
                                'replace': replace_text,
                                'regex': use_regex
                            }

                            # Use the text replacer with just this operation
                            temp_replacer = TextReplacer([temp_op], self.formatter)

                            if temp_replacer.replace_text_in_paragraph(paragraph):
                                cells_modified += 1
                                modified = True
                                self._logger.debug(f"Replaced in table '{table_heading}' (index {target_table_index}), row {row_idx}, cell {cell_idx}")

            if modified:
                self._logger.info(f"Replaced '{search_text}' with '{replace_text}' in {cells_modified} cell(s) of table '{table_heading}'")

            return modified

        except Exception as e:
            self._logger.error(f"Error replacing text in table: {e}")
            return False
