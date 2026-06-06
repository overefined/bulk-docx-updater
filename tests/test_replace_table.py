"""
Unit tests for whole-table replacement (replace_table).

Unlike replace_table_cell, replace_table swaps the entire <w:tbl> element, so the
replacement table may have a different shape, orientation, or docxtpl loop tags.
"""
import json
import pytest
from pathlib import Path
from docx import Document

from src.document_processor import DocxBulkUpdater
from src.config import load_operations_from_json


# Replacement table using only the w: prefix with NO xmlns declarations, to
# exercise the namespace-injection fallback.
NEW_TABLE_XML = (
    '<w:tbl>'
    '<w:tblPr><w:tblW w:w="9630" w:type="dxa"/></w:tblPr>'
    '<w:tblGrid><w:gridCol w:w="4815"/><w:gridCol w:w="4815"/></w:tblGrid>'
    '<w:tr>'
    '<w:tc><w:tcPr><w:tcW w:w="4815" w:type="dxa"/></w:tcPr>'
    '<w:p><w:r><w:t>SWAPPED TABLE</w:t></w:r></w:p></w:tc>'
    '<w:tc><w:tcPr><w:tcW w:w="4815" w:type="dxa"/></w:tcPr>'
    '<w:p><w:r><w:t>{{ ftir_qaqc.mdc_co }}</w:t></w:r></w:p></w:tc>'
    '</w:tr>'
    '</w:tbl>'
)


def _make_doc(tmp_path):
    """Two tables; the first carries a recognizable signature cell."""
    doc = Document()

    t0 = doc.add_table(rows=2, cols=3)
    t0.rows[0].cells[0].text = "Gas"
    t0.rows[0].cells[1].text = "Reading"
    t0.rows[0].cells[2].text = "MDC"
    t0.rows[1].cells[0].text = "{{ ftir_qaqc.reading1_co }}"

    t1 = doc.add_table(rows=1, cols=1)
    t1.rows[0].cells[0].text = "untouched"

    path = tmp_path / "doc.docx"
    doc.save(path)
    return path


def _table_texts(path):
    d = Document(path)
    return ['\n'.join(c.text for r in t.rows for c in r.cells) for t in d.tables]


def test_replace_table_by_match(tmp_path):
    path = _make_doc(tmp_path)
    ops = [{"op": "replace_table", "match": "reading1_co", "replace": NEW_TABLE_XML}]

    assert DocxBulkUpdater(ops).modify_docx(path) is True

    texts = _table_texts(path)
    assert any("SWAPPED TABLE" in t for t in texts)
    assert not any("reading1_co" in t for t in texts)
    # second table is left alone, and table count is preserved (1:1 swap)
    assert any("untouched" in t for t in texts)
    assert len(texts) == 2


def test_replace_table_by_index(tmp_path):
    path = _make_doc(tmp_path)
    ops = [{"op": "replace_table", "table_index": 0, "replace": NEW_TABLE_XML}]

    assert DocxBulkUpdater(ops).modify_docx(path) is True
    texts = _table_texts(path)
    assert "SWAPPED TABLE" in texts[0]
    assert "untouched" in texts[1]


def test_replace_table_by_header(tmp_path):
    path = _make_doc(tmp_path)
    ops = [{"op": "replace_table", "table_header": "Gas, Reading, MDC", "replace": NEW_TABLE_XML}]

    assert DocxBulkUpdater(ops).modify_docx(path) is True
    assert any("SWAPPED TABLE" in t for t in _table_texts(path))


def test_replace_table_no_match_is_noop(tmp_path):
    path = _make_doc(tmp_path)
    ops = [{"op": "replace_table", "match": "does-not-exist", "replace": NEW_TABLE_XML}]

    assert DocxBulkUpdater(ops).modify_docx(path) is False
    assert not any("SWAPPED TABLE" in t for t in _table_texts(path))


def test_replace_table_preserves_declared_namespaces(tmp_path):
    """A <w:tbl> that already declares its own namespaces parses without injection."""
    path = _make_doc(tmp_path)
    xml = NEW_TABLE_XML.replace(
        '<w:tbl>',
        '<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">',
    )
    ops = [{"op": "replace_table", "match": "reading1_co", "replace": xml}]

    assert DocxBulkUpdater(ops).modify_docx(path) is True
    assert any("SWAPPED TABLE" in t for t in _table_texts(path))


def test_replace_table_via_json_config_with_replace_file(tmp_path):
    path = _make_doc(tmp_path)

    xml_file = tmp_path / "new_table.xml"
    xml_file.write_text(NEW_TABLE_XML, encoding="utf-8")

    config_file = tmp_path / "config.json"
    config_file.write_text(
        json.dumps({"replace_table": {"match": "reading1_co", "replace_file": "new_table.xml"}}),
        encoding="utf-8",
    )

    operations, _settings = load_operations_from_json(config_file)
    assert operations[0]["op"] == "replace_table"
    # replace_file should be resolved to inline replace XML
    assert "replace" in operations[0]
    assert "{{ ftir_qaqc.mdc_co }}" in operations[0]["replace"]

    assert DocxBulkUpdater(operations).modify_docx(path) is True
    assert any("SWAPPED TABLE" in t for t in _table_texts(path))
